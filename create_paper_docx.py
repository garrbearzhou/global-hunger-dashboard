from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Margins: 1 inch all around ──
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Default font ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.0

IMG_DIR = '/Users/27zhou/Documents/Research Project/Images'

def add_para(text, bold=False, italic=False, size=11, align='left', space_after=0, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.0
    return p

def add_body(text, space_after=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    return p

def add_section_header(text, space_before=22):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(11)
    p.paragraph_format.line_spacing = 1.0
    return p

def add_sub_header(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    return p

def add_caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    return p

def add_image(path, width=Inches(6.5)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(path, width=width)
    return p

def shade_cell(cell, color='D9E2F3'):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None, header_color='D9E2F3'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        shade_cell(cell, header_color)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.space_before = Pt(1)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

# ═══════════════════════════════════════════════════════════
# PAGE 1 HEADER
# ═══════════════════════════════════════════════════════════
add_para('WORLD FOOD PRIZE FOUNDATION', bold=True, size=11, align='center', space_after=2)
add_para('Global Youth Institute \u2014 Global Challenge', size=11, align='center', space_after=16)
add_para('Bangladesh \u2014 Climate Change and Food Security in a Delta at Risk', bold=True, size=13, align='center', space_after=16)

add_para('Garrett Zhou', size=11, align='center', space_after=2)
add_para('Durham Academy Upper School', size=11, align='center', space_after=2)
add_para('Chapel Hill, North Carolina, United States of America', size=11, align='center', space_after=2)
add_para('Topic: Bangladesh \u2014 Climate Change  \u2022  March 27, 2026', size=11, align='center', space_after=22)

# ═══════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════
add_section_header('Introduction')
add_body('Most Americans will experience only a few natural disasters in their lifetimes \u2014 a hurricane, a tornado \u2014 and even then, the damage is often temporary. In Bangladesh, the stakes are categorically different. A single cyclone can devastate entire coastal communities, wiping out a season\u2019s harvest and contaminating freshwater sources for months. Situated in the low-lying Ganges-Brahmaputra delta, Bangladesh is uniquely exposed: its flat terrain, dense population, and dependence on rain-fed agriculture create a perfect storm of climate vulnerability that few nations can match. As climate change intensifies these threats, the consequences for Bangladesh\u2019s food security grow more severe. This paper quantifies how climate change is undermining food security in Bangladesh through statistical modeling and proposes a data-supported policy recommendation for the path forward.')

# ═══════════════════════════════════════════════════════════
# COUNTRY BACKGROUND
# ═══════════════════════════════════════════════════════════
add_section_header('Country Background')
add_body('Located in South Asia, with a land area similar to North Carolina and bordered by India, Myanmar, and the Bay of Bengal, lies Bangladesh (MapFight). Bangladesh is home to the confluence of three great rivers \u2014 the Ganges, the Brahmaputra, and the Meghna \u2014 which together form the world\u2019s largest river delta (Roome). Two-thirds of Bangladesh\u2019s population of 173 million live and depend upon the delta, and, consequently, these people also live less than 15 feet above sea level (Climate Reality Project).')
add_body('Bangladesh\u2019s GDP of $475 billion places it 35th globally \u2014 but GDP per capita ranks just 155th in the world (Worldometer). Bangladesh has a population density of 3,538 people per square mile, the tenth highest worldwide (World Bank, \u201cPopulation, Total for Bangladesh\u201d). Despite these challenges, living conditions have improved over the past few decades, with undernourishment falling from 16.4% a decade ago to about 11% today (United Nations Food Systems Summit).')
add_body('A major component of Bangladesh\u2019s development is international aid from the World Bank, USAID, and the Green Climate Fund. Bangladesh has suffered from political turmoil. Former Prime Minister Sheikh Hasina and the Awami League were ousted in mid-2024; an interim government led by Dr. Muhammad Yunus stabilized the country until the Bangladesh National Party (BNP) took power in February 2026 (\u201cGovernment of Bangladesh\u201d). Bangladesh ranked 149th out of 180 countries on Transparency International\u2019s Corruption Perceptions Index in 2024, and the BNP\u2019s rivalry with the Awami League puts existing climate policies \u2014 including the Bangladesh Delta Plan 2100 \u2014 at risk (Transparency International; Immigration and Refugee Board of Canada).')

# ═══════════════════════════════════════════════════════════
# A TYPICAL BANGLADESHI FAMILY
# ═══════════════════════════════════════════════════════════
add_section_header('A Typical Bangladeshi Family')
add_body('The typical family in Bangladesh resides in a rural area and comprises two parents and two to three children (ArcGIS). You have access to electricity, just like 99% of fellow Bangladeshis (Ritchie). After waking up, your breakfast consists of rice, vegetables, and perhaps some cereal, sourced from your own subsistence farm or the local market. Cooked food such as eggs or meat is extremely rare; when cooking occurs, it uses wood, crop waste, or straw. Only 28% of Bangladeshis cook with gas or electricity. The barriers are both economic and cultural: Bangladesh relies on volatile petroleum imports, and cooking with biomass is culturally seen as producing more flavorful food (Ritchie).')
add_body('After breakfast, you go to wash your hands \u2014 but like most rural Bangladeshis, you lack access to basic handwashing facilities with water and soap at home. Over one-third of the country lacks access to clean drinking water (\u201cDrinking Water and Sanitation\u201d). Your family, however, is relatively fortunate. Your father leaves for work in the service sector at the local market, where he makes roughly $10 a day \u2014 enough to ensure a comfortable lifestyle compared to many of your peers (\u201cAverage Salary in Bangladesh\u201d). Your mother tends your family\u2019s humble farm, watering rice, potatoes, jute, and some tropical fruits. You, like many of your neighbors, farm only for your family.')

# ═══════════════════════════════════════════════════════════
# CHALLENGE AND IMPACT
# ═══════════════════════════════════════════════════════════
add_section_header('Challenge and Impact')
add_body('The climate in Bangladesh follows a monsoon structure. Temperatures typically range from 27\u201331 \u00b0C but vary by geographic location (Jihan et al.). The country receives approximately 2,000 mm of rainfall per year, but 70\u201380% falls during the four-month monsoon season, creating an extreme cycle of drought and deluge (Bangladesh Meteorological Department). In the northeast, rainfall can exceed 5,000 mm, while the western interior receives as little as 1,200 mm \u2014 making water management simultaneously a problem of scarcity and surplus.')
add_body('Flooding is not an anomaly in Bangladesh. During the last half century, at least eight extreme flood events have affected 50% or more of the country\u2019s land area (World Bank, \u201cFlood Management\u201d). Over half the population is exposed to high flood risk \u2014 the highest proportion of any country in the world (LSE, \u201cTackling Flooding\u201d). Climate change is accelerating this pattern. Peak river flow is expected to rise 16\u201336% by 2070\u20132099, and annual precipitation is projected to increase by almost 30%, mostly during the post-monsoon season (Jihan et al.).')
add_body('Sea level rise compounds the flooding threat. Bangladesh experiences a relative sea-level rise of 4\u20137.8 mm per year, roughly double the global average (ICCCAD, \u201cRising Sea Level\u201d). By 2050, an estimated 17% of Bangladesh\u2019s territory could be submerged, resulting in a loss of 30% of agricultural land and forcing millions of coastal residents to migrate inland (\u201cBangladesh\u2019s 17% Areas Might Be Submerged\u201d).')
add_body('The rivers that sustain Bangladesh\u2019s delta originate in the Hindu Kush Himalayan glaciers that are disappearing. Between 1990 and 2020, the Ganges basin lost 21% of its glacier area, and the Brahmaputra basin lost 16%, with the overall rate of ice loss doubling since 2000 (ICIMOD, \u201cHKH Glacier Outlook\u201d). This creates a dual threat: accelerated melt is increasing peak river flows, intensifying the monsoon; while in the long term, as glaciers shrink past critical thresholds, dry-season water availability will decline sharply.')
add_body('Given a massive river delta, low elevation, and poor infrastructure, it should come as no surprise that Bangladesh is one of the most climate-vulnerable countries in the world. In 2024, the World Risk Index ranked Bangladesh as the 9th most vulnerable country worldwide to extreme weather (Climate Reality Project). The ND-GAIN Index ranks Bangladesh 174th out of 185 countries (35.5/100), and its disaster preparedness sub-score is 1.000 \u2014 the worst possible, placing it dead last at 185th out of 185 (Notre Dame Global Adaptation Initiative).')

add_sub_header('Global Regression Model')
add_body('Climate change impacts food security everywhere. By running an Ordinary Least Squares (OLS) regression between climate vulnerability and undernourishment across 131 countries, I found that climate vulnerability explains 48% of the variation in undernourishment rates.')
add_body('Using this regression, the equation Undernourishment (%) = \u221222.03 + 75.65 \u00d7 (Climate Vulnerability) was obtained. My model predicts that a country with Bangladesh\u2019s vulnerability score should have an undernourishment rate of 21%. However, Bangladesh has an actual value of 11.9%. This means that although Bangladesh\u2019s climate vulnerability would typically result in much higher hunger rates, other factors \u2014 including international aid, prior investments under Bangladesh Delta Plan (BDP) 2100, and decades of Green Revolution gains in rice self-sufficiency \u2014 are currently suppressing it.')

add_sub_header('Historical Trend Analysis (2002\u20132023)')
add_body('I examined this adaptation buffer over time, tracking the gap between predicted and actual undernourishment from 2002 to 2023 (Figure 1; Table 1).')

# ── FIGURE 1: vulnerability_graph.png ──
add_image(f'{IMG_DIR}/vulnerability_graph.png', width=Inches(6.3))
add_caption('Figure 1. Bangladesh\u2019s undernourishment and climate vulnerability from 2002 to 2023')

# ── TABLE 1: Vulnerability & Adaptation Buffer Summary ──
t1_headers = ['Phase', 'Years', 'Vulnerability', 'Predicted Under.', 'Actual Under.', 'Adaptation Buffer']
t1_rows = [
    ['Stagnation', '2002\u20132007', '~0.560', '~20.4%', '~15.5%', '~5 pp'],
    ['Collapse', '2007\u20132012', '0.562 \u2192 0.563', '~20.5%', '17.2% \u2192 18.2%', '5 \u2192 1.2 pp'],
    ['Recovery', '2012\u20132020', '0.555 \u2192 0.528', '19.9% \u2192 17.9%', '18.2% \u2192 10.9%', '1.7 \u2192 7 pp'],
    ['Warning', '2020\u20132023', '0.528 \u2192 0.568', '17.9% \u2192 20.9%', '10.9% \u2192 11.9%', '7 \u2192 9 pp'],
]
add_caption('Table 1. Bangladesh Climate Vulnerability and Adaptation Buffer, 2002\u20132023')
add_table(t1_headers, t1_rows)

add_body('Between 2002 and 2007, there were no major climate shocks, and climate vulnerability hovered around 0.560. The adaptation buffer \u2014 the difference between predicted and actual undernourishment \u2014 remained around 5 percentage points. From 2007 to 2012, following Cyclones Sidr and Aila, the buffer collapsed to 1.2 percentage points while undernourishment rose 2.8 percentage points. From 2012 to 2020, the BDP 2100 and international aid helped Bangladesh recover \u2014 undernourishment fell to 10.9% and the buffer rebuilt to 7 percentage points (FAO, \u201cSuite of Food Security Indicators\u201d). Now the warning signs are back. In 2021, vulnerability jumped from 0.528 to 0.557 \u2014 the largest single-year spike in the dataset \u2014 erasing a decade of resilience progress. The pattern is clear: vulnerability rises first, and undernourishment follows. Compounding this, in 2025, USAID cuts forced Bangladesh to suspend multiple climate development projects, including a $3.2 billion decade-long commitment (Mahmud).')

# ═══════════════════════════════════════════════════════════
# EXPLORING SOLUTIONS
# ═══════════════════════════════════════════════════════════
add_section_header('Exploring Solutions')
add_sub_header('Known Solutions')
add_body('Bangladesh is not short on ideas for fighting hunger, but it is short on a strategy that puts the right ideas together. Thirteen evidence-based solutions and policies, spanning food production, direct nutrition interventions, and damage prevention, were selected for this research (see Appendix Table 1).')
add_body('The most direct existing solution is school feeding. The WFP reaches 15 million ultra-poor Bangladeshis through fortified rice and biscuit programs, improving enrollment by 4.2% and cutting dropout rates by 7.5% (WFP, \u201cSchool Feeding Evaluation\u201d). A second set of solutions involves climate-resilient agriculture including 1) saline-tolerant varieties \u2014 BRRI dhan97 and dhan99 \u2014 that can produce an additional two tons per hectare even in heavily saline soil (CGIAR, \u201cSalt-Tolerant Rice\u201d), 2) flood-tolerant Sub1 rice leading to 55% higher profits and 15% higher rice consumption (CGIAR, \u201cFlood-Tolerant Rice\u201d), 3) Alternate Wetting and Drying (AWD) irrigation \u2014 saves 27% of water while increasing yields by 500 kg per hectare (IRRI). A third category is damage prevention including hermetic storage bags, flood forecasting, and farmer insurance.')

add_sub_header('Policy Optimization Model')
add_body('Each of these three categories is promising on its own, but the real question is how to allocate limited resources across many competing priorities for maximum impact. Bangladesh has roughly $8 billion budget per year available for climate adaptation \u2014 derived from the Bangladesh Delta Plan (BDP) 2100\u2019s annual allocation of 2.5% of GDP (\u201cBangladesh\u2019s Climate Resilience\u201d), guaranteed commitments from the World Bank, Asian Development Bank, and Green Climate Fund. With 18 million undernourished citizens each consuming roughly 600\u2013800 fewer calories per day than needed, the total caloric deficit that must be closed is approximately 4.6 trillion calories per year (FAO, \u201cSuite of Food Security Indicators\u201d).')
add_body('To answer this, I built a constrained optimization model across the 13 evidence-based policies. I converted every policy\u2019s output into a single common unit \u2014 additional people lifted out of undernourishment \u2014 using caloric equivalencies (see Appendix B.2). I modeled each policy\u2019s impact as a concave power function with diminishing returns calibrated from IRRI field trials, World Bank evaluations, and FAO scaling analyses (see Appendix B.4). Seven policy pairs received synergy bonuses for documented complementarity \u2014 for example, pairing saline-tolerant rice with AWD irrigation produces compounding gains (see Appendix B.5).')
add_body('Bangladesh\u2019s political landscape introduces constraints that a purely technical model would miss. I applied a 40% effectiveness discount to all BDP 2100-linked policies and required at least 35% of the budget to flow to BNP-aligned policies to reflect political feasibility. People lifted from undernourishment were valued at $285 per person per year in direct food cost savings (WFP, \u201cBangladesh Market Monitor\u201d).')
add_body('I ran the model as both a single-year allocation and a multi-year simulation: $8 billion per year for 5 years (one BNP term), with impact measured over 20 years. The multi-year model distinguishes between capital policies \u2014 one-time investments whose benefits persist for decades \u2014 and operational policies requiring continuous funding. Full model parameters, policy data sheets, synergy evidence, and simulation results are provided in the Appendix.')
add_body('In a single year, the model closes the entire undernourishment deficit with a 23% resilience buffer, funding 10 of 13 policies (see Table 2). The multi-year simulation reveals that all 11 capital policies can be fully funded within just two years at $11.2 billion total. Over 20 years, capital investments deliver 3.2 times more food security per dollar than operational programs. The model measures impact in the number of people fed. Feeding programs are still essential \u2014 they close the immediate caloric gap while infrastructure ramps up \u2014 but lasting food security must be built on durable investments, not temporary programs.')

# ═══════════════════════════════════════════════════════════
# MY RECOMMENDATION
# ═══════════════════════════════════════════════════════════
add_section_header('My Recommendation')
add_body('The model\u2019s optimal allocation tells a clear story: invest heavily in permanent infrastructure during years one and two, then maintain direct nutrition programs through year five. The data shows why this sequence matters and what it would mean for the 18 million Bangladeshis currently going hungry.')
add_body('In the single-year allocation model, the single largest allocation \u2014 $1.8 billion \u2014 goes to fortified school feeding. The limitation is that these benefits vanish the moment funding stops. This is why feeding must be paired with capital investments that outlast any single government\u2019s budget. Saline-tolerant rice receives $1.5 billion because it addresses the root cause of hunger along Bangladesh\u2019s 580-kilometer coastline. The model projects 45 million person-years of food security from this single investment over 20 years \u2014 the highest long-term return of any policy tested.')
add_body('Post-harvest storage ($1.2 billion) targets the most infuriating dimension of Bangladesh\u2019s hunger problem: food that has already been grown but never reaches anyone\u2019s plate. Over 11% of Bangladesh\u2019s rice harvest is currently lost to insects, moisture, and mold during storage \u2014 enough to feed approximately 2.3 million people per year. Distributing hermetic storage bags nationwide eliminates a source of food loss that is entirely preventable, producing 38 million person-years of food security at 32,014 people per dollar invested over 20 years \u2014 the highest per-dollar efficiency in the model.')

# ── TABLE 2: Single-Year Optimal Allocation ──
add_caption('Table 2. Single-Year Optimal Policy Allocation ($8.0B Budget)')
t2_headers = ['Policy', 'Allocation', 'People-Equiv', 'Tags']
t2_rows = [
    ['Fortified feeding', '$1,800M', '11,472,365', '\u2014'],
    ['Saline-tolerant rice', '$1,500M', '2,964,259', 'BDP'],
    ['Renewable irrigation', '$1,200M', '1,200,477', 'BNP'],
    ['Post-harvest storage', '$1,200M', '2,278,814', '\u2014'],
    ['AWD irrigation', '$1,000M', '1,871,215', 'BNP, BDP'],
    ['Flood-tolerant rice', '$499M', '565,935', 'BDP'],
    ['Crop diversification', '$200M', '152,220', '\u2014'],
    ['Farmer insurance', '$200M', '98,995', 'BNP'],
    ['River/canal excavation', '$200M', '21,860', 'BNP'],
    ['Tree planting', '$200M', '17,480', 'BNP'],
    ['Base subtotal', '$7,999M', '20,643,620', ''],
    ['Synergy bonus', '', '+1,519,946', ''],
    ['GRAND TOTAL', '$7,999M', '22,163,566', ''],
]
add_table(t2_headers, t2_rows)

add_body('The multi-year simulation reveals the most important strategic insight (Figure 2). All 11 capital policies can be fully funded within two years at a total cost of $11.2 billion. After year two, only feeding and insurance require continued spending of $2.8 billion per year. The total investment for the 5 years is $24.4 billion, within Bangladesh\u2019s $8 billion annual budget. During this period, the entire deficit is closed: 24.6 million people-equivalents of food security are generated annually against a target of 18 million, providing a 23% resilience buffer (Table 3). Over 20 years, the model projects 218 million person-years of food security at $112 per person-year \u2014 well below the $285 food cost benchmark (WFP, \u201cBangladesh Market Monitor\u201d).')

# ── FIGURE 2: allocation_graph.png ──
add_image(f'{IMG_DIR}/allocation_graph.png', width=Inches(6.3))
add_caption('Figure 2. 5-year cumulative investment for all 13 policies.')

add_body('The critical finding is what happens in year six, when the BNP\u2019s term ends, and operational budgets become uncertain. Impact drops 51% overnight as feeding and insurance programs lose funding. But capital investments keep producing. At year 20, capital alone still covers 32.3% of the deficit, feeding 5.8 million people from investments made 15 years earlier. This is the central argument for front-loading: a single five-year term of disciplined capital investment creates food security that persists for a generation.')

# ── TABLE 3: 20-Year Impact Trajectory ──
add_caption('Table 3. 20-Year Impact Trajectory \u2014 Capital vs. Operational Policies')
t3_headers = ['Year', 'Capital', 'Operational', 'Synergy', 'Total', '% of Deficit']
t3_rows = [
    ['1', '9,072,260', '11,571,360', '1,682,768', '22,326,387', '100%'],
    ['2', '10,289,472', '11,693,724', '2,628,528', '24,611,724', '100%'],
    ['3\u20135', '10,289,486', '11,693,724', '2,628,546', '24,611,755', '100%'],
    ['6', '9,833,229', '0', '2,156,163', '11,989,392', '66.4%'],
    ['7', '9,398,700', '0', '1,948,052', '11,346,751', '62.9%'],
    ['10', '8,214,713', '0', '1,438,207', '9,652,919', '53.5%'],
    ['15', '6,583,707', '0', '870,223', '7,453,930', '41.3%'],
    ['20', '5,296,290', '0', '528,536', '5,824,826', '32.3%'],
]
add_table(t3_headers, t3_rows)

# ═══════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════
add_section_header('Conclusion')
add_body('Bangladesh was dealt a difficult hand. A very low elevation for most of the country, coupled with the world\u2019s largest river delta and decades of political turmoil, ensures a road of challenges on the way to climate resiliency. However, the model results show that these challenges are solvable. My model demonstrates that $11.2 billion in capital investments over two years can permanently lift 18 million people out of undernourishment, and the benefits persist for a generation \u2014 5.8 million people remain food-secure at year 20 from investments made during a single government term. As we have seen in historical trends, a period of stagnation is not stability \u2014 it is accumulated fragility.')

# ═══════════════════════════════════════════════════════════
# ACKNOWLEDGEMENTS
# ═══════════════════════════════════════════════════════════
add_section_header('Acknowledgements')
add_body('I am grateful to Professor Hannah Jacobs of Duke University for her guidance and support for this project. I am also grateful to Dr. Eric Monson and Drew Keener for their help in the planning stages. Finally, I am grateful to my parents for supporting me throughout my research exploration process.')

# ═══════════════════════════════════════════════════════════
# WORKS CITED  (page break)
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_section_header('Works Cited', space_before=0)

works_cited = [
    'Ahmed, Sharif, and Humnath Bhandari. \u201cDiversifying Cropping for Sustainable Farming: Challenges and Opportunities in Bangladesh.\u201d Rice Today, International Rice Research Institute, ricetoday.irri.org/diversifying-cropping-for-sustainable-farming-challenges-and-opportunities-in-bangladesh/. Accessed 6 Mar. 2026.',
    '\u201cAverage Household Size \u2014 Bangladesh.\u201d ArcGIS, Esri, www.arcgis.com/home/item.html?id=692cee7e5a5e47dd86531ab0c6a00cff. Accessed 6 Mar. 2026.',
    '\u201cAverage Salary in Bangladesh.\u201d Wage Indicator Foundation, wage.is/bangladesh/. Accessed 6 Mar. 2026.',
    'Bangladesh Meteorological Department. \u201cMonsoon Rainfall.\u201d BMD, www.bmd.gov.bd/p/Monsoon-Rainfall. Accessed 6 Mar. 2026.',
    '\u201cBangladesh Agricultural Insurance Solutions Appraisal Technical Report.\u201d World Bank, documents.worldbank.org/en/publication/documents-reports/documentdetail/418491545057956149/bangladesh-agricultural-insurance-solutions-appraisal-technical-report. Accessed 6 Mar. 2026.',
    '\u201cBangladesh GDP (2025).\u201d Worldometer, www.worldometers.info/gdp/bangladesh-gdp/. Accessed 6 Mar. 2026.',
    '\u201cBangladesh Market Monitor \u2014 April 2024.\u201d World Food Programme, Apr. 2024, www.wfp.org/publications/wfp-bangladesh-market-monitor.',
    '\u201cBangladesh School Feeding USDA McGovern-Dole Grant (2020\u20132023) Evaluation.\u201d World Food Programme, www.wfp.org/publications/bangladesh-school-feeding-usda-mcgovern-dole-grant-2020-2023-evaluation. Accessed 6 Mar. 2026.',
    '\u201cBangladesh vs. North Carolina Size Comparison.\u201d MapFight, mapfight.xyz/compare/bd-vs-us.nc/. Accessed 6 Mar. 2026.',
    '\u201cBangladesh\u2019s Climate Resilience Hinges on Tripling Delta Plan Budget.\u201d The Daily Ittefaq, en.ittefaq.com.bd/13253/bangladesh%E2%80%99s-climate-resilience-hinges-on-tripling. Accessed 6 Mar. 2026.',
    '\u201cBD\u2019s 17% Areas Might Be Submerged, 30% of Farmland Lost by 2050.\u201d The Financial Express, 1 Oct. 2024, today.thefinancialexpress.com.bd/metro-news/bds-17pc-areas-might-be-submerged-30pc-of-farm-land-lost-by-2050-1727721385.',
    '\u201cCorruption Perceptions Index 2024.\u201d Transparency International, 11 Feb. 2025, www.transparency.org/en/cpi/2024/index/bgd.',
    '\u201cCountry Nutrition Profiles: Bangladesh.\u201d Global Nutrition Report, PATH, globalnutritionreport.org/resources/nutrition-profiles/asia/southern-asia/bangladesh/. Accessed 6 Mar. 2026.',
    '\u201cDrinking Water and Sanitation Facilities in Rural Bangladesh.\u201d The Financial Express, thefinancialexpress.com.bd/views/views/drinking-water-and-sanitation-facilities-in-rural-bangladesh-1633533122. Accessed 6 Mar. 2026.',
    '\u201cEffectiveness of Hermetic Bag Storage for On-Farm Storage of Maize and Rice.\u201d ASABE, elibrary.asabe.org/abstract.asp?JID=5&AID=47946&T=1. Accessed 6 Mar. 2026.',
    '\u201cEnergizing Finance: Clean Cooking in Bangladesh.\u201d Sustainable Energy for All, 19 Nov. 2020, www.seforall.org/publications/energizing-finance-understanding-the-landscape-2020/energizing-finance-clean-cooking-in-bangladesh.',
    '\u201cThe Extreme Poverty in Bangladesh: An Overview.\u201d BRAC University, dspace.bracu.ac.bd/xmlui/bitstream/handle/10361/12503/Title-1%20The%20Extreme%20Poverty%20in%20Bangladesh%20An%20Overview.pdf.',
    '\u201cThe Farmers Who Till the Harsh Coastal Lands of Bangladesh and the Salt-Tolerant Rice Varieties That Changed Their Story.\u201d CGIAR, www.cgiar.org/news-events/news/the-farmers-who-till-the-harsh-coastal-lands-of-bangladesh-and-the-salt-tolerant-rice-varieties-that-changed-their-story/. Accessed 6 Mar. 2026.',
    '\u201cFlood Management \u2014 Bangladesh Case Study.\u201d World Bank, www.floodmanagement.info/publications/casestudies/cs_bangladesh_full.pdf. Accessed 6 Mar. 2026.',
    '\u201cFlood-Tolerant Rice Improves Climate Resilience, Profitability, and Household Consumption in Bangladesh.\u201d CGIAR, cgspace.cgiar.org/items/b67df9fe-d69a-4fc6-85bc-8789d0d1c4a0. Accessed 6 Mar. 2026.',
    'Food and Agriculture Organization of the United Nations. \u201cPrevalence of Undernourishment (% of Population).\u201d FAOSTAT, www.fao.org/faostat/en/#data/FS. Accessed 6 Mar. 2026.',
    '\u201cGovernment of Bangladesh.\u201d Facts and Details, factsanddetails.com/south-asia/Bangladesh/Government_Justice_Military_Bangladesh/entry-8191.html. Accessed 6 Mar. 2026.',
    'Grantham Research Institute on Climate Change and the Environment. Tackling Flooding in Bangladesh in a Changing Climate. London School of Economics, 2023.',
    '\u201cHindu Kush Himalaya Glaciers Losing Ice at Double the Rate since 2000, New ICIMOD Report Confirms.\u201d ICIMOD, 21 Mar. 2026, www.icimod.org/press-releases/hindu-kush-himalaya-glaciers-losing-ice-at-double-the-rate-since-2000-new-icimod-report-confirm/.',
    'Hossain, Sajjad, and Md. Abbas. \u201cBNP Promises a Govt Answerable Directly to Citizens.\u201d The Daily Star, 7 Feb. 2026, www.thedailystar.net/news/national-election-2026/news/bnp-promises-govt-answerable-directly-citizens-4099891.',
    '\u201cHow the Climate Crisis Is Impacting Bangladesh.\u201d The Climate Reality Project, www.climaterealityproject.org/blog/how-climate-crisis-impacting-bangladesh. Accessed 6 Mar. 2026.',
    '\u201cHumanitarian Response Plan for Cyclone and Monsoon Floods in Bangladesh, June 2024 to March 2025.\u201d United Nations Bangladesh, 26 Sept. 2024, bangladesh.un.org/en/279755-humanitarian-response-plan-cyclone-and-monsoon-floods-bangladesh-june-2024-march-2025.',
    '\u201cHundred-Year Storm Tides Will Occur Every Few Decades in Bangladesh, Scientists Report.\u201d ScienceDaily, 11 Apr. 2025, www.sciencedaily.com/releases/2025/04/250411175457.htm.',
    'Immigration and Refugee Board of Canada. Bangladesh: Bangladesh Nationalist Party (BNP), Including Its Structure, Leaders, Membership. U.S. Department of Justice, 2015, www.justice.gov/sites/default/files/pages/attachments/2015/10/20/bgd105262.e.pdf.',
    'Jihan, Md. Akik Tanjil, et al. \u201cClimate Change Scenario in Bangladesh: Historical Data Analysis and Future Projection Based on CMIP6 Model.\u201d Scientific Reports, vol. 15, 2025, article 7856, doi:10.1038/s41598-024-81250-z.',
    'Mahmud, Lubaba. \u201cBangladesh at a Turning Point: Climate Commitments Amid Political Change.\u201d Spheres of Influence, 1 June 2025, spheresofinfluence.ca/bangladesh-at-a-turning-point-climate-commitments-amid-political-change/.',
    'Nahr, M. Samin Sajid. \u201cBangladesh Up 7 Notches on Economic Freedom Index, Score Unchanged.\u201d The Business Standard, 28 Feb. 2024.',
    'Notre Dame Global Adaptation Initiative. \u201cND-GAIN Country Index Data.\u201d University of Notre Dame, gain.nd.edu/our-work/country-index/download-data/. Accessed 6 Mar. 2026.',
    '\u201cPer Capita Caloric Intake in Bangladesh.\u201d The Business Standard. Accessed 6 Mar. 2026.',
    '\u201cRapid Himalayan Ice Melt Puts Bangladesh at Risk: UN.\u201d The Business Standard. Accessed 2026.',
    '\u201cRising Sea Level: Challenges Ahead for Bangladesh.\u201d International Centre for Climate Change and Development, icccad.net/blog/rising-sea-level-challenges-ahead-for-bangladesh/. Accessed 6 Mar. 2026.',
    'Ritchie, Hannah. \u201cAccess to Clean Cooking Fuels in Bangladesh Is Far Lower Than in Its Asian Neighbors.\u201d Our World in Data, 3 Oct. 2025.',
    'Roome, John. \u201cImplementing Bangladesh Delta Plan 2100: Key to Boost Economic Growth.\u201d World Bank Blogs, 9 June 2021.',
    '\u201cSaving Water: Alternate Wetting and Drying (AWD).\u201d IRRI Rice Knowledge Bank. Accessed 6 Mar. 2026.',
    '\u201cSilent Killer: Saltwater Threatens Women\u2019s Health in Coastal Bangladesh.\u201d Global Health Now, May 2024.',
    'United Nations Food Systems Summit. National Pathway: Bangladesh. United Nations Food Systems Hub, 2025.',
    'Victora, Cesar G., et al. \u201cMaternal and Child Undernutrition: Consequences for Adult Health and Human Capital.\u201d The Lancet, vol. 371, no. 9609, 2008, pp. 340\u201357.',
    '\u201cWorld Bank Supports Bangladesh in Flood Risk Reduction and Recovery.\u201d World Bank, 14 May 2025.',
    'World Bank. \u201cPopulation, Total for Bangladesh.\u201d FRED, Federal Reserve Bank of St. Louis. Accessed 6 Mar. 2026.',
]

for entry in works_cited:
    p = doc.add_paragraph()
    run = p.add_run(entry)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)

# ═══════════════════════════════════════════════════════════
# APPENDIX  (page break)
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
add_para('Appendix \u2014 Policy Optimization Model: Full Data and Methodology', bold=True, size=13, align='center', space_after=16)

# ── A ──
add_sub_header('A. Model 1: Global Regression')
add_body('Equation: Undernourishment (%) = \u221222.03 + 75.65 \u00d7 Climate Vulnerability')
add_body('Data: 131 countries with complete data for all 8 variables (2022 or the latest available year).')
add_body('Sources: ND-GAIN Country Index (University of Notre Dame); FAOSTAT Suite of Food Security Indicators (FAO); World Bank World Development Indicators; EM-DAT International Disaster Database; Our World in Data.')
add_body('Key result: R\u00b2 = 0.48. Climate vulnerability alone explains 48% of the variation in undernourishment rates across 131 countries. Bangladesh\u2019s predicted undernourishment is 21%, but actual is 11.9% \u2014 an adaptation buffer of 9.1 percentage points.')

# ── B.1 ──
add_sub_header('B. Model 2: Policy Optimization \u2014 13 Policies Tested')
add_sub_header('B.1 Policy List with Cost-Effectiveness Data')

b1_headers = ['#', 'Policy', 'Category', 'Cost Basis', 'Impact Basis', 'Max ($M)', 'b', 'Sources']
b1_rows = [
    ['1', 'Fortified rice & school feeding', 'Direct nutrition', '$0.12/child/day', 'Enrollment +4.2%, dropout \u22127.5%', '1,800', '0.80', 'WFP eval; PLOS One'],
    ['2', 'Saline-tolerant rice', 'Production', '~$5,000/farm', '+2 tons/ha in saline zones', '1,500', '0.80', 'CGIAR; PLOS One'],
    ['3', 'Post-harvest storage', 'Production', 'Recovered in 1 harvest', 'Losses 11.38% \u2192 0.92%', '1,200', '0.70', 'ASABE hermetic bag'],
    ['4', 'AWD irrigation', 'Production', '$67\u201397/ha', '+500 kg/ha; 27% water savings', '1,000', '0.70', 'IRRI Rice KB'],
    ['5', 'Flood-tolerant rice (Sub1)', 'Production', '~$5,000/farm', '+1\u20133 tons/ha; 55% profit', '1,800', '0.70', 'CGIAR/IRRI'],
    ['6', 'Renewable energy irrigation', 'Infrastructure', 'ADB $400M+', 'Eliminates volatile fuel costs', '1,200', '0.75', 'ADB'],
    ['7', 'Crop diversification', 'Production', '~$120M/5yr', '75% arable land is rice', '1,000', '0.65', 'Frontiers; MDPI'],
    ['8', 'Farmer insurance', 'Protection', '25% premium', 'Prevents land abandonment', '1,000', '0.50', 'World Bank; ADB'],
    ['9', 'River/canal excavation', 'Infrastructure', '~$1.7M/km', '20% flood duration reduction', '800', '0.60', 'USQ (BCR 4.35)'],
    ['10', 'Tree planting', 'Infrastructure', '$1.50\u201310/tree', 'Erosion control, 350K+ jobs', '900', '0.50', 'BNP manifesto'],
    ['11', 'Aquaculture expansion', 'Production', '~$12,700/ha', 'Production 6x over 25 yrs', '600', '0.75', 'BanglaJOL'],
    ['12', 'Cyclone early warning', 'Protection', 'Low \u2014 FFWC', 'Avoids $73\u201385M/flood', '500', '0.85', 'Springer 2024'],
    ['13', 'Flood forecasting', 'Protection', '$270M (B-STRONG)', 'BCR 79\u2013213', '700', '0.85', 'World Bank'],
]
add_table(b1_headers, b1_rows)

# ── B.2 ──
add_sub_header('B.2 Conversion to Common Unit (People Fed per $1M)')
add_body('All policies were converted to a common metric:')
add_body('Production policies: People fed = (tons produced \u00d7 3,600,000 kcal/ton) \u00f7 (2,393 kcal/day \u00d7 365 days/year). Example: 1 ton of rice = 3,600,000 kcal \u2192 3,600,000 \u00f7 873,445 = 4.12 people per ton per year.')
add_body('Direct nutrition policies: People fed = program headcount \u00d7 coverage factor.')
add_body('Protection policies: Marginal people fed = base production \u00d7 P(shock/year) \u00d7 P(loss prevented) \u00f7 annual kcal per person.')

# ── B.3 ──
add_sub_header('B.3 Caloric Deficit Calculation')
add_body('Undernourished population: 18,050,486 (FAO, 2022). Bangladesh extreme poverty caloric threshold: 1,805 kcal/day (BBS). National average caloric intake: 2,393 kcal/day (BBS HIES 2022). Estimated daily deficit per undernourished person: ~700 kcal. Total annual deficit: 18,050,486 \u00d7 700 \u00d7 365 = 4.612 trillion kcal.')

# ── B.4 ──
add_sub_header('B.4 Diminishing Returns Model')
add_body('Each policy\u2019s impact is modeled as a concave power function: impact(x) = (a / b) \u00d7 x^b, where x = spending in $M, a = marginal impact at first unit of spending (people fed per $M), b = diminishing returns exponent (0 < b < 1).')

b4_headers = ['b range', 'Policy type', 'Rationale']
b4_rows = [
    ['0.80\u20130.85', 'Standardized delivery (feeding, forecasting, early warning)', 'Milder drop-off at scale'],
    ['0.70\u20130.75', 'Agricultural production (rice, AWD, storage, aquaculture, solar)', 'Terrain, access, behavioral constraints'],
    ['0.60\u20130.65', 'Behavioral change (crop diversification, canal excavation)', 'Requires cultural shift'],
    ['0.50', 'Highest resistance (tree planting, farmer insurance)', 'Steepest diminishing returns'],
]
add_table(b4_headers, b4_rows)

# ── B.5 ──
add_sub_header('B.5 Synergy Terms')
add_body('Seven policy pairs received synergy bonuses. Synergy formula: S(i,j) = \u03b1 \u00d7 \u221a(f_i \u00d7 f_j) \u00d7 (Impact_i + Impact_j), where f = deployment fraction (spend / max budget), \u03b1 = synergy coefficient.')

b5_headers = ['Pair', '\u03b1', 'Evidence']
b5_rows = [
    ['Cyclone early warning \u00d7 Flood forecasting', '0.35', 'FFWC + Cyclone Preparedness Programme'],
    ['Post-harvest storage \u00d7 Crop diversification', '0.25', 'FAO value chain studies'],
    ['Flood-tolerant rice \u00d7 AWD irrigation', '0.20', 'IRRI compound adoption trials'],
    ['Saline-tolerant rice \u00d7 AWD irrigation', '0.20', 'BRRI coastal zone trials'],
    ['Farmer insurance \u00d7 Flood-tolerant rice', '0.15', 'World Bank \u2014 removes adoption barrier'],
    ['Aquaculture \u00d7 Crop diversification', '0.15', 'Dept. of Fisheries \u2014 rice-fish farming'],
    ['Tree planting \u00d7 River/canal excavation', '0.10', 'BDP2100 \u2014 riparian cover prevents re-siltation'],
]
add_table(b5_headers, b5_rows)

# ── B.6 ──
add_sub_header('B.6 Timeline Multipliers')
add_body('Single-year snapshots undervalue infrastructure. These multipliers convert year-1 impact to a 10-year-equivalent:')

b6_headers = ['Policy type', 'Multiplier', 'Rationale']
b6_rows = [
    ['Direct nutrition', '1.00', 'Immediate impact, only while funded'],
    ['Production', '1.15', 'Ramps up over 2\u20133 years, then persists'],
    ['Protection', '1.40', 'Prevents cumulative losses across ~3 shock cycles/decade'],
    ['Infrastructure', '1.50', 'Years to mature but compounds over a full decade'],
]
add_table(b6_headers, b6_rows)

# ── B.7 ──
add_sub_header('B.7 Political and Budget Constraints')

b7_headers = ['Constraint', 'Value', 'Source']
b7_rows = [
    ['BDP 2100 annual allocation', '2.5% of GDP (~$11.9B)', 'The Daily Ittefaq; BDP2100'],
    ['BDP 2100 political risk discount', '40% reduction (\u2192 $7.1B effective)', 'BNP\u2013Awami League rivalry'],
    ['International aid (guaranteed)', '~$900M\u2013$1.5B/year', 'World Bank; ADB; GCF'],
    ['BNP policy floor', '35% of budget to BNP-aligned', 'Political feasibility'],
    ['Minimum per policy', '$200M if funded', 'Prevents token allocations'],
    ['Implementation efficiency', '20% discount', 'World Bank/IMF standard'],
    ['Working annual budget', '~$8.0B', ''],
]
add_table(b7_headers, b7_rows)

# ── B.8 ──
add_sub_header('B.8 Food Cost Valuation')
add_body('People lifted from undernourishment are valued at $285 per person per year in direct food cost savings, based on the WFP Bangladesh Market Monitor\u2019s national food basket cost of BDT 2,844 per person per month (April 2024). The full recommended diet costs ~$340/year per FAO Cost of Recommended Diet estimates.')

# ── B.9 ──
add_sub_header('B.9 Himalayan Glacial Melt and Model Implications')
add_body('The Ganges and Brahmaputra rivers that form Bangladesh\u2019s delta originate in the Hindu Kush Himalayan glaciers. Between 1990 and 2020, these glaciers lost 12% of their total area and 9% of their ice reserves, with the rate of loss doubling since 2000 (ICIMOD, \u201cHKH Glacier Outlook,\u201d 2026). The Ganges basin lost 21% of its glacier area, and the Brahmaputra lost 16% over 30 years. Currently, 65% of the Brahmaputra\u2019s and 70% of the Ganges\u2019 dry-season water supply comes from glacial melt (\u201cRapid Himalayan Ice Melt\u201d).')
add_body('This trend affects the model in two ways:')
add_body('1. Near-term flood intensification: Accelerated melt increases peak monsoon river flows. Flood levels in the Brahmaputra could surge 80% by 2075. This strengthens the case for flood-tolerant rice (Sub1), flood forecasting, and river/canal excavation \u2014 all policies already in the model. The 10-year timeline multipliers for protection and infrastructure policies (1.40 and 1.50) may be conservative given this acceleration.')
add_body('2. Long-term dry-season water scarcity: As glaciers shrink past critical thresholds, dry-season irrigation water will decline. This makes AWD irrigation \u2014 which saves 27% of water \u2014 and renewable-powered irrigation pumps even more valuable over the 20-year measurement period. The model\u2019s current efficiency rankings already place AWD irrigation 3rd in 20-year per-dollar impact (28,438 person-years per $M); glacial melt trends suggest this ranking will hold or improve over time.')
add_body('No policy parameters were changed as a result of this analysis. The glacier melt trend reinforces existing allocations rather than requiring new policies. However, it provides additional justification for the model\u2019s central finding: capital infrastructure investments that persist for decades are more valuable than operational programs, because the climate threats they protect against are intensifying over time.')

# ── C ──
doc.add_page_break()
add_sub_header('C. Single-Year Model Results ($8.0B Budget)')
add_sub_header('C.1 Optimal Allocation')

c1_headers = ['Policy', 'Allocation', 'People-Equiv', 'Tags']
c1_rows = [
    ['Fortified feeding', '$1,800M', '11,472,365', '\u2014'],
    ['Saline-tolerant rice', '$1,500M', '2,964,259', 'BDP'],
    ['Renewable irrigation', '$1,200M', '1,200,477', 'BNP'],
    ['Post-harvest storage', '$1,200M', '2,278,814', '\u2014'],
    ['AWD irrigation', '$1,000M', '1,871,215', 'BNP, BDP'],
    ['Flood-tolerant rice', '$499M', '565,935', 'BDP'],
    ['Crop diversification', '$200M', '152,220', '\u2014'],
    ['Farmer insurance', '$200M', '98,995', 'BNP'],
    ['River/canal excavation', '$200M', '21,860', 'BNP'],
    ['Tree planting', '$200M', '17,480', 'BNP'],
    ['Base subtotal', '$7,999M', '20,643,620', ''],
    ['Synergy bonus', '', '+1,519,946', ''],
    ['GRAND TOTAL', '$7,999M', '22,163,566', ''],
]
add_table(c1_headers, c1_rows)

add_body('Policies NOT funded: Cyclone early warning, flood forecasting, aquaculture (marginal impact too low to compete at $8B scale).')

add_sub_header('C.2 Key Metrics')
add_body('Deficit closed: 100% (22.2M people-equiv vs 18.05M target). Resilience buffer: 23% (4.1M surplus capacity). BNP share of spending: 35.0% (meets floor exactly). Policies funded: 10 of 13. Synergy contribution: 7.4% of base impact. Dollar value of gains: $5.14B/year (at $285/person food cost).')

# ── D ──
add_sub_header('D. Multi-Year Simulation: 5-Year Investment, 20-Year Impact')
add_sub_header('D.1 Setup')
add_body('Investment period: 5 years at $8B/year = $40B available. Measurement period: 20 years. Capital policies (11 total): one-time investment, benefits persist with 2\u20138% annual decay. Operational policies (2: fortified feeding, farmer insurance): require annual funding; zero impact when unfunded.')

add_sub_header('D.2 Investment Phase')
d2_headers = ['Year', 'Spent', 'Key Allocation']
d2_rows = [
    ['1', '$8.0B', 'Full mix: feeding $1.8B, saline rice $1.5B, storage $1.2B, solar $1.2B, AWD $1.0B, flood rice $0.5B, plus 4 more'],
    ['2', '$8.0B', 'Maxes remaining capital: flood rice to $1.8B, diversification to $1.0B, trees $0.9B, canals $0.8B, forecasting, warning, aquaculture. Plus feeding $1.8B, insurance $1.0B'],
    ['3\u20135', '$2.8B/yr', 'All capital complete. Only feeding ($1.8B) + insurance ($1.0B)'],
]
add_table(d2_headers, d2_rows)

add_body('Total capital invested: $11.2B. Total operational: $13.2B. Grand total: $24.4B of $40B available.')

add_sub_header('D.3 20-Year Impact Trajectory')
d3_headers = ['Year', 'Capital', 'Operational', 'Synergy', 'Total', '% of Deficit']
d3_rows = [
    ['1', '9,072,260', '11,571,360', '1,682,768', '22,326,387', '100%'],
    ['2', '10,289,472', '11,693,724', '2,628,528', '24,611,724', '100%'],
    ['3\u20135', '10,289,486', '11,693,724', '2,628,546', '24,611,755', '100%'],
    ['6', '9,833,229', '0', '2,156,163', '11,989,392', '66.4%'],
    ['7', '9,398,700', '0', '1,948,052', '11,346,751', '62.9%'],
    ['10', '8,214,713', '0', '1,438,207', '9,652,919', '53.5%'],
    ['15', '6,583,707', '0', '870,223', '7,453,930', '41.3%'],
    ['20', '5,296,290', '0', '528,536', '5,824,826', '32.3%'],
]
add_table(d3_headers, d3_rows)

add_body('During years 1\u20135, the deficit is fully closed (100%). When funding stops in year 6, operational impact vanishes. Capital investments continue but decay slowly. By year 20, capital still covers 32.3% of the deficit \u2014 5.8 million people remain food-secure from investments made 15+ years earlier.')

add_sub_header('D.4 20-Year Efficiency Ranking')
d4_headers = ['Policy', 'Invested', '20-Year Person-Years', 'Per $M (20yr)']
d4_rows = [
    ['Post-harvest storage', '$1,200M', '38,416,726', '32,014'],
    ['Saline-tolerant rice', '$1,500M', '45,049,235', '30,033'],
    ['AWD irrigation', '$1,000M', '28,437,721', '28,438'],
    ['Renewable irrigation', '$1,200M', '20,237,890', '16,865'],
    ['Flood-tolerant rice', '$1,800M', '20,290,073', '11,272'],
    ['Fortified feeding', '$9,000M', '57,361,824', '6,374'],
    ['Crop diversification', '$1,000M', '5,441,917', '5,442'],
]
add_table(d4_headers, d4_rows)

add_sub_header('D.5 Capital vs Operational \u2014 20-Year Comparison')
d5_headers = ['', 'Capital', 'Operational']
d5_rows = [
    ['Total invested', '$11,200M (45.9%)', '$13,200M (54.1%)'],
    ['Total person-years', '160,212,798 (64.5%)', '58,346,257 (23.5%)'],
    ['Person-years per $M', '14,305', '4,420'],
    ['Synergy person-years', '29,719,576 (12.0%)', '\u2014'],
]
add_table(d5_headers, d5_rows)

add_body('Capital investments produce 3.2x more person-years per dollar than operational programs over 20 years.')

add_sub_header('D.6 Key Findings')
add_body('All 11 capital policies reach full deployment by year 2. Only $24.4B of $40B is productively deployable \u2014 surplus should fund education, healthcare, or climate resilience reserves. Capital investments deliver 127.5M person-years AFTER funding stops vs 90.3M DURING the 5-year period. When operational funding stops in year 6, impact drops 51% overnight (the \u201cvulnerability cliff\u201d). Cost per person-year across 20 years: $112, well below the $285 food cost benchmark. At year 20, capital alone still covers 32.3% of the deficit \u2014 5.8 million people remain food-secure from investments made 15+ years earlier.')

# ── E ──
add_sub_header('E. Structural Climate Vulnerability Effects')
add_body('The model measures impact in the number of people fed, but several policies also reduce Bangladesh\u2019s structural climate vulnerability \u2014 effects that compound over time but do not appear as \u201cpeople fed\u201d in the model\u2019s output.')
add_body('Saline-tolerant rice directly counteracts saltwater intrusion \u2014 one of the primary mechanisms through which climate change destroys food production. AWD irrigation ($1.0 billion) reduces water dependence during increasingly erratic monsoons, meaning that when rainfall patterns shift, farmers are not helpless \u2014 a critical advantage given that 65% of the Brahmaputra\u2019s dry-season flow originates from Himalayan glaciers that have lost 21% of their area in three decades and are melting at double the rate they were in 2000 (ICIMOD, \u201cHKH Glacier Outlook\u201d; \u201cRapid Himalayan Ice Melt\u201d). Flood-tolerant rice protects harvests during the monsoon floods that climate models project will intensify by 30% over the coming decades (Jihan et al.). River and canal excavation ($200 million) reduces flood duration by 20%, lowering the physical damage that each climate shock inflicts. Tree planting ($200 million) stabilizes riverbanks, prevents erosion, and creates a biological buffer against storm surges. Each of these effects reduces the probability and severity of the next climate shock, strengthening the case for capital over operational investments.')
add_body('The Farmers Card proposed by the BNP \u2014 bundling insurance, subsidies, loans, and fair market prices \u2014 addresses the behavioral side of climate vulnerability. When farmers know they will be compensated for flood losses, they are more likely to invest in improved seeds, better storage, and new techniques. Without insurance, a rational farmer avoids risk, sticks with traditional varieties, and farms conservatively. Insurance unlocks the adoption of every other policy in the model. This is why the synergy term between farmer insurance and flood-tolerant rice (\u03b1 = 0.15) is significant: insurance removes the financial barrier that prevents farmers from planting improved varieties they know perform better.')

# ── F ──
add_sub_header('F. Cross-Country Comparisons and Additional Climate Data')
add_body('School Feeding \u2014 India Comparison: India runs a similar midday meal scheme reaching over 100 million children, making it the largest school feeding program in the world. Bangladesh\u2019s version is effective but far smaller in scale \u2014 and it has a critical weakness: the meals stop the moment the funding does. India\u2019s program is government-funded as a permanent entitlement, demonstrating that sustained political commitment is required for operational programs to deliver long-term impact.')
add_body('AWD Irrigation \u2014 Vietnam Comparison: Vietnam adopted AWD irrigation nationally and saw dramatic reductions in both water use and methane emissions, proving that this approach works at scale. AWD irrigation takes on additional urgency in Bangladesh given the Himalayan glacial melt. A technique that saves 27% of irrigation water is not merely efficient \u2014 it is insurance against a future where the water supply itself is shrinking.')
add_body('Additional Climate Statistics (cut from body for word count):')
add_body('\u2022 Bangladesh\u2019s climate follows a traditional monsoon structure with four seasons: winter, pre-monsoon, monsoon (June to October), and post-monsoon. Temperatures typically range from 27\u201331 \u00b0C but can rise to 40 \u00b0C in the West (Jihan et al.).')
add_body('\u2022 Average annual rainfall: approximately 2,000 mm, but 70\u201380% falls during the four-month monsoon season. In the northeast, rainfall can exceed 5,000 mm; in the western interior, as little as 1,200 mm (Bangladesh Meteorological Department).')
add_body('\u2022 During the last half century, at least eight extreme flood events have affected 50% or more of the country\u2019s land area (World Bank, \u201cFlood Management\u201d).')
add_body('\u2022 Peak river flow is expected to rise 16\u201336% by 2070\u20132099, and annual precipitation is projected to increase by almost 30% (Jihan et al.).')
add_body('\u2022 According to the World Bank, Bangladesh\u2019s GDP from agriculture is projected to be 3.1% lower each year due to climate change \u2014 an annual loss of $570 million (Mahmud).')
add_body('\u2022 For a country where two-thirds of the population already lives less than 15 feet above sea level, even modest sea level increases can be catastrophic.')

add_body('BNP Manifesto Alignment Detail (cut from body): The Awami League, which held power for 15 years before being ousted in 2024, accelerated environmental degradation. The BNP won the 2026 election and outlined several environmental goals in their manifesto: planting 250 million trees, ensuring 20% of electricity nationwide comes from renewable sources, and creating a modern waste system. They pledged to support farmers with a Farmers Card that ensures fair prices, subsidies, loans, insurance, and state-managed markets. The BNP promised a $1 trillion economy by 2034, doubling the current market (Hossain and Abbas). Whether these promises translate to action will determine whether Bangladesh\u2019s adaptation buffer holds \u2014 or collapses again. The model\u2019s alignment with BNP priorities \u2014 tree planting, river excavation, renewable energy, and farmer insurance all appear in the optimal allocation \u2014 means that implementing these policies is not only technically optimal but politically feasible within the current government\u2019s stated agenda.')

# ═══════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════
output_path = '/Users/27zhou/Documents/Research Project/Garrett Zhou WFP Bangladesh Paper V8.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
