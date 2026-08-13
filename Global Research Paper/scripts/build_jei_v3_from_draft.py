#!/usr/bin/env python3
"""
Build a JEI-ready V3 manuscript from the student's V2 draft content.

Fixes applied for JEI pre-review readiness:
  - Restore missing R-squared / Greek symbols (plain-text safe)
  - Remove draft placeholders; embed Tables 1-4; cite figures passively
  - Strip script/file-path notes from Methods
  - JEI format: TNR 11 pt, 1.5 spacing, 1-inch margins, required section order
  - Keep student voice; only mechanical + clarity edits
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Pt

PAPER = Path(__file__).resolve().parents[1]
OUT = (
    PAPER
    / "drafts"
    / "Climate Vulnerability Does Not Equal Hunger - Measuring the Global Adaptation Buffer (V3_JEI).docx"
)

TITLE = "Climate Vulnerability Does Not Equal Hunger: Measuring the Global Adaptation Buffer"

SUMMARY = (
    "Climate vulnerability is widely used as a proxy for hunger risk, yet countries with similar "
    "vulnerability scores report very different rates of undernourishment. We introduce the "
    "adaptation buffer: the gap between hunger predicted from a cross-country regression on climate "
    "vulnerability and actual hunger rates. Using ND-GAIN, FAO, World Bank, EM-DAT, and ACLED data "
    "for 143 countries, we tested whether vulnerability is a good predictor of hunger and whether "
    "standard development indicators explain who beats and who falls short of that prediction "
    "(4, 11, 12, 13, 14, 15). We hypothesized that countries with stronger adaptive capacity would "
    "show larger positive buffers. Climate vulnerability alone explains roughly 42% of cross-country "
    "variation in hunger rates (p < 0.001), and the association survived after controlling for income, "
    "rurality, disasters, and conflict. The buffer ranged from +14.6 to -38.8 percentage points; "
    "nations such as Bangladesh (+9.1) and Senegal (+12.1) overperformed while Haiti (-38.8) "
    "underperformed. Standard covariates explained only 6% of the variation in the buffer. Monte "
    "Carlo simulation (20,000 draws) showed 66 countries were robust over-performers and 42 robust "
    "under-performers. Simulating buffer dynamics into the future, a +9.1-point buffer collapsed "
    "below zero within 20 years in 41% of trials under severe shocks, but only in 4% of trials once "
    "resilience investments reduced average shock damage and sped up recovery. The buffer is a "
    "measurable outcome-based indicator of effective climate adaptation that aid distribution "
    "frameworks could track alongside vulnerability."
)

INTRODUCTION = [
    (
        "Bangladesh and Afghanistan are two nations with similar ND-GAIN climate vulnerability scores "
        "(Bangladesh: 0.569; Afghanistan: 0.588 on a 0-1 scale) (4, 12), but the picture of "
        "undernourishment looks vastly different. Whereas Afghanistan's undernourishment rate is "
        "roughly 28.1% according to the FAO, Bangladesh's is only 10.4% (11). If climate vulnerability "
        "is a reliable indicator of undernourishment, then the two nations should have similar levels "
        "of hunger. However, Afghanistan has almost three times as much hunger as Bangladesh, even "
        "with similar climate vulnerability scores. This is not an isolated example; globally, this "
        "pattern repeats across many different pairings."
    ),
    (
        "Climate change is one of the biggest threats to global food security through yield losses, "
        "extreme events, and disrupted livelihoods (1, 2). According to FAO, extreme weather driven "
        "by climate change is among the principal drivers of undernourishment worldwide, alongside "
        "conflict and economic downturns (3). Historically, countries have been ranked by their "
        "climate vulnerability, and those rankings have been used for aid allocation plans and "
        "adaptation finance. Yet the empirical question of whether climate vulnerability maps cleanly "
        "onto hunger outcomes at the global scale has not received the direct attention that the "
        "policy use of climate vulnerability indexes would suggest."
    ),
    (
        "The primary metrics were obtained from ND-GAIN, Notre Dame's open-source decision-support "
        "tool, and FAO (4, 11, 12). ND-GAIN's Country Index is the most widely used composite measure "
        "of national climate vulnerability worldwide, while its readiness metric measures economic, "
        "governance, and social capacity to deploy adaptation investments (4, 12). Composite "
        "vulnerability indices depend on weighting and aggregation choices that can impact "
        "cross-country rankings (10). The country index ranges from 0 to 1, with higher scores "
        "indicating greater climate vulnerability. From FAO, we obtained undernourishment percentage "
        "data, which we define as the percentage of a population below the minimum dietary energy "
        "requirement (7). However, there has been a persistent gap between adaptation planning and "
        "documented outcomes, as evidenced by the adaptation-tracking literature (5, 6). Average "
        "income has become a dominant correlate of undernourishment (8), with conflict emerging as a "
        "leading driver of acute food crises (3, 9). What researchers are missing is an outcome-based "
        "measure of whether countries beat or fall short of climate-predicted hunger, one that can be "
        "easily cross-applied to other metrics."
    ),
    (
        "To the best of our knowledge, this is the first formal definition of the climate adaptation "
        "buffer: the difference between undernourishment predicted from climate vulnerability alone "
        "and observed undernourishment. To calculate it, we fit a cross-country linear regression of "
        "undernourishment on vulnerability and, for each country, calculate the difference between "
        "the predicted undernourishment value and the actual undernourishment percentage. If a "
        "country has a positive adaptation buffer, it has less hunger than its climate vulnerability "
        "alone would predict (an overperformer). A negative buffer means the country has "
        "higher-than-expected hunger (an underperformer). We devised four core research questions: "
        "(1) How strongly does vulnerability predict hunger, and does it survive controls for other "
        "factors such as income, disasters, and conflict? (2) Which countries or types of countries "
        "systematically beat or fall short of the climate-predicted benchmark? (3) Can standard "
        "development indicators (especially readiness) explain the buffer, and how uncertain are the "
        "rankings? (4) How easily does a positive buffer collapse under random shocks, and what can "
        "protect it?"
    ),
    (
        "Our hypothesis was that countries with higher ND-GAIN readiness will show larger positive "
        "adaptation buffers because they are typically more developed and have better systems for "
        "reducing undernourishment than countries with lower ND-GAIN readiness. To test our "
        "hypothesis, we used a set of 143 countries in an ordinary least squares (OLS) regression to "
        "analyze trends among our set. Then, we performed a Monte Carlo simulation with 20,000 random "
        "draws to test buffer stability. We discovered that climate vulnerability is a strong but "
        "incomplete indicator of hunger, explaining much but not all of the variation in hunger in "
        "our dataset. Fifty-eight percent of cross-country variation in hunger remained unexplained "
        "by vulnerability, but the association withstood controls for multiple covariates. We also "
        "found that buffers vary widely across countries, with an overall range of +14.6 pp to "
        "-38.8 pp, and that there is little correlation between vulnerability rankings and buffer "
        "rankings. Covariate factors do not explain the buffer well either, as only 6% of buffer "
        "variation is explained by income, readiness, disasters, and conflict. Finally, simulating a "
        "+9.1 pp buffer over twenty years produces a 41.3% chance of collapse in a baseline scenario "
        "but only a 4% chance of collapse with resilience investment."
    ),
]

RESULTS = [
    (
        "Experiment 1. First, we tested whether vulnerability was a good predictor of "
        "undernourishment. Specifically, we tested whether ND-GAIN climate vulnerability alone "
        "explained variation in cross-country undernourishment (4, 12). This step allowed us to "
        "calculate every adaptation buffer. We conducted an OLS regression using 2022 data for 143 "
        "countries to examine the relationship between ND-GAIN vulnerability and FAO undernourishment "
        "values (7, 11). We obtained an intercept of -20.04 with a standard error of 2.59 and a "
        "p-value < 0.001, and a vulnerability coefficient of 69.42 with a standard error of 7.12 and "
        "a p-value < 0.001. This coefficient means that a 0.1-point increase in a country's climate "
        "vulnerability score corresponds to roughly a 6.9 percentage-point increase in that country's "
        "undernourishment. We also obtained an R-squared of 0.423. This is a strong gradient, but 58% "
        "of the variation remains unexplained; this residual motivates the adaptation buffer. For "
        "example, Bangladesh has an ND-GAIN vulnerability score V = 0.569, which yields a predicted "
        "undernourishment value of 19.4%. However, Bangladesh's actual undernourishment is 10.4%, "
        "indicating a buffer of 9 percentage points (pp) (Table 1, Figure 1)."
    ),
    (
        "Experiment 2. Next, we tested whether the relationship between vulnerability and hunger is "
        "simply an artifact of poverty or conflict. We added ln(GDP per capita) and rural population "
        "share from World Bank World Development Indicators (13), EM-DAT disaster counts (14), and "
        "ACLED conflict fatalities (15) to the Experiment 1 specification. After running the OLS "
        "regression, we obtained a vulnerability coefficient of 40.07 with a standard error of 11.48 "
        "and a p-value < 0.001, which is lower than in Experiment 1 but still highly significant. The "
        "coefficient shrinks from about 69 to about 40, roughly a 40% decrease, meaning that roughly "
        "40% of the raw gradient co-moves with the observables (income, rurality, disasters, and "
        "conflict). Ln(GDP per capita) had a coefficient of -2.84, a standard error of 1.04, and a "
        "p-value of 0.006, indicating that higher income is associated with lower hunger. Rural share "
        "was not statistically significant (p approximately 0.29). The same applied to disaster count "
        "(p = 0.63) and conflict fatalities, specifically ln(1 + conflict fatalities) (p = 0.83). "
        "This model had an R-squared of 0.464 from 139 complete cases; four countries were excluded "
        "due to missing covariate data. We concluded that climate vulnerability survives controls and "
        "is not merely an artifact of these covariates. However, insignificant coefficients in this "
        "linear specification do not prove those factors are inconsequential; they may reflect "
        "measurement noise or functional-form limitations (Table 2)."
    ),
    (
        "Experiment 3. We then observed which countries or types of countries systematically beat or "
        "fall short of the climate-only hunger prediction. For each country, we computed the "
        "adaptation buffer as predicted minus actual undernourishment, using values from Experiment "
        "1. The buffer standard deviation was approximately 7.5 percentage points, ranging from "
        "+14.6 (Kiribati) to -38.8 (Haiti). The median buffer was around +1.0 pp, and the middle 50% "
        "ranged from -2.6 pp to +3.9 pp. Niger had the highest vulnerability score in the dataset "
        "(V = 0.636) but overperformed by 11.2 pp. By contrast, Haiti had a relatively moderate "
        "vulnerability score (V = 0.510) but underperformed by 38.8 pp, the largest underperformance "
        "by far. Overall, three core trends emerged: agrarian countries such as Bangladesh, Senegal, "
        "and Nepal tended to overperform their predicted hunger; countries affected by conflict "
        "tended to underperform; and in the tails of the distribution, vulnerability ranking was not "
        "correlated with buffer ranking (Table 3, Figure 2)."
    ),
    (
        "Experiment 4. We next tested our hypothesis that ND-GAIN readiness correlates with larger "
        "buffers, and whether the buffer is just repackaged development data. We regressed the "
        "adaptation buffer for 139 countries on ND-GAIN readiness (4, 12), ln(GDP per capita), rural "
        "share, disaster count, and conflict fatalities. We obtained an R-squared of 0.057, "
        "indicating that only about 6% of buffer variation was explained. The ND-GAIN readiness "
        "coefficient was 13.35, with a standard error of 8.69 and a p-value of 0.12. Although the "
        "result is not significant at the 5% level, readiness has the largest coefficient and "
        "directionally supports our hypothesis. None of the other variables were significant at the "
        "5% level; rural share was closest (p = 0.13). Therefore, we concluded that adaptation "
        "buffers cannot be inferred from standard indicators and must be computed directly."
    ),
    (
        "Experiment 5. One regression run is not enough to confidently classify a country as an "
        "overperformer or underperformer. We therefore ran a Monte Carlo simulation, repeating the "
        "analysis 20,000 times, each time resampling countries, refitting the regression, adding "
        "measurement noise to reported FAO hunger rates (7, 11), and recomputing all buffers. Across "
        "the 20,000 samples, the average slope was roughly 69.6, nearly the same as in Experiment 1, "
        "indicating that the main gradient is stable. In 90% of simulations, the slope ranged from "
        "58.3 to 81.5, and the intercept ranged from -24.4 to -16.0. The R-squared stayed between "
        "0.34 and 0.53, indicating that vulnerability still explains roughly 34-53% of the variation "
        "in hunger across reruns. After the simulation, each country was assigned three values: the "
        "mean buffer over 20,000 runs, the 90% interval, and the percentage of runs with a positive "
        "buffer. We found 66 robust over-performing nations where the buffer was positive in at least "
        "90% of reruns and 42 robust under-performing nations where the buffer was positive in at "
        "most 10% of reruns. Nations such as Bangladesh, Haiti, Senegal, and Niger are not statistical "
        "flukes; their classifications survive substantial uncertainty (Table 3)."
    ),
    (
        "Experiment 6. Buffers are not static; they can collapse. Between 2007 and 2009, Bangladesh "
        "experienced two major cyclones (Sidr and Aila), which were associated with a partial "
        "collapse of its buffer. To account for buffer collapse, we simulated 20,000 twenty-year "
        "trajectories of a +9.1 pp buffer under random annual shocks (+9.1 pp is Bangladesh's "
        "observed Monte Carlo mean buffer). The buffer is defined as collapsed if it falls below "
        "zero at least once during the simulation. This is a different Monte Carlo experiment from "
        "Experiment 5. Both the baseline and resilience scenarios have the same 20% annual chance of "
        "a shock; the key differences are shock size and recovery speed. In the baseline scenario, "
        "the buffer collapsed within 20 years in 41.3% of reruns and within 10 years in 21.6%. The "
        "expected number of years with a collapsed buffer was 1.01, and the median buffer at year 20 "
        "was roughly +6.6 pp. In the resilience scenario, the buffer collapsed within 20 years in "
        "only 4.4% of reruns and within 10 years in 2.0%. The expected number of years underwater was "
        "0.05, and the median buffer at year 20 was +8.2 pp. There was roughly a tenfold reduction in "
        "collapse risk with identical shock frequency, underscoring the impact of resilience "
        "investment. Defending the buffer matters even when climate exposure is unchanged "
        "(Table 4, Figure 3)."
    ),
    (
        "Robustness. Finally, we tested whether the main findings were sensitive to data choices. "
        "For R1, the vulnerability coefficient remained 68.7, with an R-squared of 0.442, when using "
        "FAO undernourishment data averaged over 2020-2022 instead of 2022-2024 (11). For R2, the "
        "gradient replicated with the Global Data Lab climate vulnerability index rather than "
        "ND-GAIN (R-squared = 0.408; N = 92) (16). Using the World Risk Index vulnerability "
        "component also produced a positive, significant gradient (17). For R3, climate "
        "vulnerability significantly predicted WHO child stunting prevalence (R-squared = 0.624), "
        "indicating that the pattern is not exclusive to FAO undernourishment data (18). For R5, "
        "after dropping nations with populations under 1 million, Bangladesh's adaptation buffer "
        "remained +10.5 and ranked sixth. These checks indicate that the core vulnerability-hunger "
        "association and Bangladesh's outperformance are robust to alternative data specifications."
    ),
]

DISCUSSION = [
    (
        "We concluded that climate vulnerability does not equate directly to hunger. Vulnerability "
        "accounts for roughly 42% of the variation in undernourishment across 143 countries, a "
        "substantial but still incomplete share. We argue that the remaining 58% is structured and "
        "policy-relevant rather than random noise. The adaptation buffer names that residual as a "
        "trackable quantity. It reveals food-security performance relative to a nation's climate "
        "burden, unlike adaptation tracking, which counts plans and investments (5, 6). Because "
        "standard development indicators explain only about 6% of buffer variation, the buffer must "
        "be monitored directly rather than inferred from income or readiness alone. A positive "
        "buffer is equivalent to beating climate odds in actual hunger outcomes."
    ),
    (
        "The largest underperformers were dominated by political violence and state fragility, not "
        "climate exposure alone. For example, Haiti's buffer is -38.8 and Syria's buffer is -25.2. "
        "These nations, along with other strong underperformers such as Liberia, the Central African "
        "Republic, and Afghanistan, face political instability and weak governance. They are not "
        "always the most climate-vulnerable countries according to ND-GAIN. Several of the most "
        "climate-exposed countries have large positive buffers, including Niger, Bangladesh, and "
        "Senegal. We speculate that conflict and governance collapse destroy food-security "
        "performance faster than climate exposure alone predicts, consistent with evidence that "
        "conflict is a leading driver of acute food crises and that climate and conflict stresses "
        "interact (3, 9). Therefore, vulnerability-weighted aid formulas may miss countries where "
        "hunger most exceeds climate fundamentals. Severe underperformers may need conflict-sensitive "
        "food security interventions rather than climate adaptation alone."
    ),
    (
        "Bangladesh is a useful prototype. Decades of agricultural innovation, disaster preparedness, "
        "and social protection helped build a +9.0 pp buffer, yet that buffer partially collapsed "
        "after the 2007-2009 cyclone sequence and has since been only partially rebuilt. In our "
        "simulations, even a healthy buffer like Bangladesh's has a 41.3% twenty-year collapse risk "
        "under baseline shocks. Resilience investments that reduce shock damage and speed recovery "
        "cut that risk to 4.4%, roughly a tenfold reduction at identical shock frequency. We "
        "speculate that high-value interventions may defend existing buffers rather than only "
        "lowering index scores or hazard rankings. Examples include early warning systems, grain "
        "storage, safety nets, and stress-tolerant crop varieties. These may seem expensive upfront, "
        "but they are often cheaper than rebuilding a food system after collapse."
    ),
    (
        "Our analysis does not prove causation; all methods are observational and cross-sectional. "
        "Countries are not randomized into vulnerability levels. The buffer is also "
        "benchmark-relative: it depends on ND-GAIN vulnerability and a bivariate regression line, "
        "and composite vulnerability indices involve aggregation and weighting choices that can "
        "affect rankings (10). A different benchmark could lead to different buffer rankings. FAO "
        "also uses three-year averages of undernourishment data to smooth acute shocks and censors "
        "values at 2.5% (7). The stunting robustness check (R3) helped, but stunting is a different "
        "nutritional construct. Smaller states such as Kiribati and Samoa may partially reflect FAO "
        "estimation uncertainty, which is why we conducted R5. For the Monte Carlo assumptions, we "
        "used a 10% measurement-error rule and illustrative shock parameters; the comparison between "
        "scenarios is more important than exact percentages. Additionally, our hypothesis was not "
        "confirmed: a readiness p-value of 0.12 at the 5% level is hypothesis-generating but not "
        "statistically significant. Finally, a linear covariate relationship between climate "
        "vulnerability and conflict may understate interaction effects."
    ),
    (
        "Future work could include annual panel tracking of buffer levels and trends, richer "
        "covariates such as agricultural productivity, aid flows, governance indices, and cereal "
        "yields, and tests of whether shrinking buffers predict subsequent hunger increases "
        "(an early-warning application). Integrating buffer level and trend into climate-adaptation "
        "monitoring alongside vulnerability indices could make countries with eroding food-security "
        "gains visible before crises fully materialize."
    ),
]

METHODS = [
    (
        "Data sources and sample. We merged publicly available country-level datasets by ISO3 code "
        "for 143 countries with non-missing ND-GAIN vulnerability (2022) and FAO prevalence of "
        "undernourishment (Item 210041, 2022-2024 three-year window; values reported as <2.5% coded "
        "at 2.5%) (7, 11, 12). ND-GAIN composite vulnerability and readiness indices came from the "
        "Notre Dame Global Adaptation Initiative (4, 12). World Bank World Development Indicators "
        "provided 2022 GDP per capita and rural population share (13). EM-DAT disaster counts were "
        "summed over 2013-2022 (14). ACLED provided three-year average conflict fatalities (15). "
        "Complete cases for multivariate models totaled N = 139."
    ),
    (
        "Buffer construction. From Experiment 1, we obtained a fitted undernourishment value for "
        "each country and calculated the adaptation buffer as predicted undernourishment minus "
        "actual undernourishment. A positive buffer means a country reported less hunger than the "
        "climate-only benchmark predicted; a negative buffer means more hunger than predicted. "
        "Buffers were computed for all 143 countries in the Experiment 1 sample."
    ),
    (
        "Regression analysis. In Experiment 1, we used OLS to regress undernourishment on climate "
        "vulnerability. In Experiment 2, we added ln(GDP per capita), rural share (%), EM-DAT "
        "disaster count, and ln(1 + ACLED conflict fatalities). In Experiment 4, we regressed the "
        "adaptation buffer on those covariates, excluding climate vulnerability because it already "
        "enters the buffer through the predicted value. Analyses were implemented in Python 3 with "
        "pandas and statsmodels. Significance was assessed at alpha = 0.05."
    ),
    (
        "Monte Carlo uncertainty. To quantify uncertainty in each country's buffer, we used Monte "
        "Carlo simulation with 20,000 iterations and random seed 42 for reproducibility. In each "
        "iteration we (1) resampled 143 countries with replacement and refitted the Experiment 1 "
        "OLS line; (2) added measurement noise to each country's FAO undernourishment percentage "
        "(7, 11), drawing from a Normal(0, sigma) distribution where sigma = max(1.0, 0.1 x "
        "undernourishment percentage); and (3) recomputed every buffer as fitted hunger from the new "
        "line minus noisy actual hunger. For each country we recorded mean buffer, the 90% interval "
        "(5th-95th percentiles), and the share of runs with buffer > 0. Countries with at least 90% "
        "positive runs were classified as robust over-performers; those with at most 10% positive "
        "runs were robust under-performers; all others were ambiguous."
    ),
    (
        "Collapse simulation. To model buffer fragility over time, we simulated 20,000 twenty-year "
        "trajectories starting from B0 = +9.1 pp (Bangladesh's Monte Carlo mean buffer). Each year "
        "the buffer recovered a fraction of the gap toward its target and could take a random shock. "
        "Recovery rates were 0.25 (baseline) and 0.40 (resilience). Each year had a 20% probability "
        "of a shock; when a shock occurred, size was drawn from a gamma distribution with mean "
        "4.0 pp (baseline) or 2.4 pp (resilience). Collapse was defined as the buffer falling below "
        "0 at least once within the 20-year horizon. We reported collapse probabilities within 10 "
        "and 20 years, expected years underwater, and the median buffer at year 20."
    ),
    (
        "Robustness checks. We assessed sensitivity to alternative data specifications, reusing "
        "Experiment 1 unless noted. R1 used the 2020-2022 FAO window instead of 2022-2024 (11). "
        "R2a replaced ND-GAIN with the World Risk Index vulnerability component rescaled from "
        "0-100 to 0-1 (17). R2b used the composite World Risk Index (17). R2c used the Global Data "
        "Lab national climate vulnerability index 2022 (0-1), which reduced the sample to N = 92 "
        "(16). R2d used the ND-GAIN exposure sub-index only. R3 replaced FAO undernourishment with "
        "WHO child stunting prevalence (latest national estimate; N = 121) (18). R5 restricted the "
        "sample to countries with population greater than 1,000,000 according to the World Bank "
        "(13) (N = 125)."
    ),
    (
        "Ethics. We used only publicly available aggregated country-level data. No human subjects "
        "or vertebrate animals were involved. Analysis scripts are available upon request."
    ),
]

REFERENCES = [
    'Wheeler, Tim, and Joachim von Braun. "Climate Change Impacts on Global Food Security." Science, vol. 341, no. 6145, 2013, pp. 508-513. https://doi.org/10.1126/science.1239402.',
    "IPCC. Climate Change 2022: Impacts, Adaptation and Vulnerability. Contribution of Working Group II to the Sixth Assessment Report of the Intergovernmental Panel on Climate Change. Cambridge University Press, 2022.",
    "FAO, IFAD, UNICEF, WFP, and WHO. The State of Food Security and Nutrition in the World 2024. Food and Agriculture Organization of the United Nations, 2024.",
    "Chen, C., et al. University of Notre Dame Global Adaptation Index: Country Index Technical Report. University of Notre Dame, 2015.",
    'Ford, James D., et al. "How to Track Adaptation to Climate Change: A Typology of Approaches for National-Level Application." Ecology and Society, vol. 18, no. 3, 2013, https://doi.org/10.5751/ES-05732-180340.',
    'Berrang-Ford, Lea, et al. "A Systematic Global Stocktake of Evidence on Human Adaptation to Climate Change." Nature Climate Change, vol. 11, no. 11, 2021, pp. 989-1000. https://doi.org/10.1038/s41558-021-01170-y.',
    'Cafiero, Carlo, Sara Viviani, and Mark Nord. "Food Security Measurement in a Global Context: The Food Insecurity Experience Scale." Measurement, vol. 116, 2018, pp. 146-152. https://doi.org/10.1016/j.measurement.2017.10.065.',
    'Headey, Derek. "Developmental Drivers of Nutritional Change: A Cross-Country Analysis." World Development, vol. 42, 2013, pp. 76-88. https://doi.org/10.1016/j.worlddev.2012.07.002.',
    'Burke, Marshall, Solomon M. Hsiang, and Edward Miguel. "Climate and Conflict." Annual Review of Economics, vol. 7, 2015, pp. 577-617. https://doi.org/10.1146/annurev-economics-080614-115430.',
    'Hinkel, Jochen. "Indicators of Vulnerability and Adaptive Capacity: Towards a Clarification of the Science-Policy Interface." Global Environmental Change, vol. 21, no. 1, 2011, pp. 198-208. https://doi.org/10.1016/j.gloenvcha.2010.08.002.',
    '"Suite of Food Security Indicators (Item 210041: Prevalence of Undernourishment)." FAOSTAT. Food and Agriculture Organization of the United Nations, www.fao.org/faostat/en/#data/FS. Accessed 10 June 2026.',
    '"ND-GAIN Country Index Data." Notre Dame Global Adaptation Initiative, gain.nd.edu/our-work/country-index/download-data/. Accessed 10 June 2026.',
    '"World Development Indicators." World Bank, databank.worldbank.org. Accessed 10 June 2026.',
    '"EM-DAT: The International Disaster Database." Centre for Research on the Epidemiology of Disasters, www.emdat.be. Accessed 10 June 2026.',
    '"Conflict Event and Fatality Data." Armed Conflict Location & Event Data Project, acleddata.com. Accessed 10 June 2026.',
    '"Climate Vulnerability Index." Global Data Lab, Radboud University, globaldatalab.org. Accessed 10 June 2026.',
    "Bundnis Entwicklung Hilft and Ruhr University Bochum - IFHV. WorldRiskReport 2025. weltrisikobericht.de. Accessed 10 June 2026.",
    "UNICEF, WHO, and World Bank. Joint Child Malnutrition Estimates: stunting prevalence in children under 5. WHO Global Health Observatory. Accessed 10 June 2026.",
]

ACK = (
    "I would like to thank the North Carolina Youth Institute and the World Food Prize Foundation "
    "for inspiring me to conduct this research project. Researching Bangladesh gave me insight into "
    "how other nations face food insecurity at both a local and national level and first sparked my "
    "interest in the adaptation buffer. I would also like to thank the Food Bank of Central and "
    "Eastern North Carolina for piquing my interest in food security issues. These interests "
    "combined to create the foundation for this research project."
)

TABLE1 = [
    ["Term", "Coefficient", "Std. error", "p-value"],
    ["Constant", "-20.04", "2.59", "<0.001"],
    ["Vulnerability", "69.42", "7.12", "<0.001"],
    ["R-squared", "0.423", "(N = 143)", ""],
]

TABLE2 = [
    ["Variable", "Coefficient", "Std. error", "p-value"],
    ["Vulnerability", "40.07", "11.48", "<0.001"],
    ["ln(GDP per capita)", "-2.84", "1.04", "0.006"],
    ["Rural share (%)", "-0.047", "0.045", "0.29"],
    ["Disaster count (2013-22)", "0.12", "0.25", "0.63"],
    ["ln(1 + conflict fatalities)", "0.065", "0.31", "0.83"],
    ["R-squared", "0.464", "(N = 139)", ""],
]

TABLE3 = [
    ["Country", "Buffer (pp)", "90% interval", "P(over-performer)"],
    ["Senegal", "+12.1", "[+9.5, +14.9]", "1.00"],
    ["Niger", "+11.2", "[+7.4, +15.2]", "1.00"],
    ["Bangladesh", "+9.1", "[+6.1, +12.1]", "1.00"],
    ["Viet Nam", "+7.1", "[+5.0, +9.3]", "1.00"],
    ["Kenya", "-22.1", "[-28.4, -15.7]", "0.00"],
    ["Haiti", "-38.8", "[-47.9, -29.6]", "0.00"],
]

TABLE4 = [
    ["Scenario", "P(collapse <= 20 y)", "Expected years underwater", "Buffer yr 20 (median, pp)"],
    ["Baseline (large shocks, slow recovery)", "41.3%", "1.01", "+6.6"],
    ["Resilience investment", "4.4%", "0.05", "+8.2"],
]

FIGURE_NOTES = {
    1: (
        "Figure 1. Relationship between ND-GAIN climate vulnerability and FAO prevalence of "
        "undernourishment across 143 countries. Each point is one country. The line shows the OLS "
        "fit (R-squared = 0.423). Bangladesh is labeled as an example of a positive adaptation "
        "buffer. Submit separately as Figure1.jpeg."
    ),
    2: (
        "Figure 2. Adaptation buffer for selected countries. The buffer equals predicted minus "
        "actual undernourishment (percentage points). Positive values indicate countries "
        "outperforming their climate-only benchmark. Submit separately as Figure2.jpeg."
    ),
    3: (
        "Figure 3. Monte Carlo simulation of buffer collapse over 20 years (20,000 trials). "
        "Starting buffer +9.1 percentage points (Bangladesh). Both scenarios face the same 20% "
        "annual shock probability; resilience investment reduces mean shock size and increases "
        "recovery rate. Submit separately as Figure3.jpeg."
    ),
}


def set_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    for level in range(1, 3):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(12 if level == 1 else 11)
        h.font.bold = True


def add_centered(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_paras(doc: Document, paras: list[str]) -> None:
    for text in paras:
        doc.add_paragraph(text)


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.bold = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            if i == 0:
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.bold = True
    doc.add_paragraph()


def add_fig_note(doc: Document, num: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(FIGURE_NOTES[num])
    run.italic = True


def build() -> Path:
    assert len(TITLE) <= 110, len(TITLE)
    assert len(SUMMARY.split()) <= 250, len(SUMMARY.split())

    doc = Document()
    set_styles(doc)

    # Title page
    add_centered(doc, TITLE, bold=True, size=12)
    doc.add_paragraph()
    add_centered(doc, "Garrett Zhou1")
    add_centered(doc, "Hannah Jacobs2")
    doc.add_paragraph()
    add_centered(doc, "1 Durham Academy Upper School, Durham, North Carolina")
    add_centered(doc, "2 Duke University, Durham, North Carolina")
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "[Paste this content into the official JEI Word template using Paste Special > "
        "Unformatted Text, then re-apply JEI Styles.]"
    )
    nr.italic = True
    doc.add_page_break()

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(SUMMARY)
    wc = doc.add_paragraph()
    wc.add_run(f"[Summary word count: {len(SUMMARY.split())} / 250]").italic = True
    doc.add_page_break()

    # Introduction
    doc.add_heading("Introduction", level=1)
    add_paras(doc, INTRODUCTION)

    # Results with tables/figure notes interleaved
    doc.add_heading("Results", level=1)
    doc.add_paragraph(RESULTS[0])
    add_table(doc, "Table 1. Experiment 1 regression of undernourishment on vulnerability.", TABLE1)
    add_fig_note(doc, 1)
    doc.add_paragraph(RESULTS[1])
    add_table(doc, "Table 2. Experiment 2 multivariate regression.", TABLE2)
    doc.add_paragraph(RESULTS[2])
    add_table(
        doc,
        "Table 3. Monte Carlo buffer estimates for selected countries (Experiments 3 and 5).",
        TABLE3,
    )
    add_fig_note(doc, 2)
    doc.add_paragraph(RESULTS[3])
    doc.add_paragraph(RESULTS[4])
    doc.add_paragraph(RESULTS[5])
    add_table(doc, "Table 4. Buffer collapse simulation results (Experiment 6).", TABLE4)
    add_fig_note(doc, 3)
    doc.add_paragraph(RESULTS[6])

    # Discussion
    doc.add_heading("Discussion", level=1)
    add_paras(doc, DISCUSSION)

    # Methods
    doc.add_heading("Materials and Methods", level=1)
    add_paras(doc, METHODS)

    # References
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFERENCES, start=1):
        doc.add_paragraph(f"{i}. {ref}")

    # Acknowledgements
    doc.add_heading("Acknowledgements", level=1)
    doc.add_paragraph(ACK)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"Written: {out}")
    print(f"Title chars: {len(TITLE)} (max 110)")
    print(f"Summary words: {len(SUMMARY.split())} (max 250)")
