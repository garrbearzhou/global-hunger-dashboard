"""Expanded section content for Global_Paper_Outline.docx — bullet-point only."""

from __future__ import annotations

from docx import Document
from docx.shared import Pt


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def h4(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for j, header in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    doc.add_paragraph()


def add_title_page(doc: Document) -> None:
    t = doc.add_paragraph()
    t.alignment = 1
    run = t.add_run("Comprehensive JEI Paper Outline")
    run.bold = True
    run.font.size = Pt(16)
    bullets(doc, [
        "Paper title: Climate Vulnerability Does Not Equal Hunger: Measuring the Global Adaptation Buffer",
        "Target venue: Journal of Emerging Investigators (JEI)",
        "Formatting: 11 pt Times New Roman, 1.5 line spacing, 1-inch margins",
        "Purpose: section-by-section blueprint with word targets, numbers, and detailed bullet guidance",
        "All statistics from project pipeline (June 2026)",
        "Regenerate: python3 scripts/build_global_paper_outline.py",
    ])
    doc.add_page_break()


def add_formatting_guide(doc: Document) -> None:
    h1(doc, "Part A — Formatting & JEI Rules (read first)")
    h2(doc, "Manuscript typography")
    bullets(doc, [
        "Body text: 11 pt Times New Roman, 1.5 line spacing, 1-inch margins on all sides",
        "Headings in submission: 11–12 pt bold Times New Roman (match JEI template)",
        "Tables: 11 pt, embedded in the Word document",
        "Figure axis labels: at least 12 pt bold; other figure text 14–18 pt bold",
        "Do not put a title on the graph itself — title goes in the caption only",
        "Title page title: 12–14 pt bold; follow the official JEI template",
    ])
    h2(doc, "JEI section order (non-negotiable)")
    add_table(doc, ["Order", "Section", "Page?", "Word target"], [
        ["1", "Title Page", "Own page", "N/A"],
        ["2", "Summary", "Own page", "180–230 words (max 250)"],
        ["3", "Introduction", "Body", "550–700 words"],
        ["4", "Results", "Body", "900–1,100 words"],
        ["5", "Discussion", "Body", "500–650 words"],
        ["6", "Materials and Methods", "Body", "450–550 words"],
        ["7", "References", "End", "Not counted in 10-page limit"],
        ["8", "Acknowledgements", "End", "40–80 words"],
    ])
    h2(doc, "JEI Results writing rules")
    bullets(doc, [
        "Write in paragraphs, not bullet lists (this outline uses bullets for planning only)",
        "Each experiment: rationale → brief method → numbers → interpretation",
        "Cite figures/tables passively at the END of the sentence: (Table 1, Figure 1)",
        "Never write 'Figure 1 shows…' or 'As seen in Table 2…'",
        "Good example: 'Vulnerability alone explained 42.3% of variance (p < 0.001) (Table 1, Figure 1).'",
        "Call regressions 'Experiments' — not 'Model 1' or 'OLS Model 1' in the manuscript",
    ])
    h2(doc, "Words and concepts to AVOID in the JEI manuscript")
    bullets(doc, [
        "heteroskedasticity, HC1, clustered standard errors, bootstrap (say 'resampling' instead)",
        "KKT, Lagrangian, PCA, Lasso, Ridge, elastic net, fixed effects, panel FE",
        "seven-model hierarchy, shadow price, complementary slackness",
        "file paths, script names, '[PENDING]', draft status notes",
    ])
    h2(doc, "Words that ARE fine")
    bullets(doc, [
        "ordinary least squares regression",
        "Monte Carlo simulation / repeated random sampling",
        "coefficient, R-squared, p-value, standard error",
        "predicted value, fitted line, correlation",
        "resample, probability, 90% interval",
    ])
    doc.add_page_break()


def add_master_plan(doc: Document) -> None:
    h1(doc, "Part 0 — Master Document Plan")
    bullets(doc, [
        "Central claim: climate vulnerability is strongly associated with hunger, but a large structured residual remains",
        "Adaptation buffer = predicted hunger minus actual hunger; measures who beats or falls short of climate odds",
        "Total body target: 2,400–2,900 words (~5–6 pages)",
        "Hard JEI limit: 10 pages body excluding title page, references, and figure captions",
    ])
    h2(doc, "Internal analysis → JEI experiment labels")
    add_table(doc, ["Your script", "JEI name", "Question answered"], [
        ["Model 1 / Exp 1 OLS", "Experiment 1", "Does vulnerability predict hunger?"],
        ["Model 2", "Experiment 2", "Does vulnerability survive controls?"],
        ["Buffer rankings", "Experiment 3", "Who over/under-performs?"],
        ["Model 3 / buffer OLS", "Experiment 4", "What explains the buffer?"],
        ["global_monte_carlo.py part A", "Experiment 5", "How uncertain are rankings?"],
        ["global_monte_carlo.py part B", "Experiment 6", "When do buffers collapse?"],
    ])
    h2(doc, "Research questions — state all four in Introduction")
    bullets(doc, [
        "RQ1: How strongly does vulnerability predict hunger, and does it survive controls for income, rurality, disasters, and conflict?",
        "RQ2: Which countries systematically beat or fall short of the climate-only benchmark?",
        "RQ3: Can standard development indicators (especially readiness) explain the buffer, and how uncertain are rankings?",
        "RQ4: How easily does a positive buffer collapse under random shocks, and what protects it?",
    ])
    h2(doc, "Hypothesis — state in Introduction AND Summary")
    bullets(doc, [
        "Countries with higher ND-GAIN readiness will show larger positive adaptation buffers",
        "Must appear explicitly in both Summary and Introduction ¶4",
        "Report honestly in Results: readiness p = 0.12 — directional support only, not confirmed at 5%",
    ])
    h2(doc, "Figure/table budget")
    bullets(doc, [
        "Regenerate all assets: python3 scripts/build_jei_manuscript.py",
        "3 figures → separate JPEG uploads (figures/Figure1.jpeg, Figure2.jpeg, Figure3.jpeg)",
        "4 tables → embedded in Word manuscript (global_paper_JEI.docx)",
        "Experiment 4 has no table or figure; robustness has no table or figure",
    ])
    add_table(doc, ["Asset", "Experiment", "Where it goes", "File"], [
        ["Table 1", "Exp 1", "In Word", "Regression: constant, vulnerability, R²"],
        ["Figure 1", "Exp 1", "JPEG upload", "figures/Figure1.jpeg — scatter + OLS line"],
        ["Table 2", "Exp 2", "In Word", "Regression: vulnerability + 4 controls"],
        ["Table 3", "Exp 3 & 5", "In Word", "6 countries: buffer, 90% interval, P(over)"],
        ["Figure 2", "Exp 3", "JPEG upload", "figures/Figure2.jpeg — horizontal bar chart"],
        ["Table 4", "Exp 6", "In Word", "Baseline vs resilience collapse comparison"],
        ["Figure 3", "Exp 6", "JPEG upload", "figures/Figure3.jpeg — collapse probability bars"],
    ])
    doc.add_page_break()


def add_title_page_section(doc: Document) -> None:
    h1(doc, "Part 1 — Title Page (~½ page)")
    h2(doc, "Required content")
    bullets(doc, [
        "Title ≤ 110 characters including spaces",
        "All authors with superscript affiliation numbers",
        "Affiliation lines numbered to match superscripts",
        "Student author(s) first; mentor listed LAST",
        "Optional: equal-contribution footnote with asterisks",
        "Do NOT include Summary, hypothesis, figures, keywords, or draft notes",
    ])
    h2(doc, "Title options")
    add_table(doc, ["Title", "Chars"], [
        ["Climate Vulnerability Does Not Equal Hunger: Measuring the Global Adaptation Buffer", "83"],
        ["The Adaptation Buffer: Who Beats Their Climate Odds on Hunger?", "58"],
        ["Predicted vs. Actual Hunger Across 143 Countries: The Adaptation Buffer", "78"],
    ])
    bullets(doc, [
        "Pick the title that states your main finding or question — JEI readers scan for the 'so what'",
        "Author line format: [Student Name]¹, [Mentor Name]²",
        "Affiliation format: ¹ [High School], [City], [State]; ² [School/University], [City], [State]",
    ])
    doc.add_page_break()


def add_summary_section(doc: Document) -> None:
    h1(doc, "Part 2 — Summary (~180–230 words; MAX 250)")
    bullets(doc, [
        "Write the Summary LAST, after Results and Discussion are finalized",
        "Gets its own page immediately after the title page",
        "Do not use figure/table numbers in the Summary",
        "Avoid jargon: say 'regression' not 'OLS'; say 'resampling' not 'bootstrap'",
        "Do not use 'we argue' — state findings directly",
    ])
    h2(doc, "Required elements")
    add_table(doc, ["Element", "Include?", "Content"], [
        ["Problem", "Yes", "Vulnerability used as hunger proxy but similar-V countries differ"],
        ["Purpose", "Yes", "Test vulnerability→hunger; identify who beats prediction"],
        ["Hypothesis", "Yes", "Higher readiness → larger positive buffers"],
        ["Methods", "Yes", "Regression + Monte Carlo, N=143, ND-GAIN/FAO/WB/EM-DAT/ACLED"],
        ["Key results", "Yes", "See Block C bullets below"],
        ["Implications", "Yes", "Buffer = trackable outcome metric for aid"],
    ])
    h3(doc, "Block A — Problem (~35 words, ~2 sentences)")
    bullets(doc, [
        "Open with the policy assumption: climate vulnerability is widely used as a proxy for food insecurity",
        "State the paradox: countries with similar vulnerability scores report dramatically different undernourishment rates",
        "Give one concrete example: Bangladesh and Afghanistan have comparable ND-GAIN vulnerability but very different PoU",
        "Optional second example: Bangladesh vs Haiti for sharper contrast",
    ])
    h3(doc, "Block B — Methods + hypothesis (~40 words, ~2 sentences)")
    bullets(doc, [
        "Define adaptation buffer in plain English: hunger predicted from climate vulnerability minus actual hunger",
        "State sample size: N = 143 countries",
        "Name methods briefly: ordinary least squares regression and Monte Carlo simulation (20,000 draws)",
        "State hypothesis explicitly: higher adaptive readiness → larger positive buffers",
    ])
    h3(doc, "Block C — Key results (~70 words, ~4 sentences) — use these numbers")
    bullets(doc, [
        "Vulnerability alone: R² = 0.423, p < 0.001",
        "Vulnerability survives controls: coefficient falls from ~69 to ~40 but stays significant",
        "Buffer range: +14.6 (Kiribati) to −38.8 (Haiti); Bangladesh +9.1; Senegal +12.1",
        "Standard covariates explain only R² = 0.057 of buffer variation",
        "Monte Carlo (20,000 draws): 66 robust over-performers, 42 robust under-performers, 35 ambiguous",
        "Collapse simulation: 41.3% baseline vs 4.4% with resilience investment over 20 years",
    ])
    h3(doc, "Block D — Implication (~30 words, ~1–2 sentences)")
    bullets(doc, [
        "Adaptation buffer is a measurable, outcome-based indicator of effective adaptation",
        "Aid and adaptation finance should track buffers alongside vulnerability",
        "Goal: identify countries whose food-security gains are eroding before crises materialize",
    ])
    h2(doc, "Word-count tracker")
    bullets(doc, [
        "Block A: ~35 words",
        "Block B: ~40 words",
        "Block C: ~70 words",
        "Block D: ~30 words",
        "Total target: 175–210 words (stay under 250 hard max)",
    ])
    doc.add_page_break()


def add_introduction_section(doc: Document) -> None:
    h1(doc, "Part 3 — Introduction (~550–700 words, 4–5 paragraphs)")
    bullets(doc, [
        "Introduction = condensed literature review + your question",
        "Do NOT use a separate 'Literature Review' heading — JEI folds it into Introduction",
        "Number citations in order of first appearance; in-text format (3) at end of sentence",
    ])

    h2(doc, "Paragraph 1 — Hook (~110–130 words)")
    bullets(doc, [
        "Purpose: grab attention with a concrete country paradox before defining any new terms",
        "Open with Bangladesh vs Afghanistan: nearly identical ND-GAIN vulnerability scores",
        "Bangladesh V = 0.569, PoU = 10.4%; Afghanistan V ≈ 0.588, PoU = 28.1%",
        "State the logic: if vulnerability translated directly into hunger, they should look alike — they do not",
        "One closing sentence: this pattern repeats globally — some highly vulnerable countries report far less hunger than expected, others far more",
        "Cite FAO data when you first mention PoU (reference 11)",
        "Do NOT define the adaptation buffer yet — save for ¶4",
        "Tone: factual and precise; avoid 'shocking', 'unprecedented', or other hype words",
    ])

    h2(doc, "Paragraph 2 — Climate and food security (~120–140 words)")
    bullets(doc, [
        "Purpose: establish the field and show vulnerability→hunger is assumed in policy, not proven globally",
        "Climate threatens food security through yield losses, extreme events, disrupted livelihoods — cite (1, 2)",
        "FAO: climate extremes among principal drivers of recent global hunger increases, alongside conflict and downturns — cite (3)",
        "In practice, countries are ranked by climate vulnerability; rankings used for aid allocation and adaptation finance",
        "Gap to name: whether vulnerability maps cleanly onto hunger outcomes at global scale has received less direct attention than policy use of indices would suggest",
        "Citation 1: Wheeler & von Braun 2013 — climate → food security",
        "Citation 2: IPCC AR6 WGII 2022 — climate impacts",
        "Citation 3: FAO SOFI 2024 — hunger drivers",
    ])

    h2(doc, "Paragraph 3 — Measurement and the adaptation gap (~120–150 words)")
    bullets(doc, [
        "Purpose: introduce ND-GAIN, FAO PoU, and why outcome-based metrics matter",
        "ND-GAIN Country Index: most widely used composite national climate vulnerability measure — cite (4, 12)",
        "ND-GAIN readiness: economic, governance, and social capacity to deploy adaptation investment",
        "Define PoU on first use: percentage of population below minimum dietary energy requirement — cite (7)",
        "Define vulnerability: ND-GAIN composite 0–1 scale, higher = more vulnerable",
        "Adaptation tracking literature: persistent gap between planning and documented outcomes — cite (5, 6)",
        "Cross-country evidence: income is dominant correlate of undernourishment — cite (8)",
        "Conflict has emerged as leading driver of acute food crises — cite (3, 9)",
        "Optional one sentence on index limitations — cite (10); sets up Discussion",
        "Bridge sentence: what is missing is an outcome-based measure of whether countries beat or fall short of climate-predicted hunger",
    ])

    h2(doc, "Paragraph 4 — Your contribution (~150–180 words)")
    bullets(doc, [
        "Purpose: define buffer, state all four RQs, state hypothesis, preview methods and headline findings",
        "Define adaptation buffer: difference between undernourishment predicted from vulnerability alone and observed undernourishment",
        "Explain computation: fit cross-country regression of undernourishment on vulnerability; for each country, predicted minus actual PoU",
        "Positive buffer = less hunger than climate vulnerability alone would predict (over-performer)",
        "Negative buffer = more hunger than predicted (under-performer)",
        "Embed all four research questions (see Part 0) as prose or a short numbered list",
        "State hypothesis explicitly: higher ND-GAIN readiness → larger positive buffers",
        "Methods preview: 143 countries, ordinary least squares regression, Monte Carlo simulation with 20,000 random draws",
        "Results preview: vulnerability explains much but not all hunger variation",
        "Results preview: buffer rankings are statistically robust (66 over / 42 under in MC)",
        "Results preview: standard covariates barely predict overperformance (R² = 0.057)",
        "Results preview: resilience investments can reduce buffer-collapse probability roughly tenfold",
    ])

    h2(doc, "Paragraph 5 — Personal motivation (OPTIONAL, ~60–80 words)")
    bullets(doc, [
        "Only include if authentic and if Introduction is still under ~650 words",
        "Why you cared: Bangladesh case, climate news, food security interest",
        "What surprised you: Niger overperforms despite extreme vulnerability; Botswana underperforms despite moderate vulnerability",
        "Why public data made this feasible for a high school researcher",
        "Skip entirely if page count is tight — cut this first",
    ])

    h2(doc, "Introduction citation map")
    add_table(doc, ["#", "Source", "First use"], [
        ["1", "Wheeler & von Braun 2013", "¶2 climate-food"],
        ["2", "IPCC AR6 2022", "¶2"],
        ["3", "FAO SOFI 2024", "¶2 hunger drivers; ¶3 conflict"],
        ["4", "Chen / ND-GAIN report", "¶3"],
        ["5", "Ford et al. 2013", "¶3 adaptation tracking"],
        ["6", "Berrang-Ford et al. 2021", "¶3"],
        ["7", "Cafiero et al. 2018", "¶3 PoU"],
        ["8", "Headey 2013", "¶3 income"],
        ["9", "Burke et al. 2015", "¶3 conflict"],
        ["10", "Hinkel 2011", "¶3 optional index critique"],
        ["11", "FAOSTAT website", "¶1 PoU data"],
        ["12", "ND-GAIN website", "¶3 index data"],
    ])
    doc.add_page_break()


def add_results_section(doc: Document) -> None:
    h1(doc, "Part 4 — Results (~900–1,100 words)")
    bullets(doc, [
        "Results = WHAT you found and WHAT it means. Write as connected paragraphs.",
        "Methods = HOW you did it. Full procedure lives in Part 6 — do NOT repeat it in Results.",
        "Each experiment in Results: one-sentence setup → key numbers → interpretation → cite table/figure at end",
        "Label blocks 'Experiment 1', 'Experiment 2', etc. — never 'Model 1'",
    ])

    h2(doc, "Results vs Methods — what goes where")
    bullets(doc, [
        "RESULTS: research question, sample size (N), coefficients, p-values, R², country examples, patterns, interpretation",
        "RESULTS: one plain-English sentence on what you did (e.g. 'We regressed undernourishment on vulnerability for 143 countries')",
        "METHODS: data source URLs, merge keys, variable definitions, software names, formulas, iteration counts, random seeds",
        "METHODS: Monte Carlo step-by-step algorithm, collapse simulation equation, noise formulas, shock distributions",
        "METHODS: significance threshold (α = 0.05), robustness check specifications",
        "NEVER in Results: Python/statsmodels, ISO3 merge logic, bootstrap resampling steps, Gamma distribution parameters",
        "NEVER in Methods: interpretation of Haiti vs Niger, policy implications, 'this suggests adaptation works'",
        "Rule of thumb: if a reader needs it to REPLICATE your analysis → Methods. If they need it to UNDERSTAND your finding → Results.",
    ])

    h2(doc, "Figures & tables — full inventory (build with scripts/build_jei_manuscript.py)")
    bullets(doc, [
        "All 3 figures: separate JPEG uploads in Editorial Manager (NOT embedded in the Word body for submission)",
        "All 4 tables: embedded in the Word manuscript with captions above each table",
        "Figure files live at: Global Research Paper/figures/Figure1.jpeg, Figure2.jpeg, Figure3.jpeg",
        "Regenerate figures + tables: python3 scripts/build_jei_manuscript.py",
    ])
    add_table(doc, ["Asset", "Type", "Experiment", "In Word?", "File / source"], [
        ["Table 1", "Regression output", "Experiment 1", "Yes — embed", "Constant, vulnerability coef, R², N=143"],
        ["Figure 1", "Scatter + OLS line", "Experiment 1", "No — JPEG upload", "figures/Figure1.jpeg"],
        ["Table 2", "Regression output", "Experiment 2", "Yes — embed", "Vulnerability + 4 controls, R², N=139"],
        ["Table 3", "Country buffers + MC", "Experiments 3 & 5", "Yes — embed", "6 countries with 90% intervals"],
        ["Figure 2", "Horizontal bar chart", "Experiment 3", "No — JPEG upload", "figures/Figure2.jpeg"],
        ["Table 4", "Collapse comparison", "Experiment 6", "Yes — embed", "Baseline vs resilience scenarios"],
        ["Figure 3", "Grouped bar chart", "Experiment 6", "No — JPEG upload", "figures/Figure3.jpeg"],
    ])
    bullets(doc, [
        "Every table and figure must be cited at least once in Results (passive format at end of sentence)",
        "Experiment 4 (buffer determinants) has no dedicated table or figure — report coefficients in prose only",
        "Robustness checks (R1–R5) have no table — one short paragraph at end of Results",
    ])

    h2(doc, "Stats primer — Pearson correlation (read before Experiment 1)")
    h3(doc, "What is Pearson r? (one sentence)")
    bullets(doc, [
        "Pearson r (Pearson correlation coefficient) = a single number from −1 to +1 that summarizes how strongly two variables move together in a straight-line pattern",
        "r measures both DIRECTION (+ or −) and STRENGTH (how close to ±1) of a linear relationship",
        "r = +1: perfect positive linear relationship — when x goes up, y always goes up in proportion",
        "r = −1: perfect negative linear relationship — when x goes up, y always goes down in proportion",
        "r = 0: no linear relationship — knowing x tells you nothing about y on average",
        "In practice, r is rarely exactly ±1 or 0 — most real data fall somewhere in between",
        "Pearson r only measures LINEAR association — two variables can be strongly related in a curved way but still have r near 0",
        "Your paper uses OLS regression (Experiment 1), but r describes the same scatter in Figure 1 in one number",
    ])
    h3(doc, "The formula — piece by piece (plain English)")
    bullets(doc, [
        "Formula: r = Σ(xᵢ − x̄)(yᵢ − ȳ) / √[Σ(xᵢ − x̄)² · Σ(yᵢ − ȳ)²]",
        "xᵢ and yᵢ = each country's values (one point on the scatter plot); n = 143 countries",
        "x̄ (x-bar) = mean of all x values = average vulnerability across countries — the center of the x distribution",
        "ȳ (y-bar) = mean of all y values = average hunger across countries — the center of the y distribution",
        "Numerator Σ(xᵢ − x̄)(yᵢ − ȳ): for each country, multiply 'how far above/below average vulnerability' by 'how far above/below average hunger', then add up",
        "  Both above average → positive product; both below average → also positive (minus × minus = plus)",
        "  One above and one below → negative product",
        "  Large positive sum = countries with high x tend to have high y (x and y rise and fall together)",
        "  Large negative sum = high x tends to pair with low y (x and y move in opposite directions)",
        "Denominator √[Σ(xᵢ − x̄)² · Σ(yᵢ − ȳ)²]: scales the numerator so r always lands between −1 and +1 regardless of units",
        "  Σ(xᵢ − x̄)² = total spread of x around its mean (how much countries differ in vulnerability)",
        "  Σ(yᵢ − ȳ)² = total spread of y around its mean (how much countries differ in hunger)",
        "  Multiplying and taking the square root = 'normalize by how much each variable varies' — without this step, r could be any number",
        "Plain English: r asks 'Do countries above average on x also tend to be above average on y?' and adjusts for how spread out the data are",
    ])
    h3(doc, "In YOUR paper — what is x and what is y?")
    bullets(doc, [
        "x = ND-GAIN climate vulnerability (0–1 scale, higher = more vulnerable) — Figure 1 horizontal axis",
        "y = FAO prevalence of undernourishment (PoU, in %) — Figure 1 vertical axis",
        "Each dot = one country; r summarizes whether dots trend upward (positive), downward (negative), or scattered (near zero)",
        "Positive r expected: more vulnerable countries tend to report higher hunger — both x and y are 'bad' outcomes, so they should rise together",
        "Negative r would mean: more vulnerable countries tend to have LESS hunger — opposite of what you find",
        "Experiment 1 fits a regression line through these 143 points; Pearson r describes the same scatter pattern the line summarizes",
    ])
    h3(doc, "Link to Experiment 1 — r and R² are two views of the same relationship")
    bullets(doc, [
        "Experiment 1 reports R² = 0.423 — vulnerability explains 42.3% of cross-country variance in hunger",
        "For a simple two-variable regression (one x, one y), r² = R² — squaring r gives the same number as R²",
        "So r = √0.423 ≈ +0.65 (take the positive root because vulnerability and hunger rise together)",
        "Interpretation: moderate-to-strong positive linear association — countries with higher vulnerability tend to have higher hunger, but the link is not perfect",
        "r ≈ +0.65 is in the 'moderate-strong' range: clearly not zero, clearly not perfect, consistent with R² = 42%",
        "You do NOT need to report r separately in the JEI manuscript if you already report R² and the regression coefficient — but understanding r helps you write Results clearly",
        "Figure 1 scatter: points cluster upward-left to upward-right; r = +0.65 captures that upward tilt",
    ])
    h3(doc, "Pearson r vs R² — don't confuse them")
    bullets(doc, [
        "r (correlation coefficient): signed number from −1 to +1; tells you direction (+ or −) AND strength of linear association",
        "R² (R-squared, coefficient of determination): number from 0 to 1 (or 0% to 100%); always positive; tells you what FRACTION of y's variance is explained by x",
        "Key identity (simple bivariate case only): R² = r² — so R² = 0.423 means r² = 0.423, hence r = ±0.65",
        "Sign lives in r, not R²: R² = 0.423 could mean r = +0.65 OR r = −0.65 — you need the regression slope (or scatter plot) to know which",
        "In Experiment 1: slope = +69.42 (positive) → r is positive → r ≈ +0.65",
        "R² = 0.423 plain English: 'About 42% of the country-to-country differences in hunger can be accounted for by differences in vulnerability alone'",
        "The other 58% is unexplained by vulnerability — that residual motivates the adaptation buffer (Experiments 3–6)",
        "R² is often easier for JEI readers: 'explains 42% of variance' — you do NOT need to report r in the manuscript",
        "In Experiment 2 (multiple predictors): R² = 0.464 but there is no single r — r is only defined for two variables at a time",
        "If you mention r: say 'Pearson correlation ≈ +0.65' once in Results or Methods — optional, not required",
    ])
    h3(doc, "Worked example — reading r = +0.65 from Figure 1")
    bullets(doc, [
        "Imagine plotting all 143 countries: x-axis = vulnerability, y-axis = hunger %",
        "If r were +1.0, all points would fall exactly on an upward straight line — no scatter at all",
        "If r were 0, the cloud of points would look like a random blob with no tilt",
        "With r ≈ +0.65, points tilt upward but scatter widely — Bangladesh (moderate V, low hunger) sits below the line; Haiti (moderate V, very high hunger) sits above",
        "That scatter IS the adaptation buffer: countries that beat or fall short of the line",
        "OLS regression draws the best-fit line through that cloud; R² = 0.423 tells you how tight the cloud is around the line",
        "r close to +1 would mean almost perfect straight line — your cloud is real but noisy (many countries off the line)",
        "r near 0 would mean vulnerability and hunger unrelated — your r ≈ 0.65 rules that out",
        "r negative would mean more vulnerable → less hunger on average — opposite of your finding",
        "The 58% unexplained (1 − R²) is exactly why you need the adaptation buffer — correlation is strong but not perfect",
    ])
    h3(doc, "JEI guidance — where to mention correlation")
    bullets(doc, [
        "RESULTS: use plain English — 'associated with', 'predicts', 'positively related', 'stronger vulnerability was linked to higher hunger'",
        "RESULTS: report R² = 0.423 and/or say 'strong positive association' — cite (Table 1, Figure 1)",
        "RESULTS: optional one clause: 'Vulnerability and undernourishment were positively correlated (R² = 0.423)' — but R² alone is fine",
        "RESULTS: do NOT paste the Pearson formula",
        "METHODS (optional one sentence): 'We assessed the linear association between vulnerability and undernourishment using ordinary least squares regression; for the bivariate model, R² equals the square of the Pearson correlation coefficient'",
        "METHODS: include the full r formula only if a mentor/reviewer asks — JEI high-school audience does not require it in the main text",
        "Never claim r or R² proves causation — correlation ≠ causation (say this in Discussion limitations too)",
    ])
    h3(doc, "Good vs bad phrasing — correlation in Results")
    bullets(doc, [
        "Good: 'Climate vulnerability was positively associated with undernourishment across 143 countries (R² = 0.423, p < 0.001) (Table 1, Figure 1).'",
        "Good: 'Vulnerability and undernourishment were positively associated (R² = 0.423) (Table 1, Figure 1)'",
        "Good: 'Countries with higher vulnerability tended to report higher hunger rates'",
        "Good: 'Vulnerability alone explained 42.3% of cross-country variation in hunger — a moderate-to-strong association.'",
        "Good: 'The positive slope (69.42) indicates that more vulnerable countries report more hunger on average.'",
        "Optional: 'Pearson correlation was approximately +0.65'",
        "Bad: 'Vulnerability and hunger are highly correlated (r = 0.65)' without also giving R² or the regression — r alone is less standard in regression papers",
        "Bad: 'The Pearson formula proves vulnerability causes hunger'",
        "Bad: 'Vulnerability causes hunger because r = 0.65' — correlation does not establish causation",
        "Bad: 'The Pearson correlation proves vulnerability drives hunger' — overclaims; say 'associated with' or 'predicts'",
        "Bad: 'We calculated Σ(xᵢ − x̄)(yᵢ − ȳ)…' in Results — save for Methods or skip entirely",
        "Bad: 'r = 0.65 means 65% of countries are correlated' (misinterpretation — use R² language for variance explained)",
        "Bad: 'There is a 42% correlation between vulnerability and hunger' — R² is not 'percent correlation'; say '42% of variance explained'",
    ])

    h2(doc, "Experiment 1 — Vulnerability predicts hunger, but incompletely (160–180 words)")
    h3(doc, "IN RESULTS — write these bullets as one paragraph")
    bullets(doc, [
        "Rationale: test whether ND-GAIN vulnerability alone predicts cross-country undernourishment; this regression defines the buffer benchmark",
        "One-sentence method: 'We regressed FAO prevalence of undernourishment on ND-GAIN vulnerability for 143 countries using ordinary least squares'",
        "N = 143 countries",
        "Intercept = −20.04 (SE 2.59, p < 0.001)",
        "Vulnerability coefficient = 69.42 (SE 7.12, p < 0.001)",
        "Scale translation: a 0.1-point increase in vulnerability → 6.9 percentage points higher undernourishment",
        "R² = 0.423 — vulnerability explains 42.3% of cross-country variance",
        "Optional: Pearson r ≈ +0.65 (see Stats primer above) — same finding in correlation language; not required in manuscript",
        "Interpretation: strong gradient but 58% unexplained — this residual motivates the adaptation buffer",
        "Bangladesh example: V = 0.569, predicted PoU = 19.4%, actual PoU = 10.4%, buffer = +9.0 pp",
        "Say 'associated with' or 'predicts' — never claim causation",
        "Close sentence with: (Table 1, Figure 1)",
    ])
    h3(doc, "FIGURES & TABLES for Experiment 1")
    bullets(doc, [
        "Table 1 (in Word): columns = Term | Coefficient | Std. error | p-value; rows = Constant (−20.04), Vulnerability (69.42), R-squared (0.423, N=143)",
        "Figure 1 (JPEG upload): scatter plot, 143 blue points, red OLS fit line, R² = 0.423 in legend",
        "Figure 1 X-axis: ND-GAIN vulnerability (0–1); Y-axis: Prevalence of undernourishment (%)",
        "Figure 1: label Bangladesh with arrow; bold axis labels ≥ 12 pt; no title on the graph",
        "Figure 1 file: figures/Figure1.jpeg (regenerate via build_jei_manuscript.py)",
        "Figure 1 caption (in Word below placeholder): 'Figure 1: Relationship between ND-GAIN climate vulnerability and FAO prevalence of undernourishment across 143 countries (N = 143). Each point is one country. The line shows the ordinary least squares fit (R² = 0.423). Bangladesh is labeled as an example of a positive adaptation buffer.'",
        "In Results text, refer to Figure 1 only at end of sentence — never 'Figure 1 shows…'",
    ])
    h3(doc, "IN METHODS — do NOT put in Results (see Part 6 ¶3)")
    bullets(doc, [
        "OLS regression formula: undernourishment ~ vulnerability",
        "Software: Python 3, pandas, statsmodels",
        "Data: FAO PoU Item 210041 (2022–2024 window), ND-GAIN vulnerability 2022",
        "Significance threshold α = 0.05",
    ])
    h3(doc, "Common mistakes")
    bullets(doc, [
        "Claiming vulnerability 'causes' hunger",
        "Writing 'Figure 1 shows the relationship'",
        "Skipping the Bangladesh worked example",
        "Putting software names or data URLs in Results",
    ])

    h2(doc, "Experiment 2 — Climate–hunger link survives controls (140–160 words)")
    h3(doc, "IN RESULTS — write these bullets as one paragraph")
    bullets(doc, [
        "Rationale: test whether vulnerability–hunger link is just a poverty or conflict artifact",
        "One-sentence method: 'We added ln(GDP per capita), rural population share, disaster counts, and conflict fatalities to the Experiment 1 specification (N = 139)'",
        "Vulnerability coefficient = 40.07 (SE 11.48, p < 0.001) — down from 69.42 in Experiment 1 but still highly significant",
        "Key substantive point: coefficient shrinks ~69 → ~40; roughly 40% of raw gradient co-moves with observables",
        "ln(GDP per capita) = −2.84 (SE 1.04, p = 0.006) — higher income associated with lower hunger",
        "Rural share (%) = −0.047 (p = 0.29) — not significant",
        "Disaster count = 0.12 (p = 0.63) — not significant in this linear model",
        "ln(1 + conflict fatalities) = 0.065 (p = 0.83) — not significant in this linear model",
        "R² = 0.464, N = 139 complete cases (4 countries drop for missing covariates)",
        "Interpretation: vulnerability survives controls — not just a poverty artifact",
        "Caution: non-significant disaster/conflict does NOT prove shocks don't matter — may reflect measurement noise",
        "Close with: (Table 2)",
    ])
    h3(doc, "FIGURES & TABLES for Experiment 2")
    bullets(doc, [
        "Table 2 (in Word): columns = Variable | Coefficient | Std. error | p-value",
        "Table 2 rows: Vulnerability (40.07), ln(GDP per capita) (−2.84), Rural share (−0.047), Disaster count (0.12), ln(1 + conflict fatalities) (0.065), R-squared (0.464, N=139)",
        "No figure for Experiment 2",
        "If over page limit: fold Table 2 into prose and keep only vulnerability + GDP coefficients",
    ])
    h3(doc, "IN METHODS — do NOT put in Results (see Part 6 ¶3)")
    bullets(doc, [
        "Control variable definitions: ln(GDP pc) from World Bank WDI 2022, rural % from WDI, disasters summed 2013–2022 from EM-DAT, conflict fatalities 3-year average from ACLED",
        "Complete-case sample: N = 139 after dropping countries with missing covariates",
        "Same OLS framework and software as Experiment 1",
    ])

    h2(doc, "Experiment 3 — Global buffer rankings (170–200 words)")
    h3(doc, "IN RESULTS — write these bullets as one paragraph")
    bullets(doc, [
        "Rationale: identify which countries systematically beat or fall short of the climate-only hunger benchmark",
        "One-sentence method: 'For each country we computed the adaptation buffer as predicted minus actual undernourishment from Experiment 1'",
        "Buffer SD = 7.5 percentage points",
        "Buffer range: +14.6 (Kiribati) to −38.8 (Haiti)",
        "Median buffer ≈ +1.0 pp; middle 50% approximately −2.6 to +3.9",
        "Top overperformers to name: Senegal (+12.1), Niger (+11.2), Bangladesh (+9.0), Nepal (+8.7), Cambodia (+8.1)",
        "Niger key point: V = 0.636 (among most vulnerable globally) yet buffer = +11.2 — beats climate odds",
        "Bottom underperformers to name: Haiti (−38.8), Syria (−25.2), Kenya (−22.1), Madagascar (−20.7), Botswana (−14.1)",
        "Haiti key point: worst underperformer despite only moderate vulnerability (V = 0.510) — conflict dominates",
        "Botswana key point: V = 0.431 (moderate) but buffer = −14.1 — underperforms despite lower climate burden",
        "Pattern A — agrarian overperformers: Bangladesh, Senegal, Nepal, Cambodia, Philippines, Viet Nam (+7.1)",
        "Pattern B — conflict underperformers: Haiti, Syria, Liberia, CAR, Afghanistan",
        "Pattern C — vulnerability ranking ≠ buffer ranking in the tails",
        "Island caveat: Kiribati (+14.6) and Samoa (+12.0) may partly reflect FAO estimation in small states — mention briefly",
        "Close with: (Table 3, Figure 2)",
    ])
    h3(doc, "FIGURES & TABLES for Experiment 3")
    bullets(doc, [
        "Table 3 (in Word): shared with Experiment 5 — columns = Country | Buffer (pp) | 90% interval | P(over-performer)",
        "Table 3 countries (6 rows): Senegal (+12.1), Niger (+11.2), Bangladesh (+9.1), Viet Nam (+7.1), Kenya (−22.1), Haiti (−38.8)",
        "Table 3 also includes MC 90% intervals and P(over-performer) — cite in Experiment 5 as well",
        "Figure 2 (JPEG upload): horizontal bar chart of 8 countries",
        "Figure 2 countries (in order): Senegal, Niger, Bangladesh, Viet Nam, Kenya, Haiti, Botswana, Madagascar",
        "Figure 2 colors: teal/green (#2a9d8f) for positive buffer, red/orange (#e76f51) for negative",
        "Figure 2: vertical black line at zero; X-axis = Adaptation buffer (percentage points)",
        "Figure 2 file: figures/Figure2.jpeg; data from output/global_mc_buffer_uncertainty.csv",
        "Figure 2 caption: 'Figure 2: Adaptation buffer for selected countries. The buffer equals predicted minus actual undernourishment (percentage points). Positive values indicate countries outperforming their climate-only benchmark. Data from 143-country cross-section, 2022.'",
    ])
    h3(doc, "IN METHODS — do NOT put in Results (see Part 6 ¶2)")
    bullets(doc, [
        "Buffer formula: fitted PoU from Experiment 1 minus observed PoU",
        "Positive buffer = less hunger than climate-only benchmark predicts",
        "Computed for all 143 countries in the Experiment 1 sample",
    ])

    h2(doc, "Experiment 4 — Standard covariates barely explain the buffer (130–150 words)")
    h3(doc, "IN RESULTS — write these bullets as one paragraph")
    bullets(doc, [
        "Rationale: test hypothesis that ND-GAIN readiness correlates with larger buffers; ask if buffer is just repackaged development data",
        "One-sentence method: 'We regressed the adaptation buffer on readiness, ln(GDP per capita), rural share, disaster count, and conflict fatalities (N = 139)'",
        "R² = 0.057 — only ~6% of buffer variation explained; this is a finding, not a failure",
        "ND-GAIN readiness coefficient = +13.35 (SE 8.69, p = 0.12) — largest coefficient, directionally supports hypothesis",
        "Readiness NOT significant at 5% level — call result hypothesis-generating, not confirmed",
        "ln(GDP per capita) = 0.41 (p = 0.71) — not significant",
        "Rural share = 0.065 (p = 0.13) — not significant",
        "Disaster count = −0.055 (p = 0.83) — not significant",
        "ln(1 + fatalities) = 0.049 (p = 0.87) — not significant",
        "Practical implication: buffers cannot be inferred from standard indicators — must compute directly",
        "No table or figure for Experiment 4 — coefficients stay in prose only",
    ])
    h3(doc, "FIGURES & TABLES for Experiment 4")
    bullets(doc, [
        "No dedicated table or figure",
        "If space allows, could add a small coefficient table — but JEI manuscript keeps this in text only",
    ])
    h3(doc, "IN METHODS — do NOT put in Results (see Part 6 ¶3)")
    bullets(doc, [
        "Regression: buffer ~ ln(GDP pc) + readiness + rural % + disasters + ln(1+fatalities)",
        "Vulnerability excluded from right-hand side (buffer already net of vulnerability)",
        "N = 139 complete cases",
    ])

    h2(doc, "Monte Carlo primer — read this before Experiments 5 & 6")
    h3(doc, "What is Monte Carlo? (one sentence)")
    bullets(doc, [
        "Monte Carlo = run the same analysis thousands of times with randomness, then summarize what usually happens",
        "Named after the Monte Carlo casino — it's about repeated random trials, not gambling",
        "You already use this idea informally: 'if I flipped this coin 10,000 times, how often would I get heads?'",
    ])
    h3(doc, "Why do we need it here?")
    bullets(doc, [
        "Experiment 1 gives ONE best-fit line through 143 countries — but that line would change slightly if you had different countries or slightly different hunger data",
        "Each country's buffer = predicted hunger minus actual hunger — so buffer uncertainty comes from (a) the line wobbling and (b) FAO hunger numbers not being perfectly exact",
        "Monte Carlo quantifies: 'How often does Bangladesh still look like an over-performer when we account for that uncertainty?'",
        "Experiment 6 uses the same idea for a different question: 'If shocks hit every year, how often does a healthy buffer fall below zero?'",
    ])
    h3(doc, "Key vocabulary (plain English)")
    bullets(doc, [
        "Draw / trial / iteration / simulation = one full repeat of the random process (we do 20,000)",
        "Resample with replacement = pick 143 countries randomly from the list, allowing duplicates (like drawing names from a hat and putting each name back)",
        "Bootstrap = resample the real data many times and refit the regression each time — see how much the answers wobble",
        "Slope (β₁ or b₁) = the number from Experiment 1: how much hunger rises when vulnerability rises (69.42 in the original fit)",
        "Intercept (β₀ or b₀) = where the line crosses the y-axis (−20.04 in Experiment 1)",
        "Measurement noise = pretend each country's reported hunger rate could be a little higher or lower than FAO published, because all statistics have error",
        "90% interval = the middle 90% of outcomes across 20,000 repeats: throw away the lowest 5% and highest 5%",
        "If 90% interval is [58.3, 81.5], it means: in 95% of simulations the slope landed below 81.5, and in 95% it landed above 58.3",
        "P(over-performer) or P(buffer > 0) = fraction of the 20,000 simulations where that country's buffer was positive",
        "P = 1.00 means the buffer was positive in all 20,000 runs — very confident it's a real over-performer",
        "P = 0.00 means the buffer was never positive in 20,000 runs — very confident it's a real under-performer",
        "P between 0.10 and 0.90 = ambiguous — could reasonably be either side of zero",
        "Robust over-performer = P ≥ 0.90 (positive buffer in at least 18,000 of 20,000 simulations)",
        "Robust under-performer = P ≤ 0.10 (positive buffer in at most 2,000 of 20,000 simulations)",
        "Collapse = buffer drops below zero at least once during the 20-year simulation (like Bangladesh after bad cyclones)",
    ])
    h3(doc, "Experiment 5 vs Experiment 6 — don't mix them up")
    bullets(doc, [
        "Experiment 5 (Part A): uncertainty about the LINE and each country's BUFFER RIGHT NOW — uses real 143-country data",
        "Experiment 6 (Part B): uncertainty about the FUTURE — simulates one country's buffer year by year under random shocks",
        "Exp 5 answers: 'Is Bangladesh really beating its climate odds, or could that be noise?'",
        "Exp 6 answers: 'Even if Bangladesh has a +9 buffer today, how fragile is it over the next 20 years?'",
    ])
    h3(doc, "Worked example — Bangladesh in Experiment 5 (read this until it clicks)")
    bullets(doc, [
        "Experiment 1 says: Bangladesh buffer = +9.0 pp (predicted hunger 19.4%, actual 10.4%)",
        "But what if the regression line were slightly steeper? Then predicted hunger would be higher → buffer smaller",
        "What if the line were slightly flatter? Then predicted hunger lower → buffer larger",
        "What if FAO's 10.4% is a bit off — maybe true hunger is 11%? Then buffer drops",
        "Monte Carlo tries ALL those 'what ifs' 20,000 times by randomly changing the line AND adding noise to hunger",
        "After 20,000 tries, Bangladesh's buffer ranged from about +6.1 to +12.1 — always positive",
        "P(over) = 1.00 means 20,000 out of 20,000 tries were positive → we're very sure Bangladesh over-performs",
        "Haiti went from about −48 to −30 every time — always negative → P = 0.00",
    ])
    h3(doc, "Worked example — slope 90% interval (the confusing one)")
    bullets(doc, [
        "The slope is the 69.42 from Experiment 1: 'each 1.0 unit of vulnerability adds ~69 pp hunger'",
        "When we resample countries 20,000 times, we get 20,000 slightly different slopes",
        "Those 20,000 slopes might range from about 50 to 90 — most cluster near 69",
        "90% interval [58.3, 81.5] = throw away the most extreme 10% (5% low, 5% high); what's left is this range",
        "Plain English for the paper: 'The vulnerability–hunger relationship stayed in a similar range across 20,000 resampled datasets'",
        "You use this to show Experiment 1's main finding is STABLE — not a fluke of which 143 countries you happened to have",
    ])

    h2(doc, "Experiment 5 — Monte Carlo uncertainty (160–180 words)")
    h3(doc, "What this experiment does (step by step, in plain English)")
    bullets(doc, [
        "Step 0 — Start from Experiment 1: you already have one regression line (hunger ~ vulnerability) and one buffer per country",
        "Step 1 — Repeat 20,000 times: randomly pick 143 countries FROM your list WITH replacement (some countries appear twice, some not at all)",
        "Step 2 — Each time, fit a NEW regression line on that random sample → get a slightly different slope and intercept",
        "Step 3 — Each time, add small random error to every country's hunger rate (bigger error for countries with higher hunger)",
        "Step 4 — Each time, recompute every country's buffer using the new line and the noisy hunger rates",
        "Step 5 — After 20,000 repeats, you have 20,000 buffer values per country → summarize with average, 90% interval, and P(positive)",
    ])
    h3(doc, "IN RESULTS — write these bullets as one paragraph")
    bullets(doc, [
        "Rationale: a single regression gives one buffer per country, but we want to know how trustworthy that ranking is",
        "Plain-English method sentence for Results: 'To assess uncertainty, we repeated the buffer calculation 20,000 times, each time randomly resampling countries, refitting the vulnerability regression, adding realistic noise to reported hunger rates, and recomputing every country's buffer'",
        "Do NOT put the 5-step algorithm in Results — one sentence only; details go in Methods",
    ])
    h3(doc, "Numbers to report — the regression line itself (optional 1 sentence in Results)")
    bullets(doc, [
        "Original Experiment 1 slope = 69.42 — this is the 'best guess' for how much hunger rises with vulnerability",
        "Across 20,000 resamples, the slope averaged ~69.6 — almost the same, so the main gradient is stable",
        "Slope 90% interval = [58.3, 81.5] — PLAIN ENGLISH: in 20,000 reruns, the slope was between 58.3 and 81.5 about 90% of the time",
        "You do NOT need to explain percentiles in the paper — just say 'the slope remained between 58.3 and 81.5 in 90% of simulations'",
        "Intercept 90% interval = [−24.4, −16.0] — same idea for the intercept (optional; skip if tight on words)",
        "R² 90% interval = [0.34, 0.53] — vulnerability still explains roughly 34–53% of hunger variation across reruns (optional)",
    ])
    h3(doc, "Numbers to report — country buffer uncertainty (main focus)")
    bullets(doc, [
        "For each country you get THREE summary numbers from the 20,000 simulations:",
        "  (1) MC mean buffer = average buffer across all 20,000 runs (e.g. Bangladesh +9.1)",
        "  (2) 90% interval = range where the buffer landed in the middle 90% of runs (e.g. Bangladesh [+6.1, +12.1])",
        "  (3) P(over-performer) = what fraction of runs had a positive buffer (e.g. Bangladesh 1.00 = 100%)",
        "Bangladesh: mean +9.1, interval [+6.1, +12.1], P = 1.00 → always positive, never close to zero",
        "Haiti: mean −38.8, interval [−47.9, −29.6], P = 0.00 → always negative",
        "Senegal: +12.1, P = 1.00; Niger: +11.2, P = 1.00 — also rock-solid over-performers",
        "Kenya: −22.1, P = 0.00 — rock-solid under-performer",
    ])
    h3(doc, "Numbers to report — how many countries in each bucket")
    bullets(doc, [
        "66 robust over-performers: P(over) ≥ 0.90 — buffer was positive in at least ~18,000 of 20,000 simulations",
        "42 robust under-performers: P(over) ≤ 0.10 — buffer was positive in at most ~2,000 of 20,000 simulations",
        "35 ambiguous: P between 0.10 and 0.90 — too close to zero to call confidently",
        "Key message for Results: Bangladesh, Haiti, Senegal, Niger are NOT statistical flukes — their classifications survive heavy uncertainty",
    ])
    h3(doc, "How to explain this in Results without jargon")
    bullets(doc, [
        "Good: 'Bangladesh's buffer remained positive in all 20,000 simulations (Table 3)'",
        "Good: 'Haiti's buffer remained negative in all 20,000 simulations'",
        "Good: 'The vulnerability–hunger slope stayed between 58.3 and 81.5 in 90% of simulations, consistent with Experiment 1'",
        "Bad: 'We performed bootstrap resampling with replacement' (save for Methods)",
        "Bad: 'The 5th–95th percentile of the slope distribution…' (too technical for JEI Results)",
        "Close Results paragraph with: (Table 3)",
    ])
    h3(doc, "FIGURES & TABLES for Experiment 5")
    bullets(doc, [
        "Table 3 (in Word): same table as Experiment 3 — cite again here for MC columns",
        "Table 3 columns: Country | Buffer (pp) | 90% interval | P(over-performer)",
        "Table 3 example row — Bangladesh: +9.1 | [+6.1, +12.1] | 1.00",
        "Table 3 example row — Haiti: −38.8 | [−47.9, −29.6] | 0.00",
        "No separate figure for Experiment 5 — uncertainty is in Table 3, not a graph",
        "Optional figure you did NOT include: histogram of Bangladesh's 20,000 buffer values — skip to save figure budget",
    ])
    h3(doc, "IN METHODS — do NOT put in Results (see Part 6 ¶4)")
    bullets(doc, [
        "B = 20,000 iterations; random seed = 42 (same seed → same results if someone reruns your code)",
        "Step 1: bootstrap resample 143 countries with replacement; refit Experiment 1 OLS line",
        "Step 2: add Normal(0, σ) noise to each country's PoU, where σ = max(1.0, 0.10 × PoU)",
        "Step 3: recompute buffer for every country: (new predicted hunger) − (noisy actual hunger)",
        "90% interval = 5th percentile to 95th percentile of the 20,000 buffer values per country",
        "P(over-performer) = count(buffer > 0) / 20,000",
        "Classification: P ≥ 0.90 = robust over-performer; P ≤ 0.10 = robust under-performer; else ambiguous",
        "Output files: output/global_mc_buffer_uncertainty.csv (per country), output/global_mc_slope_distribution.csv (slope/intercept/R²)",
    ])

    h2(doc, "Experiment 6 — Buffer collapse simulation (160–180 words)")
    h3(doc, "What this experiment does (step by step, in plain English)")
    bullets(doc, [
        "This is a DIFFERENT Monte Carlo from Experiment 5 — it simulates the FUTURE, not resampling past data",
        "Imagine Bangladesh starts with a +9.1 pp buffer today (from Experiment 1 / MC)",
        "Simulate 20 years, one year at a time, repeated 20,000 times",
        "Each year TWO things happen: (1) buffer drifts back toward +9.1 if it was knocked down, and (2) maybe a bad shock hits",
        "Shock = random bad year (cyclone, drought, conflict) that knocks several percentage points off the buffer",
        "20% shock probability per year = roughly 1 in 5 years has a shock, same in both scenarios",
        "Baseline scenario: shocks are BIG (average 4.0 pp) and recovery is SLOW (buffer heals 25% of the gap per year)",
        "Resilience scenario: shocks are SMALLER (average 2.4 pp) and recovery is FASTER (40% of the gap per year)",
        "Collapse = at any point during the 20 years, buffer drops below zero (hunger worse than climate predicts)",
        "Count how many of the 20,000 simulated futures had at least one collapse",
    ])
    h3(doc, "Key vocabulary for Experiment 6")
    bullets(doc, [
        "B₀ (starting buffer) = +9.1 pp — Bangladesh's buffer at year 0",
        "Recovery rate ρ (rho) = how fast the buffer bounces back toward +9.1 after a shock (0.25 = slow, 0.40 = fast)",
        "Shock size = how many percentage points get knocked off the buffer in a bad year (4.0 baseline vs 2.4 resilience)",
        "P(collapse within 20 years) = in what fraction of 20,000 simulated futures did the buffer ever go negative",
        "41.3% baseline = in about 4 out of 10 simulated futures, Bangladesh lost its positive buffer at least once",
        "4.4% resilience = in only about 1 in 25 simulated futures did that happen — roughly tenfold safer",
        "Expected years underwater = average number of years the buffer spent below zero (1.01 baseline vs 0.05 resilience)",
        "Median buffer year 20 = where the buffer ended up after 20 years in the typical simulation (+6.6 baseline vs +8.2 resilience)",
        "IMPORTANT: these are ILLUSTRATIVE scenarios with assumed shock sizes — not a precise forecast for Bangladesh",
    ])
    h3(doc, "IN RESULTS — write these bullets as one paragraph")
    bullets(doc, [
        "Rationale: motivate with real example — Bangladesh's buffer partially collapsed after 2007–2009 cyclones before partial rebuilding",
        "One-sentence method: 'We simulated 20,000 twenty-year trajectories of a +9.1 percentage-point buffer under random annual shocks'",
        "Define collapse for the reader: buffer falls below zero at least once during the 20-year period",
        "Both scenarios have the SAME 20% annual chance of a shock — the only difference is shock size and recovery speed",
        "Baseline: P(collapse within 20 years) = 41.3%",
        "Baseline: P(collapse within 10 years) = 21.6% (optional if space tight)",
        "Baseline: expected years underwater = 1.01; median buffer at year 20 = +6.6 pp",
        "Resilience investment: P(collapse within 20 years) = 4.4%",
        "Resilience: P(collapse within 10 years) = 2.0% (optional)",
        "Resilience: expected years underwater = 0.05; median buffer at year 20 = +8.2 pp",
        "Main comparison: 41.3% → 4.4% = roughly tenfold reduction in collapse risk at identical shock frequency",
        "Frame as: defending the buffer (smaller shocks, faster recovery) matters even when climate exposure is unchanged",
        "Do NOT list Gamma distribution, ρ symbol, or the update equation in Results",
        "Close with: (Table 4, Figure 3)",
    ])
    h3(doc, "How to explain Experiment 6 in Results without jargon")
    bullets(doc, [
        "Good: 'Under baseline shock conditions, a +9.1 pp buffer had a 41.3% chance of turning negative at least once within twenty years'",
        "Good: 'When shocks were smaller and recovery faster, that risk fell to 4.4%'",
        "Good: 'Both scenarios faced the same 20% annual probability of a shock'",
        "Bad: 'We drew shocks from a Gamma distribution with mean 4.0' (Methods only)",
        "Bad: 'Bangladesh will collapse with 41% probability' (too strong — say 'a buffer of this size' or 'under these assumptions')",
    ])
    h3(doc, "FIGURES & TABLES for Experiment 6")
    bullets(doc, [
        "Table 4 (in Word): columns = Scenario | P(collapse ≤ 20 y) | Expected years underwater | Buffer yr 20 (median, pp)",
        "Table 4 row 1: Baseline (large shocks, slow recovery) — 41.3%, 1.01, +6.6",
        "Table 4 row 2: Resilience investment — 4.4%, 0.05, +8.2",
        "Figure 3 (JPEG upload): grouped bar chart, two bars side by side",
        "Figure 3 bars: Baseline = 41.3% (red/orange #e76f51), Resilience investment = 4.4% (teal #2a9d8f)",
        "Figure 3 Y-axis: P(buffer falls below 0 within 20 years, %); value labels on top of each bar",
        "Figure 3 file: figures/Figure3.jpeg; data from output/global_mc_collapse_dynamics.csv",
        "Figure 3 caption: 'Figure 3: Monte Carlo simulation of buffer collapse over 20 years (N = 20,000 trials). Starting buffer +9.1 percentage points (Bangladesh). Both scenarios face the same 20% annual shock probability; resilience investment reduces mean shock size and increases recovery rate.'",
    ])
    h3(doc, "IN METHODS — do NOT put in Results (see Part 6 ¶5)")
    bullets(doc, [
        "T = 20 years; 20,000 trials per scenario",
        "Yearly update equation: B_{t+1} = B_t + ρ(target − B_t) − shock + noise",
        "Shock probability = 0.20/year; shock size ~ Gamma(mean 4.0 baseline, 2.4 resilience)",
        "Recovery rate ρ = 0.25 baseline, 0.40 resilience",
        "Output file: output/global_mc_collapse_dynamics.csv",
    ])

    h2(doc, "Robustness checks — primer (read before writing the last Results paragraph)")
    h3(doc, "What is a robustness check? (plain English)")
    bullets(doc, [
        "A robustness check = rerun the analysis with ONE thing changed, to see if your main conclusion still holds",
        "Think of it as: 'Was our result just an accident of this specific dataset choice?'",
        "You are NOT running six new experiments — you write ONE short paragraph at the end of Results summarizing them",
        "No table, no figure for robustness in the JEI manuscript — numbers go in prose only",
        "Goal: show reviewers (and yourself) the headline finding is not fragile",
    ])
    h3(doc, "What are we trying to defend?")
    bullets(doc, [
        "Main finding from Experiment 1: climate vulnerability predicts hunger (β ≈ 69, R² ≈ 0.42, p < 0.001)",
        "Main finding from Experiment 3: buffers vary a lot — Bangladesh over-performs, Haiti under-performs",
        "Robustness asks: do these still look true if we change the hunger data window, the vulnerability index, the outcome measure, or drop tiny countries?",
    ])
    h3(doc, "Headline numbers to compare against (Experiment 1 baseline)")
    bullets(doc, [
        "Baseline (main paper): ND-GAIN vulnerability → FAO PoU 2022–2024, N = 143, β = 69.42, R² = 0.423",
        "If a robustness check gives similar β and R² → supports main finding",
        "If it gives a very different R² but same direction/significance → partial support — explain what that means",
        "If it fails entirely → you would mention honestly (none of ours fail the main story)",
    ])

    h3(doc, "R1 — Alternate FAO hunger window")
    bullets(doc, [
        "Question: 'Is our result just because we used FAO's 2022–2024 hunger average?'",
        "What changed: PoU window 2020–2022 instead of 2022–2024 (still ND-GAIN vulnerability 2022, same 143 countries)",
        "What stayed the same: Experiment 1 setup (hunger ~ vulnerability OLS)",
        "Results: β = 68.7 (vs 69.4 baseline), R² = 0.442 (vs 0.423 baseline), N = 143, p < 0.001",
        "What it means: almost identical — the vulnerability–hunger gradient does NOT depend on which FAO 3-year window you pick",
        "Verdict: SUPPORTS main finding — cite in Results as your first robustness sentence",
    ])

    h3(doc, "R2a — WRI 2025 vulnerability component (alternate index)")
    bullets(doc, [
        "Question: 'Is this only true for ND-GAIN, or do other vulnerability indices show a similar pattern?'",
        "What changed: replace ND-GAIN vulnerability with WRI 2025 vulnerability component (scaled 0–1)",
        "What stayed the same: FAO PoU outcome, Experiment 1 regression",
        "Results: β = 31.48, R² = 0.289, N = 143, p < 0.001",
        "What it means: WRI vulnerability still significantly predicts hunger, but explains less variance (29% vs 42%) — different indices measure different things",
        "Verdict: SUPPORTS direction (vulnerability still matters); weaker R² is expected when swapping indices — optional mention if space; not required in the 80-word paragraph",
    ])

    h3(doc, "R2b — WRI 2025 composite World Risk Index (alternate index)")
    bullets(doc, [
        "Question: 'What if we use a broad disaster-risk score instead of a climate-vulnerability score?'",
        "What changed: WRI composite index (exposure + vulnerability + lack of coping capacity combined)",
        "Results: β = 4.56, R² = 0.002, N = 143, p = 0.53 (NOT significant)",
        "What it means: the broad composite WRI index does NOT predict hunger — composite risk ≠ undernourishment",
        "Why this SUPPORTS your paper: your main index (ND-GAIN) captures something hunger-relevant that a generic risk composite misses",
        "Plain English for Results: 'pure composite disaster-risk scores did not predict hunger (R² = 0.002)'",
        "Verdict: SUPPORTS using ND-GAIN specifically — optional in short paragraph; good for Discussion/limitations",
    ])

    h3(doc, "R2c — Global Data Lab climate vulnerability (alternate index)")
    bullets(doc, [
        "Question: 'Does the vulnerability–hunger link replicate with a completely independent climate vulnerability measure?'",
        "What changed: Global Data Lab national climate vulnerability index 2022 (scaled 0–1) instead of ND-GAIN",
        "Results: β = 39.88, R² = 0.408, N = 92, p < 0.001",
        "What it means: different index, smaller sample (92 countries have GDL data), but still explains ~41% of hunger variation — very close to baseline 42%",
        "Verdict: STRONG SUPPORT — cite in Results as second robustness sentence",
    ])

    h3(doc, "R2d — ND-GAIN exposure sub-index only (narrower measure)")
    bullets(doc, [
        "Question: 'Is the link driven only by physical climate exposure, not the full ND-GAIN composite?'",
        "What changed: ND-GAIN exposure sub-index only (drops sensitivity/adaptive capacity components)",
        "Results: β = 50.60, R² = 0.153, N = 143, p < 0.001",
        "What it means: exposure alone predicts hunger significantly but explains only 15% of variance — full ND-GAIN vulnerability (42%) captures more than raw exposure",
        "Verdict: SUPPORTS that climate exposure matters, but full vulnerability index is stronger — optional mention; skip in 80-word paragraph",
    ])

    h3(doc, "R3 — WHO child stunting as alternate outcome")
    bullets(doc, [
        "Question: 'Is this just a quirk of FAO's prevalence-of-undernourishment metric?'",
        "What changed: outcome = WHO child stunting prevalence (latest available per country) instead of FAO PoU",
        "What stayed the same: ND-GAIN vulnerability as predictor, Experiment 1 regression form",
        "Results: β = 110.82, R² = 0.624, N = 121, p < 0.001",
        "What it means: vulnerability predicts stunting even MORE strongly (62% of variance) — the link is not a PoU artifact",
        "Caution: stunting and PoU measure different things (children vs whole population, chronic malnutrition vs calorie intake) — do not say they are the same",
        "Verdict: STRONG SUPPORT — cite in Results as third robustness sentence",
    ])

    h3(doc, "R4 — Standardized Experiment 2 (NOT in JEI manuscript)")
    bullets(doc, [
        "What it is: rerun Experiment 2 with all predictors z-scored (mean 0, SD 1) to compare effect sizes on a common scale",
        "Results: vulnerability still significant (β = 3.67 per 1 SD, p < 0.001); ln(GDP pc) still negative (p = 0.006); disasters/conflict still NS",
        "Verdict: confirms Experiment 2 sign pattern — kept in script/output but DROPPED from JEI paper to save space",
        "Do NOT mention R4 in the JEI manuscript unless a reviewer asks",
    ])

    h3(doc, "R5 — Drop small states (population < 1 million)")
    bullets(doc, [
        "Question: 'Are buffer rankings driven by tiny island countries with unreliable FAO estimates (Kiribati, Samoa)?'",
        "What changed: exclude all countries with population < 1 million; refit Experiment 1 and recompute buffers",
        "Results: β = 75.51, R² = 0.458, N = 125 (vs 143); Bangladesh buffer = +10.5, rank 6 (vs +9.0, rank 9 in full sample)",
        "Top over-performers after drop: Senegal (+13.4), Sudan (+13.2), Niger (+13.1), Mauritania (+13.0), Myanmar (+11.5), Bangladesh (+10.5)",
        "Kiribati (+14.6) and Samoa (+12.0) disappear from top of list — they were small-state artifacts",
        "What it means: the vulnerability gradient is slightly stronger without micro-states; Bangladesh still a top over-performer",
        "Verdict: SUPPORTS main country stories — cite Bangladesh rank ~6 in Results closing sentence",
    ])

    h3(doc, "Which checks to put in the ~80–100 word Results paragraph")
    bullets(doc, [
        "Lead with THREE strongest: R1 (FAO window), R2c (GDL index), R3 (stunting outcome)",
        "Add ONE line on R5: Bangladesh buffer +10.5, rank ~6 when small states excluded",
        "Skip in short paragraph: R2a, R2b, R2d (save for Discussion/limitations if needed)",
        "End with one concluding sentence: core gradient and Bangladesh over-performance stable across specifications",
        "Total: 4–5 sentences, ~80–100 words — no heading, no 'Experiment 7'",
    ])

    h2(doc, "Robustness paragraph — IN RESULTS (~80–100 words)")
    h3(doc, "What to write (bullet → combine into one paragraph)")
    bullets(doc, [
        "Open: 'We tested whether the main findings were sensitive to data choices' (or similar — one clause only)",
        "R1: vulnerability coefficient remained 68.7 (R² = 0.442) when using FAO PoU averaged over 2020–2022 instead of 2022–2024",
        "R2c: the gradient replicated with the Global Data Lab climate vulnerability index (R² = 0.408, N = 92)",
        "R3: vulnerability significantly predicted WHO child stunting prevalence (R² = 0.624), indicating the pattern is not specific to FAO PoU",
        "R5: excluding countries with population below one million left Bangladesh's buffer at +10.5 percentage points (rank 6)",
        "Close: 'These checks indicate that the core vulnerability–hunger association and Bangladesh's over-performance are robust to alternative data specifications'",
    ])
    h3(doc, "Good vs bad phrasing in Results")
    bullets(doc, [
        "Good: 'The vulnerability–hunger association was stable across alternative FAO hunger windows (R² = 0.442)'",
        "Good: 'The pattern replicated with an independent vulnerability index (Global Data Lab, R² = 0.408)'",
        "Good: 'Vulnerability also predicted child stunting (R² = 0.624), suggesting the result is not an artifact of FAO measurement'",
        "Bad: 'We ran robustness checks R1 through R5' (too mechanical — explain what you changed)",
        "Bad: 'Results were robust' (vague — give at least two numbers)",
        "Bad: 'WRI composite proved ND-GAIN is better' (overclaims — R2b only shows composite ≠ hunger)",
        "Bad: 'Stunting proves causation' (stunting is a different outcome, not causal proof)",
    ])

    h3(doc, "IN METHODS — one-line spec per check (see Part 6 ¶7)")
    bullets(doc, [
        "R1: Experiment 1 with FAO PoU 3-year window ending 2022 (2020–2022 average)",
        "R2a: Experiment 1 with WRI 2025 vulnerability component (/100) replacing ND-GAIN",
        "R2b: Experiment 1 with WRI 2025 composite World Risk Index (/100) replacing ND-GAIN",
        "R2c: Experiment 1 with Global Data Lab national climate vulnerability 2022 (/100)",
        "R2d: Experiment 1 with ND-GAIN exposure sub-index only",
        "R3: Experiment 1 with WHO child stunting (latest national estimate) replacing PoU",
        "R4: Experiment 2 with z-scored regressors (not reported in JEI manuscript)",
        "R5: Experiment 1 on countries with World Bank population ≥ 1,000,000; buffers recomputed",
        "Script: global_robustness_checks.py; output: global_robustness_summary.csv",
    ])
    doc.add_page_break()


def add_discussion_section(doc: Document) -> None:
    h1(doc, "Part 5 — Discussion (~500–650 words, 4–5 paragraphs)")
    bullets(doc, [
        "Only discuss results already reported in Results — no new numbers",
        "Label speculation explicitly: 'We speculate that…'",
        "Do not repeat Results verbatim — interpret and connect to policy and literature",
    ])

    h2(doc, "Paragraph 1 — Main finding (~120 words)")
    h3(doc, "What this paragraph is FOR (read first)")
    bullets(doc, [
        "Purpose: answer 'So what?' — translate your Results into a clear takeaway for a non-expert reader",
        "This is the FIRST Discussion paragraph — it restates your headline claim, it does NOT introduce new experiments",
        "Discussion ¶1 = interpretation; Results = numbers. Do not re-list coefficients or country tables here",
        "Do NOT put limitations here (that's ¶4), conflict examples (¶2), or collapse simulation (¶3)",
        "Tone: confident but not causal — 'associated with,' 'explains variation,' not 'proves' or 'causes'",
        "Target length: ~100–130 words, ~4–6 sentences",
    ])
    h3(doc, "Discussion vs Results — what moves from Results into ¶1")
    bullets(doc, [
        "FROM Results (already reported) → reinterpret in plain language, do not re-copy tables",
        "Experiment 1: R² = 0.423 → 'vulnerability explains 42% of cross-country hunger variation'",
        "Experiment 1: 58% unexplained → 'the remaining variation is structured and policy-relevant, not random noise'",
        "Experiment 3: buffer SD 7.5 pp, range −38.8 to +14.6 → optional one phrase; don't re-list countries",
        "Experiment 4: R² = 0.057 for buffer determinants → 'standard covariates explain only ~6% of buffer variation'",
        "Experiment 4: readiness p = 0.12 → 'buffer cannot be inferred from readiness alone' (not 'hypothesis confirmed')",
        "DO NOT introduce in ¶1: Monte Carlo 66/42 split, 41% vs 4% collapse, Haiti/Syria names — those belong in later ¶s or Summary",
    ])
    h3(doc, "The headline claim — 'Climate vulnerability does not equal hunger'")
    bullets(doc, [
        "Open with this sentence or a close variant — it's your paper's title and central message",
        "Plain English: countries with similar climate vulnerability scores can have very different hunger rates",
        "You already proved the statistical link in Experiment 1 (R² = 0.423) — vulnerability DOES predict hunger",
        "The twist: predicting ~42% is not the same as vulnerability BEING hunger — a large, meaningful gap remains",
        "Bangladesh vs Afghanistan (from Introduction) is the intuition: same ballpark vulnerability, very different PoU",
        "Do not contradict Experiment 1 — you are nuancing it, not denying the gradient",
    ])
    h3(doc, "The adaptation buffer — what it IS and why it matters")
    bullets(doc, [
        "Define again in one sentence: buffer = hunger predicted from vulnerability alone minus actual hunger",
        "Positive buffer = country beats its climate odds (less hunger than vulnerability alone would predict)",
        "Negative buffer = country falls short (more hunger than climate burden alone would predict)",
        "Key move in Discussion: the buffer NAMES the leftover 58% as a measurable quantity — not just 'error' or 'noise'",
        "Policy angle: if you only track vulnerability, you miss who is actually outperforming or underperforming on hunger OUTCOMES",
        "Contrast with adaptation tracking literature (cite 5, 6): most frameworks count plans, policies, dollars — not whether hunger outcomes beat climate predictions",
        "Your contribution in one line: the buffer is an outcome-based measure of effective adaptation, not planned adaptation",
    ])
    h3(doc, "The value-added analogy (optional but powerful — one sentence)")
    bullets(doc, [
        "Education uses 'value-added' to compare schools: not raw test scores, but scores vs. what demographics predict",
        "Same logic here: not raw hunger rate, but hunger vs. what climate vulnerability predicts",
        "Helps JEI readers who know neither econometrics nor climate policy grasp why a residual metric is useful",
        "Keep to one sentence — do not turn ¶1 into a literature essay",
        "Example: 'Like value-added metrics in education, the buffer benchmarks food-security outcomes against circumstances rather than raw hunger rates'",
    ])
    h3(doc, "Why standard data cannot replace the buffer (tie to Experiment 4)")
    bullets(doc, [
        "Experiment 4 showed R² = 0.057 — income, readiness, disasters, conflict explain only ~6% of who over/under-performs",
        "Plain English: you cannot look at a country's GDP or readiness score and reliably guess its buffer",
        "Therefore: the buffer must be COMPUTED and MONITORED directly each year — not inferred from development indicators",
        "This is a finding, not a failure — unexplained variation means the buffer carries information those indicators miss",
        "One sentence for ¶1: 'Because standard covariates explain only a small fraction of buffer variation, the buffer must be tracked directly rather than inferred from income or readiness alone'",
    ])
    h3(doc, "Suggested sentence order (~120 words total)")
    bullets(doc, [
        "Sentence 1 — Headline: 'Climate vulnerability does not equal hunger.'",
        "Sentence 2 — How much vulnerability explains: 'Across 143 countries, vulnerability explains 42% of cross-country variation in undernourishment, but the remaining variation is structured and policy-relevant.'",
        "Sentence 3 — Introduce buffer: 'The adaptation buffer makes that residual visible by benchmarking actual hunger against what climate exposure alone would predict.'",
        "Sentence 4 — Outcome vs plans (optional): 'Unlike adaptation tracking that counts plans and investment, the buffer measures revealed food-security performance relative to climate burden (5, 6).'",
        "Sentence 5 — Cannot infer from covariates: 'Because standard development indicators explain only about 6% of buffer variation, the buffer must be monitored directly rather than inferred from income or readiness alone.'",
        "Sentence 6 — Value-added analogy (optional if over word count): one sentence comparing to education benchmarking",
    ])
    h3(doc, "Numbers you may cite in ¶1 (all already in Results)")
    bullets(doc, [
        "42% / R² = 0.423 — vulnerability's share of hunger variation",
        "58% — unexplained share (1 − 0.423); call it 'structured' not 'random'",
        "~6% / R² = 0.057 — share of buffer variation explained by standard covariates (Experiment 4)",
        "N = 143 — optional; usually skip in Discussion unless emphasizing global scope",
        "Skip in ¶1: 66/42 MC counts, ±7.5 SD, Bangladesh +9.0, 41.3%/4.4% collapse — save for later paragraphs",
    ])
    h3(doc, "Good vs bad phrasing for ¶1")
    bullets(doc, [
        "Good: 'Climate vulnerability does not equal hunger'",
        "Good: 'The remaining variation is structured and policy-relevant'",
        "Good: 'The adaptation buffer benchmarks actual outcomes against climate-predicted hunger'",
        "Good: 'Outcome-based indicator of effective adaptation'",
        "Good: 'Associated with' / 'explains variation in' — cautious causal language",
        "Bad: 'Vulnerability does not affect hunger' (contradicts Experiment 1)",
        "Bad: 'We proved adaptation works' (overclaim — observational study)",
        "Bad: 'Readiness causes larger buffers' (p = 0.12 — not confirmed)",
        "Bad: Re-listing Table 1 coefficients or naming Haiti/Bangladesh in detail (save for ¶2–3)",
        "Bad: Starting with limitations ('Although our study has flaws…') — limitations go in ¶4",
    ])
    h3(doc, "What comes AFTER ¶1 (do not cram into paragraph 1)")
    bullets(doc, [
        "¶2 — Conflict vs climate: Haiti, Syria, Niger, Botswana patterns; aid misdirection",
        "¶3 — Buffers as assets: Bangladesh cyclones, 41% vs 4% collapse, defending vs rebuilding",
        "¶4 — Limitations: not causal, benchmark-relative, PoU measurement, MC assumptions",
        "¶5 — Future work: panel tracking, richer covariates, buffer monitoring in adaptation finance",
    ])

    h2(doc, "Paragraph 2 — Conflict vs climate (~110 words)")
    bullets(doc, [
        "Largest underperformers dominated by political violence and state fragility, not climate exposure alone",
        "Name examples: Haiti (−38.8), Syria (−25.2), Liberia, CAR, Afghanistan",
        "These are not always the most climate-vulnerable countries on ND-GAIN",
        "Conversely, several most climate-exposed countries hold large positive buffers: Niger (+11.2), Bangladesh (+9.0), Senegal (+12.1)",
        "Speculate (label as speculation): conflict and governance collapse destroy food-security performance faster than climate exposure alone predicts",
        "Policy implication: vulnerability-weighted aid formulas may miss countries where hunger most exceeds climate fundamentals",
        "Underperformers may need conflict-sensitive food security intervention, not just climate adaptation",
    ])

    h2(doc, "Paragraph 3 — Buffers as assets (~110 words)")
    bullets(doc, [
        "Bangladesh as prototype: decades of agricultural innovation, disaster preparedness, social protection built +9 pp buffer",
        "Real-world fragility: buffer partially collapsed after 2007–2009 cyclone sequence, then partially rebuilt",
        "Collapse simulation: even healthy buffer has 41.3% twenty-year collapse risk under baseline shocks",
        "Resilience investments reduce that to 4.4% — roughly tenfold reduction at identical shock frequency",
        "Highest-value interventions may defend existing buffers, not only reduce hazard scores",
        "Examples of buffer-defending interventions: early warning, grain storage, safety nets, stress-tolerant crop varieties",
        "Frame as asset management: defending buffer cheaper than rebuilding food system after collapse",
    ])

    h2(doc, "Paragraph 4 — Limitations (~130–150 words)")
    bullets(doc, [
        "Not causal: all analyses observational cross-section; countries not randomized into vulnerability levels",
        "Benchmark-relative: buffer depends on ND-GAIN vulnerability and bivariate regression line; different benchmark → different buffer rankings",
        "PoU measurement: FAO 3-year average smooths acute shocks; values censored at 2.5%; stunting check (R3) helps but different nutritional construct",
        "Small states: Kiribati/Samoa top rankings may partly reflect FAO estimation uncertainty; drop out when population < 1M (R5)",
        "Monte Carlo assumptions: 10% measurement error and shock parameters are stated assumptions; comparison between scenarios matters more than exact percentages",
        "Hypothesis not confirmed: readiness p = 0.12 at 5% level — treat as hypothesis-generating",
        "Conflict–climate entanglement: linear covariates understate interaction effects",
        "Ecological fallacy: country-level patterns; cannot infer individual-level relationships",
        "JEI wants specific limitations — avoid generic 'human error may have occurred'",
    ])
    add_table(doc, ["Limitation", "Where to address"], [
        ["Not causal", "Discussion ¶4 + cautious language throughout"],
        ["Benchmark-relative", "Discussion ¶4 + Methods buffer construction"],
        ["PoU measurement", "Discussion ¶4 + Methods + R3 in Results"],
        ["Small states", "Results Exp 3 + R5 + Discussion"],
        ["MC assumptions", "Methods + Discussion"],
        ["Readiness p = 0.12", "Results Exp 4 + Discussion"],
        ["Ecological fallacy", "Discussion future work"],
    ])

    h2(doc, "Paragraph 5 — Future work + closing (~80–100 words)")
    bullets(doc, [
        "Annual panel tracking of buffer level and trend over time",
        "Richer covariates: agricultural productivity, aid flows, governance indices, cereal yields",
        "Test whether shrinking buffers predict subsequent hunger increases (early warning system)",
        "Alternative outcome: IPC Phase 3+ acute food insecurity",
        "Closing proposal: integrate buffer level and trend into climate-adaptation monitoring alongside vulnerability",
        "Goal: make countries with eroding food-security gains visible before crises fully materialize",
        "End on forward-looking policy note, not a summary of Results",
    ])
    doc.add_page_break()


def add_methods_section(doc: Document) -> None:
    h1(doc, "Part 6 — Materials and Methods (~450–550 words)")
    h2(doc, "What is Methods? — start here if you're confused")
    bullets(doc, [
        "Methods answers ONE question: 'How would someone else repeat my analysis and get the same numbers?'",
        "Results = WHAT you found ('vulnerability explains 42% of hunger variation')",
        "Methods = HOW you did it ('we merged FAO and ND-GAIN data by country code, then ran OLS regression in Python')",
        "JEI puts Methods AFTER Results and Discussion — weird order, but that's their rule",
        "Write Methods in PARAGRAPHS (prose), not bullet lists — the outline uses bullets to plan; your manuscript does not",
        "Target: ~450–550 words total, split into 6–7 short paragraphs",
        "You already have a draft in global_paper_JEI.docx — you can start there and expand using the bullets below",
        "Do NOT interpret findings in Methods — no 'this shows Bangladesh adapts well' or 'conflict matters'",
        "Do NOT cite Table 1, Figure 1, etc. in Methods — only in Results",
    ])
    h3(doc, "Your job — write these 7 paragraphs in order")
    add_table(doc, ["¶", "Topic", "Answers the question", "~Words"], [
        ["1", "Data", "Where did each dataset come from? How many countries?", "120"],
        ["2", "Buffer", "How did you compute the adaptation buffer?", "60"],
        ["3", "Regression", "What regressions did you run? What software?", "100"],
        ["4", "Monte Carlo (Exp 5)", "How did the 20,000-iteration uncertainty simulation work?", "100"],
        ["5", "Collapse sim (Exp 6)", "How did the 20-year shock simulation work?", "100"],
        ["6", "Robustness", "What one-thing-changed checks did you run?", "80"],
        ["7", "Ethics + stats", "Human subjects? Significance level? Code available?", "40"],
    ])
    h3(doc, "Step-by-step workflow")
    bullets(doc, [
        "Step 1: Open global_paper_JEI.docx → scroll to 'Materials and Methods' — there is already a draft",
        "Step 2: For each paragraph below, turn the outline bullets into 2–4 sentences of prose",
        "Step 3: Weave citations into sentences: (7, 11) not a separate reference list in Methods",
        "Step 4: Check you did NOT repeat Results numbers as findings — only describe procedure",
        "Step 5: Paste into JEI official template when ready; verify Methods is still ≤10 pages total body",
        "Tip: Methods is the EASIEST section to write — it's mostly copying your script logic into plain English",
    ])
    h3(doc, "Results vs Methods — quick examples")
    add_table(doc, ["Topic", "Results (findings)", "Methods (procedure)"], [
        ["Exp 1", "'Vulnerability explained 42% of variation (R² = 0.423)'", "'We regressed undernourishment on vulnerability using OLS'"],
        ["Buffer", "'Bangladesh had a +9.0 pp buffer'", "'Buffer = predicted PoU minus actual PoU from Experiment 1'"],
        ["Monte Carlo", "'Bangladesh was positive in 100% of simulations'", "'We repeated 20,000 iterations: resample, refit, add noise, recompute'"],
        ["Collapse", "'41.3% vs 4.4% collapse risk'", "'We simulated 20-year paths with 20% annual shock probability'"],
    ])
    bullets(doc, [
        "JEI places Methods AFTER Results — do not move it earlier",
        "Methods = HOW you did it. No interpretation, no policy implications, no country-pattern discussion",
        "Embed data sources in sentences with citations — avoid reagents-style bullet lists in the manuscript",
        "Cross-reference: each Methods paragraph below maps to the 'IN METHODS' bullets in Part 4",
    ])

    h2(doc, "Results vs Methods — Methods-side rules")
    bullets(doc, [
        "METHODS contains: data URLs, merge logic, variable definitions, formulas, software, seeds, algorithms, robustness specs",
        "METHODS does NOT contain: coefficient interpretation, 'Haiti underperforms because…', R² interpretation, policy recommendations",
        "METHODS does NOT repeat Results numbers — state procedure once; Results already reported the findings",
        "If you already said it in Results as a one-sentence setup, expand the full version here",
        "Do not cite figure/table numbers in Methods — cite them only in Results",
    ])

    h2(doc, "Paragraph 1 — Data sources and sample (~120 words)")
    h3(doc, "Data primer — what the Data paragraph is (read if confused)")
    bullets(doc, [
        "The Data paragraph answers: 'Where did every number in my paper come from?'",
        "You are NOT analyzing raw data in Methods — you are DESCRIBING which public datasets you downloaded and how you combined them",
        "Think of it like a recipe ingredients list: name each source, what variable you took, what year/window, how you matched countries together",
        "One paragraph of prose (~120 words) — not a table, not a bullet list in the final manuscript",
        "Every dataset gets a citation in parentheses: (7, 11) for FAO, (4, 12) for ND-GAIN, etc.",
    ])
    h3(doc, "The 5 data sources — what each one is and what you pulled")
    add_table(doc, ["Source", "What it measures", "Variable you use", "Year / window", "Cite"], [
        ["FAO FAOSTAT", "Hunger", "Prevalence of undernourishment (%) — Item 210041", "2022–2024 3-year average", "7, 11"],
        ["ND-GAIN", "Climate exposure + capacity", "Vulnerability index (0–1, higher = worse)", "2022", "4, 12"],
        ["ND-GAIN", "Adaptation capacity", "Readiness index (0–1, higher = better)", "2022", "4, 12"],
        ["World Bank WDI", "Development", "GDP per capita (USD); rural population (%)", "2022", "13"],
        ["EM-DAT", "Disasters", "Count of natural disasters per country", "Summed 2013–2022", "14"],
        ["ACLED", "Conflict", "Average conflict fatalities", "3-year average", "15"],
    ])
    h3(doc, "What each variable means in plain English")
    bullets(doc, [
        "Prevalence of undernourishment (PoU): % of people who don't get enough food energy — your OUTCOME (y-axis in Figure 1)",
        "ND-GAIN vulnerability: composite score of how climate-exposed and sensitive a country is — your main PREDICTOR (x-axis in Figure 1)",
        "ND-GAIN readiness: how prepared a country is to deploy adaptation investment — used in Experiment 4 only",
        "GDP per capita: income level — control in Experiment 2 (you use ln(GDP) in the regression)",
        "Rural population %: share of people living in rural areas — control in Experiments 2 and 4",
        "Disaster count: how many natural disasters hit the country 2013–2022 — control in Experiments 2 and 4",
        "Conflict fatalities: average yearly deaths from political violence (3-year window) — control in Experiments 2 and 4; you use ln(1 + fatalities)",
    ])
    h3(doc, "How countries were matched together (the merge)")
    bullets(doc, [
        "Every dataset uses a different country name format — FAO says 'Bangladesh', World Bank might spell it differently",
        "Solution: merge everything on ISO3 — a 3-letter country code (Bangladesh = BGD, United States = USA)",
        "Step 1: Start with World Bank country list to build a name → ISO3 lookup table",
        "Step 2: Pull FAO hunger data; map each FAO country name to ISO3",
        "Step 3: Pull ND-GAIN vulnerability and readiness by ISO3 directly",
        "Step 4: Pull World Bank GDP and rural % by ISO3 for 2022",
        "Step 5: Sum EM-DAT disasters 2013–2022 per ISO3",
        "Step 6: Pull ACLED 3-year average fatalities per ISO3",
        "Step 7: Inner merge on ISO3 — keep only countries that have BOTH hunger AND vulnerability (this defines N=143)",
        "Excluded: regional aggregates ('World', 'Africa') — country-level units only",
    ])
    h3(doc, "Sample sizes — why 143 vs 139")
    bullets(doc, [
        "N = 143: countries with non-missing FAO hunger AND ND-GAIN vulnerability — used in Experiments 1, 3, 5, 6",
        "N = 139: same 143 minus 4 countries missing GDP, rural %, disasters, or conflict data — used in Experiments 2 and 4",
        "The 4 dropped countries: whichever lack complete World Bank or ACLED rows (you don't need to name them in Methods unless asked)",
        "Say in Methods: 'Complete cases for multivariate models: N = 139'",
    ])
    h3(doc, "Timing and data quirks to mention")
    bullets(doc, [
        "FAO PoU is a 3-year average (2022–2024) — FAO smooths hunger estimates over multiple years",
        "ND-GAIN vulnerability is a single calendar year (2022) — slight timing mismatch, acknowledged in Methods",
        "FAO reports values below 2.5% as '<2.5' — you coded these as 2.5% (standard practice)",
        "Countries with zero disasters or zero conflict fatalities: filled as 0, not dropped",
        "GDP per capita: used natural log ln(GDP) in regressions because income effects are nonlinear",
    ])
    h3(doc, "What to write — stitch these facts into one paragraph")
    bullets(doc, [
        "Sentence 1: 'We merged publicly available country-level datasets by ISO3 code for 143 countries…'",
        "Sentence 2: Name FAO PoU (Item 210041, 2022–2024 window, <2.5 coded at 2.5) with cite (7, 11)",
        "Sentence 3: Name ND-GAIN vulnerability and readiness 2022 with cite (4, 12)",
        "Sentence 4: Name World Bank GDP pc and rural % 2022 with cite (13)",
        "Sentence 5: Name EM-DAT disasters 2013–2022 and ACLED 3-year fatalities with cite (14, 15)",
        "Sentence 6: 'Complete cases for multivariate models: N = 139'",
        "Do NOT list file paths or script names in the JEI manuscript",
    ])
    h3(doc, "Example Data paragraph (from global_paper_JEI.docx — you can adapt this)")
    bullets(doc, [
        "'We merged publicly available country-level datasets by ISO3 code for 143 countries with non-missing ND-GAIN vulnerability (2022) and FAO prevalence of undernourishment (Item 210041, 2022–2024 three-year window; values reported as <2.5% coded at 2.5%) (7, 11, 12). ND-GAIN composite vulnerability and readiness indices came from the Notre Dame Global Adaptation Initiative (4, 12). World Bank World Development Indicators provided 2022 GDP per capita and rural population share (13). EM-DAT disaster counts were summed over 2013–2022 (14). ACLED provided three-year average conflict fatalities (15). Complete cases for multivariate models: N = 139.'",
    ])
    h3(doc, "Data paragraph — common mistakes")
    bullets(doc, [
        "Bad: listing datasets as bullets in the manuscript (JEI wants prose paragraphs)",
        "Bad: forgetting citations for each source",
        "Bad: saying 'we downloaded data from the internet' without naming FAOSTAT, ND-GAIN, etc.",
        "Bad: explaining what hunger IS (that's Introduction) — Methods only says which FAO item you used",
        "Bad: reporting R² or coefficients in the Data paragraph (that's Results)",
        "Bad: naming individual countries in Methods Data section",
    ])
    bullets(doc, [
        "NOT in Results: full data provenance, merge keys, timing details, censoring rules",
        "Merge publicly available country-level datasets by ISO3 code",
        "Exclude regional aggregates; country-level units only",
        "FAO FAOSTAT PoU Item 210041: 2022–2024 three-year average; values below 2.5% coded as 2.5 — cite (7, 11)",
        "ND-GAIN vulnerability and readiness: 2022 annual values on 0–1 scale — cite (4, 12)",
        "World Bank WDI: GDP per capita and rural population share, 2022 — cite (13)",
        "EM-DAT: disaster count summed over 2013–2022 — cite (14)",
        "ACLED: conflict fatalities, 3-year average — cite (15)",
        "Bivariate / MC sample: N = 143 countries with non-missing vulnerability and PoU",
        "Multivariate sample: N = 139 complete cases (4 countries drop for missing covariates)",
        "Timing: vulnerability = calendar 2022; PoU = 2022–2024 average (modest stagger acknowledged)",
        "R1 robustness (mentioned briefly in Results): re-run with PoU averaged over 2020–2022",
    ])

    h2(doc, "Paragraph 2 — Buffer construction (~60 words) → supports Experiments 3–6")
    bullets(doc, [
        "NOT in Results: formal definition of fitted values, sign convention",
        "From Experiment 1 OLS regression, obtain fitted undernourishment value for each country",
        "Adaptation buffer = predicted PoU minus actual PoU",
        "Positive buffer: country reported less hunger than climate-only benchmark predicts",
        "Negative buffer: country reported more hunger than benchmark predicts",
        "Computed for all 143 countries in Experiment 1 sample",
    ])

    h2(doc, "Paragraph 3 — Regression analysis (~100 words) → supports Experiments 1, 2, 4")
    bullets(doc, [
        "NOT in Results: software package names, significance threshold, full list of control definitions",
        "Experiment 1: OLS regression of undernourishment ~ vulnerability",
        "Experiment 2: add ln(GDP per capita), rural share (%), EM-DAT disaster count, ln(1 + ACLED conflict fatalities)",
        "Experiment 4: OLS regression of buffer ~ ln(GDP pc) + readiness + rural % + disasters + ln(1+fatalities); vulnerability excluded",
        "Software: Python 3, pandas, statsmodels",
        "Significance threshold: α = 0.05",
    ])

    h2(doc, "Paragraph 4 — Monte Carlo uncertainty (~100 words) → supports Experiment 5")
    bullets(doc, [
        "NOT in Results: step-by-step algorithm, noise formula, classification thresholds, seed, output file paths",
        "Opening sentence for Methods: 'To quantify uncertainty in each country's buffer, we used Monte Carlo simulation with 20,000 iterations'",
        "B = 20,000 iterations; random seed = 42 (reproducibility)",
        "Each iteration — Step 1: randomly resample 143 countries with replacement (bootstrap); refit Experiment 1 OLS line → new slope and intercept",
        "Each iteration — Step 2: add measurement noise to each country's PoU: draw from Normal(0, σ) where σ = max(1.0, 0.10 × PoU)",
        "Each iteration — Step 3: recompute every country's buffer = (fitted hunger from new line) − (noisy actual hunger)",
        "After 20,000 iterations, per country: mean buffer, 90% interval (5th–95th percentile of the 20,000 values), P(buffer > 0)",
        "P(buffer > 0) = proportion of iterations where buffer was positive",
        "Classification: robust over-performer if P ≥ 0.90; robust under-performer if P ≤ 0.10; otherwise ambiguous",
        "Slope 90% interval reported in Results = 5th–95th percentile of the 20,000 refitted slopes",
        "Script: global_monte_carlo.py; outputs: global_mc_buffer_uncertainty.csv, global_mc_slope_distribution.csv",
    ])

    h2(doc, "Paragraph 5 — Collapse simulation (~100 words) → supports Experiment 6")
    bullets(doc, [
        "NOT in Results: update equation, Gamma distribution, recovery rate symbols, trial count",
        "Opening sentence: 'To model buffer fragility over time, we simulated 20,000 twenty-year trajectories of a +9.1 pp starting buffer'",
        "T = 20 years; 20,000 independent trials per scenario; B₀ = +9.1 pp (Bangladesh MC mean buffer)",
        "Yearly update: B_{t+1} = B_t + ρ(target − B_t) − shock + noise; target = initial B₀",
        "ρ (recovery rate): fraction of the gap to target closed each year — 0.25 baseline, 0.40 resilience",
        "Each year: 20% probability of a shock; if shock occurs, size drawn from Gamma distribution (mean 4.0 pp baseline, 2.4 pp resilience)",
        "Collapse = buffer < 0 at least once during the 20-year horizon",
        "Outcomes: P(collapse within 10 and 20 years), expected years with buffer < 0, median buffer at year 20",
        "Script: global_monte_carlo.py; output: global_mc_collapse_dynamics.csv",
        "Figure 3 shows P(collapse within 20 years) for baseline vs resilience; Table 4 lists full comparison",
    ])

    h2(doc, "Paragraph 6 — Robustness checks (~80–100 words) → summarized in Results closing paragraph")
    bullets(doc, [
        "NOT in Results: check labels (R1, R2a…), full coefficient tables, script names",
        "Opening sentence for Methods: 'We assessed sensitivity of the main findings to alternative data specifications'",
        "All checks reuse Experiment 1 (OLS: outcome ~ vulnerability) unless noted; script: global_robustness_checks.py",
        "R1 — FAO window: PoU 3-year average 2020–2022 (column ending 2022) instead of headline 2022–2024; N = 143",
        "R2a — WRI 2025 vulnerability component (0–100 rescaled to 0–1) replaces ND-GAIN vulnerability",
        "R2b — WRI 2025 composite World Risk Index (0–1) replaces ND-GAIN vulnerability",
        "R2c — Global Data Lab national climate vulnerability index 2022 (0–1); N = 92 due to coverage",
        "R2d — ND-GAIN exposure sub-index only (physical exposure component)",
        "R3 — WHO child stunting prevalence (latest national estimate per country) replaces FAO PoU; N = 121",
        "R4 — Experiment 2 with z-scored predictors (sign-stability check); not reported in JEI manuscript",
        "R5 — Experiment 1 restricted to World Bank population ≥ 1,000,000; buffers recomputed; N = 125",
        "Output: output/global_robustness_summary.csv, output/global_buffer_rankings_pop1m.csv (R5 rankings)",
    ])

    h2(doc, "Paragraph 7 — Ethics and reproducibility (~40 words)")
    bullets(doc, [
        "Aggregated publicly available country-level data only",
        "No human subjects or vertebrate animals involved",
        "Report: coefficients, standard errors, p-values, R²; significance at α = 0.05",
        "Analysis scripts available in project repository or upon request",
    ])

    h2(doc, "Methods checklist — before you submit")
    bullets(doc, [
        "[ ] Seven paragraphs in prose (not bullets)",
        "[ ] Every data source named with citation: FAO (7,11), ND-GAIN (4,12), World Bank (13), EM-DAT (14), ACLED (15)",
        "[ ] Sample sizes stated: N=143 bivariate, N=139 multivariate",
        "[ ] Buffer formula stated once",
        "[ ] All six experiments mentioned in Methods (Exp 3 = buffer construction, no separate regression)",
        "[ ] Monte Carlo 3 steps described",
        "[ ] Collapse simulation parameters stated (20 years, 20,000 trials, shock prob 0.20, two scenarios)",
        "[ ] Robustness checks listed briefly (R1, R2 alt indices, R3 stunting, R5 drop small states)",
        "[ ] No interpretation ('Haiti fails because…')",
        "[ ] No figure/table numbers cited",
        "[ ] Ethics sentence included",
    ])
    doc.add_page_break()


def add_back_matter(doc: Document) -> None:
    h1(doc, "Part 7 — References (numbered, JEI modified MLA)")
    bullets(doc, [
        "Number citations in order of FIRST appearance in text",
        "In-text format: (3). at end of sentence — period after citation parenthesis",
        "Reference list: numbered, no hanging indents, modified MLA format",
        "Include DOI or stable URL where available",
        "Access dates for web sources: 'Accessed 10 June 2026'",
    ])
    refs = [
        "Wheeler, Tim, and Joachim von Braun. \"Climate Change Impacts on Global Food Security.\" Science, vol. 341, no. 6145, 2013, pp. 508-513. https://doi.org/10.1126/science.1239402.",
        "IPCC. Climate Change 2022: Impacts, Adaptation and Vulnerability. Contribution of Working Group II to the Sixth Assessment Report. Cambridge University Press, 2022.",
        "FAO, IFAD, UNICEF, WFP, and WHO. The State of Food Security and Nutrition in the World 2024. Food and Agriculture Organization of the United Nations, 2024.",
        "Chen, C., et al. University of Notre Dame Global Adaptation Index: Country Index Technical Report. University of Notre Dame, 2015.",
        "Ford, James D., et al. \"How to Track Adaptation to Climate Change: A Typology of Approaches for National-Level Application.\" Ecology and Society, vol. 18, no. 3, 2013. https://doi.org/10.5751/ES-05732-180340.",
        "Berrang-Ford, Lea, et al. \"A Systematic Global Stocktake of Evidence on Human Adaptation to Climate Change.\" Nature Climate Change, vol. 11, no. 11, 2021, pp. 989-1000.",
        "Cafiero, Carlo, Sara Viviani, and Mark Nord. \"Food Security Measurement in a Global Context.\" Measurement, vol. 116, 2018, pp. 146-152.",
        "Headey, Derek. \"Developmental Drivers of Nutritional Change: A Cross-Country Analysis.\" World Development, vol. 42, 2013, pp. 76-88.",
        "Burke, Marshall, Solomon M. Hsiang, and Edward Miguel. \"Climate and Conflict.\" Annual Review of Economics, vol. 7, 2015, pp. 577-617.",
        "Hinkel, Jochen. \"Indicators of Vulnerability and Adaptive Capacity.\" Global Environmental Change, vol. 21, no. 1, 2011, pp. 198-208.",
        "\"Suite of Food Security Indicators (Item 210041).\" FAOSTAT. Food and Agriculture Organization, www.fao.org/faostat/en/#data/FS. Accessed 10 June 2026.",
        "\"ND-GAIN Country Index Data.\" Notre Dame Global Adaptation Initiative, gain.nd.edu/our-work/country-index/download-data/. Accessed 10 June 2026.",
        "\"World Development Indicators.\" World Bank, databank.worldbank.org. Accessed 10 June 2026.",
        "\"EM-DAT: The International Disaster Database.\" Centre for Research on the Epidemiology of Disasters, www.emdat.be. Accessed 10 June 2026.",
        "\"Conflict Event and Fatality Data.\" Armed Conflict Location & Event Data Project, acleddata.com. Accessed 10 June 2026.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    doc.add_page_break()
    h1(doc, "Part 8 — Acknowledgements (~40–80 words)")
    bullets(doc, [
        "Thank people who read drafts but are NOT co-authors",
        "State funding source OR 'This research received no external funding'",
        "Do NOT thank mentor here if mentor is a co-author",
        "Keep brief — 2–4 sentences maximum",
    ])

    h1(doc, "Part 9 — Glossary")
    add_table(doc, ["Term", "First-use definition", "Symbol"], [
        ["Prevalence of undernourishment (PoU)", "% population below minimum dietary energy requirement", "U"],
        ["ND-GAIN vulnerability", "Composite climate vulnerability, 0–1, higher = more vulnerable", "V"],
        ["ND-GAIN readiness", "National capacity to deploy adaptation investment", "—"],
        ["Adaptation buffer", "Predicted PoU minus actual PoU from vulnerability-only regression", "B"],
        ["Over-performer", "Country with positive buffer", "—"],
        ["Under-performer", "Country with negative buffer", "—"],
        ["Collapse", "Buffer falls below zero at least once in simulation", "—"],
        ["Percentage points (pp)", "Absolute difference on % scale; do not say 'percent' when you mean pp", "—"],
    ])

    h1(doc, "Part 10 — Master Numbers Cheat Sheet")
    add_table(doc, ["Category", "Item", "Value"], [
        ["Sample", "N bivariate / MC", "143"],
        ["Sample", "N multivariate", "139"],
        ["Exp 1", "Intercept", "−20.04 (SE 2.59)"],
        ["Exp 1", "Vulnerability β", "69.42 (SE 7.12), p<0.001"],
        ["Exp 1", "R²", "0.423"],
        ["Exp 1", "0.1 V → PoU", "+6.9 pp"],
        ["Exp 2", "Vulnerability β", "40.07 (SE 11.48)"],
        ["Exp 2", "ln(GDP pc)", "−2.84 (p=0.006)"],
        ["Exp 2", "R²", "0.464"],
        ["Buffer", "SD", "7.5 pp"],
        ["Buffer", "Range", "+14.6 to −38.8"],
        ["Bangladesh", "V, pred, actual, buffer", "0.569, 19.4%, 10.4%, +9.0"],
        ["Exp 4", "R²", "0.057"],
        ["Exp 4", "Readiness β, p", "13.35, p=0.12"],
        ["MC", "Draws", "20,000"],
        ["MC", "Slope 90% CI", "[58.3, 81.5]"],
        ["MC", "Over / ambiguous / under", "66 / 35 / 42"],
        ["MC", "Bangladesh P(over)", "1.00"],
        ["MC", "Haiti P(over)", "0.00"],
        ["Collapse", "P(20y) baseline vs resilience", "41.3% vs 4.4%"],
        ["Collapse", "P(10y) baseline vs resilience", "21.6% vs 2.0%"],
    ])

    h1(doc, "Part 11 — Recommended Writing Order & Time Budget")
    add_table(doc, ["Step", "Task", "Hours"], [
        ["1", "Download JEI template from emerginginvestigators.org/documents", "0.25"],
        ["2", "Write Materials and Methods", "2–3"],
        ["3", "Write Results + build Tables 1–4 and Figures 1–3", "4–5"],
        ["4", "Write Introduction using Part 3 bullets", "2–3"],
        ["5", "Write Discussion using Part 5 bullets", "2"],
        ["6", "Write Summary last; trim to <250 words", "1"],
        ["7", "Title page, Acknowledgements, renumber citations", "1–2"],
        ["8", "Paste into JEI template; verify ≤10 pages", "1"],
        ["9", "Mentor review + submit via Editorial Manager", "—"],
    ])

    h1(doc, "Part 12 — Editorial Manager Submission Fields")
    add_table(doc, ["Field", "Suggested answer"], [
        ["One-sentence hypothesis", "Countries with higher ND-GAIN readiness will show larger positive adaptation buffers relative to climate-predicted undernourishment."],
        ["How did you hear about JEI?", "[Your honest answer]"],
        ["Why interested?", "Puzzle of Bangladesh vs Afghanistan — similar vulnerability, very different hunger"],
        ["Interesting finding?", "Niger is among the most vulnerable countries but still overperforms; Haiti is the worst underperformer despite moderate vulnerability"],
        ["Difficulties?", "Merging messy World Bank data; explaining Monte Carlo simulation plainly; staying under 10 pages"],
        ["Institution", "Your high school name"],
    ])

    h1(doc, "Part 13 — What to Cut if Over 10 Pages (priority order)")
    add_table(doc, ["Priority", "Cut", "Saves ~"], [
        ["1", "Shorten robustness to 2 sentences", "80 words"],
        ["2", "Compress Introduction ¶3 literature", "80 words"],
        ["3", "Delete personal motivation ¶5", "70 words"],
        ["4", "Fold Table 2 into text (keep key coefs)", "¼ page"],
        ["5", "Reduce Table 3 to 6 countries", "¼ page"],
        ["NEVER CUT", "Buffer definition, 66/42 MC split, 41% vs 4.4% collapse, Bangladesh example", "—"],
    ])

    h1(doc, "Part 14 — Anticipated Reviewer Concerns & Responses")
    add_table(doc, ["Concern", "Where to address", "Response strategy"], [
        ["Not causal", "Discussion limitations", "State observational design upfront; say 'associated with'"],
        ["Why this vulnerability index?", "Intro + robustness", "ND-GAIN is standard; GDL replicates (R²=0.408)"],
        ["PoU is modeled/smoothed", "Methods + limitations + R3", "Acknowledge; stunting gradient even stronger"],
        ["Monte Carlo assumptions arbitrary", "Methods + limitations", "State assumptions; emphasize comparison not exact %"],
        ["Ecological fallacy (country-level)", "Discussion future work", "Descriptive global patterns; propose panel tracking"],
        ["Readiness hypothesis failed", "Results Exp 4 honestly", "p=0.12, hypothesis-generating; R²=0.057 is the finding"],
        ["Small island anomalies", "Results Exp 3 + R5", "Flag Kiribati; drop in R5; Bangladesh rises to rank ~6"],
    ])

    h1(doc, "Part 15 — Pre-Submission Checklist")
    bullets(doc, [
        "[ ] Summary ≤ 250 words and includes hypothesis",
        "[ ] Every figure cited ≥ once in Results (passive format)",
        "[ ] Every table cited ≥ once in Results",
        "[ ] All Discussion claims appear in Results first",
        "[ ] JEI official Word template used",
        "[ ] 11 pt Times New Roman, 1.5 spacing, 1-inch margins",
        "[ ] Body ≤ 10 pages (excl. title, references, figure captions)",
        "[ ] 3 figures uploaded as separate JPEG files",
        "[ ] Mentor listed last; mentor submits (not student)",
        "[ ] Submission fee paid; confirmation code saved",
        "[ ] Manuscript not submitted elsewhere (no conflicting preprint)",
        "[ ] Only one manuscript at JEI at a time",
        "[ ] References numbered in order of appearance; no hanging indents",
    ])


def build_expanded_document(doc: Document) -> None:
    add_title_page(doc)
    add_formatting_guide(doc)
    add_master_plan(doc)
    add_title_page_section(doc)
    add_summary_section(doc)
    add_introduction_section(doc)
    add_results_section(doc)
    add_discussion_section(doc)
    add_methods_section(doc)
    add_back_matter(doc)
