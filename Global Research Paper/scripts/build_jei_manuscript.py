#!/usr/bin/env python3
"""
Build a JEI-formatted Word manuscript and companion figure files.

JEI requirements implemented:
  - Section order: Title Page, Summary, Introduction, Results, Discussion,
    Materials and Methods, References, Acknowledgements
  - Times New Roman 11 pt, 1.5 line spacing, 1-inch margins
  - Numbered in-text citations (1), (2), ... matching numbered reference list
  - Tables in the Word document; figures saved separately as JPEG
  - Title <= 110 characters; Summary <= 250 words

Download the official JEI template from https://emerginginvestigators.org/documents
and paste this content into it before submission if editors require the template file.

Run from Global Research Paper/:
  python3 scripts/build_jei_manuscript.py
"""

from __future__ import annotations

import sys
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

PAPER = Path(__file__).resolve().parents[1]
OUT_DOC = PAPER / "global_paper_JEI.docx"
OUT_FIG = PAPER / "figures"
sys.path.insert(0, str(PAPER / "scripts"))

from global_models_1_to_4 import cross_section_2022, load_world_bank_backbone  # noqa: E402

# ---------------------------------------------------------------------------
# Manuscript text (JEI section order). Edit author placeholders before submitting.
# ---------------------------------------------------------------------------

TITLE = "Climate Vulnerability Does Not Equal Hunger: Measuring the Global Adaptation Buffer"

AUTHORS = [
    ("[Student Firstname Lastname]", "1"),
    ("[Mentor Firstname Lastname]", "2"),
]
AFFILIATIONS = [
    "1 [High School Name], [City], [State]",
    "2 [High School Name], [City], [State]",
]

SUMMARY = (
    "Climate vulnerability is widely used as a proxy for hunger risk, yet countries with "
    "similar vulnerability scores report very different rates of undernourishment. We introduce "
    "the adaptation buffer: the gap between hunger predicted from a cross-country regression on "
    "climate vulnerability and actual hunger. Using ND-GAIN, FAO, World Bank, EM-DAT, and ACLED "
    "data for 143 countries, we tested whether vulnerability predicts undernourishment and whether "
    "standard development indicators explain who beats that prediction. We hypothesized that "
    "countries with stronger adaptive capacity would show larger positive buffers. Vulnerability "
    "alone explained 42% of cross-country hunger variation (p < 0.001), and the association "
    "survived controls for income, rurality, disasters, and conflict. The buffer ranged from "
    "+14.6 to -38.8 percentage points; Bangladesh (+9.1) and Senegal (+12.1) overperformed while "
    "Haiti (-38.8) underperformed. Standard covariates explained only 6% of buffer variation. "
    "Monte Carlo simulation (20,000 draws) showed 66 countries were robust over-performers and 42 "
    "robust under-performers. Simulating buffer dynamics forward, a +9-point buffer collapsed below "
    "zero within 20 years in 41% of trials under severe shocks but only 4% once resilience "
    "investments reduced shock damage and sped recovery. The buffer is a measurable outcome-based "
    "indicator of effective adaptation that aid frameworks could track alongside vulnerability."
)

INTRODUCTION = [
    (
        "Bangladesh and Afghanistan have similar ND-GAIN climate-vulnerability scores (0.57 and "
        "0.59 on a 0-1 scale), yet Afghanistan's prevalence of undernourishment is 28.1% while "
        "Bangladesh's is 10.4% (11). If vulnerability mapped directly onto hunger, these countries "
        "should look alike. They do not."
    ),
    (
        "Climate change threatens food security through yield losses, extreme events, and disrupted "
        "livelihoods (1,2). The FAO reports climate extremes as a leading driver of recent global "
        "hunger increases alongside conflict and economic downturns (3). Countries are often ranked "
        "by vulnerability and those rankings are read as hunger-risk rankings, but this coupling is "
        "rarely tested globally."
    ),
    (
        "The ND-GAIN Country Index is the most widely used composite measure of national climate "
        "vulnerability, paired with a readiness index for adaptive capacity (4). Tracking studies "
        "show a persistent gap between adaptation planning and documented outcomes; outcome-based "
        "evidence of risk reduction remains rare (5,6). The FAO prevalence of undernourishment "
        "(PoU) is the standard SDG hunger indicator, estimated from food balance sheets and "
        "consumption distributions (7). Income is the dominant cross-country correlate of hunger "
        "(8), while conflict has become a leading driver of acute food crises (3,9)."
    ),
    (
        "We define the adaptation buffer as predicted undernourishment minus actual "
        "undernourishment, where prediction comes from a bivariate regression on vulnerability alone. "
        "A positive buffer means a country reports less hunger than its climate burden predicts. "
        "We asked: (1) How strongly does vulnerability predict hunger after controls? (2) Which "
        "countries beat or fall short of the climate-only benchmark? (3) Can standard covariates "
        "explain the buffer, and how uncertain are country rankings? (4) How easily does a buffer "
        "collapse under random shocks, and what protects it? We hypothesized that higher ND-GAIN "
        "readiness would correlate with larger buffers. Using ordinary least squares regression "
        "and Monte Carlo simulation on 143 countries, we found vulnerability explains much but not "
        "all hunger variation, buffers are large and structured, rankings are statistically robust, "
        "and resilience investments can sharply reduce collapse risk."
    ),
]

RESULTS = [
    (
        "Experiment 1: Vulnerability predicts hunger, but incompletely. We regressed FAO "
        "undernourishment on ND-GAIN vulnerability for 143 countries. A 0.1-point increase in "
        "vulnerability was associated with 6.9 percentage points higher undernourishment "
        "(p < 0.001). Vulnerability alone explained 42.3% of cross-country variance (R2 = 0.423) "
        "(Table 1, Figure 1). Bangladesh (V = 0.569) had predicted PoU of 19.4% but actual PoU of "
        "10.4%, giving a buffer of +9.0 percentage points."
    ),
    (
        "Experiment 2: The climate-hunger link survives controls. Adding log GDP per capita, "
        "rural population share, disaster counts (2013-2022), and conflict fatalities (N = 139), "
        "the vulnerability coefficient fell from 69.4 to 40.1 but remained significant "
        "(p < 0.001). Higher income was associated with lower hunger (p = 0.006). Disasters and "
        "conflict aggregates were not significant in this linear specification (Table 2)."
    ),
    (
        "Experiment 3: Global buffer rankings. Computing the buffer for all 143 countries yielded "
        "a distribution with standard deviation 7.5 percentage points, ranging from +14.6 (Kiribati) "
        "to -38.8 (Haiti). Top overperformers included Senegal (+12.1), Niger (+11.2), and "
        "Bangladesh (+9.0). Largest underperformers were Haiti (-38.8), Syria (-25.2), and Kenya "
        "(-22.1). Niger, among the world's most vulnerable countries, had a positive buffer, while "
        "Botswana, moderately vulnerable, had a negative buffer (-14.1). Underperformance clustered "
        "in conflict-affected states; overperformance clustered in agrarian countries with sustained "
        "agricultural and social-protection investment (Table 3, Figure 2)."
    ),
    (
        "Experiment 4: Standard covariates barely explain the buffer. Regressing the buffer on "
        "log GDP per capita, ND-GAIN readiness, rural share, disasters, and conflict fatalities "
        "(N = 139) gave R2 = 0.057. Readiness had the largest coefficient (+13.4 pp per 1.0 index "
        "point) but was not significant at the 5% level (p = 0.12). This suggests the buffer "
        "captures adaptation outcomes not reducible to standard development indicators."
    ),
    (
        "Experiment 5: Monte Carlo uncertainty. We resampled the 143 countries with replacement "
        "and added measurement noise to reported PoU (10% coefficient of variation, 1 pp floor), "
        "repeating 20,000 times. The vulnerability slope averaged 69.6 with a 90% interval of "
        "[58.3, 81.5]. Sixty-six countries had P(buffer > 0) >= 0.90 (robust over-performers) and "
        "42 had P <= 0.10 (robust under-performers). Bangladesh's buffer was positive in 100% of "
        "simulations (90% interval: +6.1 to +12.1 pp); Haiti's was negative in 100% (Table 3)."
    ),
    (
        "Experiment 6: Buffer collapse dynamics. We simulated a Bangladesh-like buffer (+9.1) "
        "forward 20 years, 20,000 times. Each year the buffer recovered toward its target and "
        "occasionally took a random shock (20% annual probability). Under baseline conditions "
        "(mean shock 4.0 pp, slow recovery), the buffer fell below zero at least once within "
        "20 years in 41.3% of trials. With resilience investments that reduced shock damage by "
        "40% and sped recovery, collapse probability fell to 4.4% for the same climate exposure "
        "(Table 4, Figure 3)."
    ),
    (
        "Robustness checks confirmed the headline gradient was stable to alternative FAO windows, "
        "alternative vulnerability indices that embed susceptibility (Global Data Lab, WRI "
        "vulnerability component), exclusion of small island states, and use of child stunting as "
        "the outcome (R2 = 0.62). Pure exposure-only indices showed much weaker associations."
    ),
]

DISCUSSION = [
    (
        "Climate vulnerability does not equal hunger. Vulnerability explains 42% of cross-country "
        "variation in undernourishment, but the remaining variation is structured and policy-relevant. "
        "The adaptation buffer makes that residual visible: it benchmarks actual hunger against what "
        "climate exposure alone would predict. Because standard covariates explain only 6% of buffer "
        "variation, the buffer must be monitored directly rather than inferred from income or "
        "readiness alone."
    ),
    (
        "Conflict destroys buffers faster than climate exposure alone would suggest. Haiti, Syria, "
        "Liberia, and the Central African Republic are among the largest underperformers despite "
        "not being the most climate-vulnerable. Conversely, Niger, Bangladesh, and Senegal hold large "
        "positive buffers. Vulnerability-only aid formulas would misdirect resources away from "
        "countries where hunger most exceeds climate fundamentals."
    ),
    (
        "Buffers are assets that can be built and lost. Bangladesh's buffer partially collapsed after "
        "the 2007-2009 cyclone sequence before partial rebuilding. Our collapse simulation showed a "
        "+9-point buffer has a 41% chance of falling below zero within 20 years under severe shocks, "
        "but only 4% once resilience investments soften shocks and speed recovery. The policy payoff "
        "lies in absorptive capacity (early warning, storage, stress-tolerant crops, safety nets) "
        "rather than in lowering hazard scores alone."
    ),
    (
        "Limitations. All analyses are observational, not causal. The buffer is defined relative to "
        "a specific vulnerability index and bivariate benchmark; alternative indices yield different "
        "buffers, though robustness checks with GDL and WRI vulnerability components replicated the "
        "gradient. FAO PoU is a modeled three-year average and may smooth short-run shocks. Monte "
        "Carlo measurement-error and shock parameters are stated assumptions; their value is in "
        "transparent uncertainty quantification and mechanism comparison, not precise country-level "
        "forecasting. Conflict and climate stress interact in ways our linear covariates do not fully "
        "capture."
    ),
    (
        "Future work could extend the buffer to annual panel tracking, enrich buffer-determinant "
        "models with agricultural productivity and aid-flow data, and test whether shrinking buffers "
        "predict subsequent hunger increases. Integrating buffer level and trend into adaptation-finance "
        "criteria could help target countries whose fragile gains are eroding before crises fully "
        "materialize."
    ),
]

METHODS = [
    (
        "Data. We merged publicly available country-level datasets by ISO3 code for 143 countries "
        "with non-missing ND-GAIN vulnerability (2022) and FAO prevalence of undernourishment "
        "(Item 210041, 2022-2024 three-year window; values reported as <2.5% coded at 2.5%) (7,11,12). "
        "ND-GAIN composite vulnerability and readiness indices came from the Notre Dame Global "
        "Adaptation Initiative (4,12). World Bank World Development Indicators provided 2022 GDP per "
        "capita and rural population share (13). EM-DAT disaster counts were summed over 2013-2022 "
        "(14). ACLED provided three-year average conflict fatalities (15). Complete cases for "
        "multivariate models: N = 139."
    ),
    (
        "Adaptation buffer. From the bivariate regression of undernourishment on vulnerability "
        "(Experiment 1), we computed fitted values and defined buffer = predicted PoU - actual PoU "
        "for each country. Positive values indicate outperformance relative to the climate-only "
        "benchmark."
    ),
    (
        "Regression analysis. We used ordinary least squares (OLS) regression for Experiments 1-4, "
        "implemented in Python 3 with statsmodels. Experiment 1: undernourishment ~ vulnerability. "
        "Experiment 2: added log GDP per capita, rural share, disaster count, and log(1 + conflict "
        "fatalities). Experiment 4: buffer ~ log GDP per capita + readiness + rural share + disasters "
        "+ log(1 + fatalities); vulnerability was excluded because it enters the buffer construction "
        "through the predicted value."
    ),
    (
        "Monte Carlo simulation. For Experiment 5, we repeated 20,000 iterations: (1) bootstrap "
        "resample of 143 countries with replacement and refit the Experiment 1 regression; (2) add "
        "normally distributed measurement noise to each country's PoU with standard deviation "
        "max(1.0, 0.10 x PoU); (3) recompute all buffers. We report 90% intervals and P(buffer > 0)."
    ),
    (
        "Collapse simulation. For Experiment 6, we simulated buffer paths over 20 years (20,000 "
        "trials) with annual recovery toward a target, random shocks (20% annual probability, gamma-"
        "distributed size), and Gaussian noise. We compared a baseline scenario (mean shock 4.0 pp, "
        "recovery rate 0.25) to a resilience-investment scenario (mean shock 2.4 pp, recovery rate "
        "0.40), starting from Bangladesh's estimated buffer (+9.1). Code and data processing scripts "
        "are available in the project repository."
    ),
    (
        "Statistical analysis. We report coefficients, standard errors, p-values, and R2. Significance "
        "was assessed at alpha = 0.05. No human subjects or vertebrate animals were used."
    ),
]

REFERENCES = [
    'Wheeler, Tim, and Joachim von Braun. "Climate Change Impacts on Global Food Security." Science, vol. 341, no. 6145, 2013, pp. 508-513. https://doi.org/10.1126/science.1239402.',
    'IPCC. Climate Change 2022: Impacts, Adaptation and Vulnerability. Contribution of Working Group II to the Sixth Assessment Report of the Intergovernmental Panel on Climate Change. Cambridge University Press, 2022.',
    'FAO, IFAD, UNICEF, WFP, and WHO. The State of Food Security and Nutrition in the World 2024. Food and Agriculture Organization of the United Nations, 2024.',
    'Chen, C., et al. University of Notre Dame Global Adaptation Index: Country Index Technical Report. University of Notre Dame, 2015.',
    'Ford, James D., et al. "How to Track Adaptation to Climate Change: A Typology of Approaches for National-Level Application." Ecology and Society, vol. 18, no. 3, 2013, https://doi.org/10.5751/ES-05732-180340.',
    'Berrang-Ford, Lea, et al. "A Systematic Global Stocktake of Evidence on Human Adaptation to Climate Change." Nature Climate Change, vol. 11, no. 11, 2021, pp. 989-1000. https://doi.org/10.1038/s41558-021-01170-y.',
    'Cafiero, Carlo, Sara Viviani, and Mark Nord. "Food Security Measurement in a Global Context: The Food Insecurity Experience Scale." Measurement, vol. 116, 2018, pp. 146-152. https://doi.org/10.1016/j.measurement.2017.10.065.',
    'Headey, Derek. "Developmental Drivers of Nutritional Change: A Cross-Country Analysis." World Development, vol. 42, 2013, pp. 76-88. https://doi.org/10.1016/j.worlddev.2012.07.002.',
    'Burke, Marshall, Solomon M. Hsiang, and Edward Miguel. "Climate and Conflict." Annual Review of Economics, vol. 7, 2015, pp. 577-617. https://doi.org/10.1146/annurev-economics-080614-115430.',
    'Hinkel, Jochen. "Indicators of Vulnerability and Adaptive Capacity: Towards a Clarification of the Science-Policy Interface." Global Environmental Change, vol. 21, no. 1, 2011, pp. 198-208. https://doi.org/10.1016/j.gloenvcha.2010.08.002.',
    '"Suite of Food Security Indicators (Item 210041: Prevalence of Undernourishment)." FAOSTAT. Food and Agriculture Organization of the United Nations, www.fao.org/faostat/en/#data/FS. Accessed 10 June 2026.',
    '"ND-GAIN Country Index Data." Notre Dame Global Adaptation Initiative, gain.nd.edu/our-work/country-index/download-data/. Accessed 10 June 2026.',
    '"World Development Indicators." World Bank, databank.worldbank.org. Accessed 10 June 2026.',
    '"EM-DAT: The International Disaster Database." Centre for Research on the Epidemiology of Disasters, www.emdat.be. Accessed 10 June 2026.',
    '"Conflict Event and Fatality Data." Armed Conflict Location & Event Data Project, acleddata.com. Accessed 10 June 2026.',
]

ACKNOWLEDGEMENTS = (
    "[Acknowledge anyone who read and commented on the manuscript before submission. "
    "State funding sources if any. Do not list the mentor here if they are a co-author.]"
)

TABLE1 = [
    ["Term", "Coefficient", "Std. error", "p-value"],
    ["Constant", "-20.04", "2.59", "<0.001"],
    ["Vulnerability", "69.42", "7.12", "<0.001"],
    ["R-squared", "0.423", "(N = 143)", ""],
]

TABLE2 = [
    ["Variable", "Coefficient", "Std. error", "p-value"],
    ["Vulnerability", "40.07", "11.48", "<0.001"],
    ["ln(GDP per capita)", "-2.84", "1.04", "0.006"],
    ["Rural share (%)", "-0.047", "0.045", "0.29"],
    ["Disaster count (2013-22)", "0.12", "0.25", "0.63"],
    ["ln(1 + conflict fatalities)", "0.065", "0.31", "0.83"],
    ["R-squared", "0.464", "(N = 139)", ""],
]

TABLE3 = [
    ["Country", "Buffer (pp)", "90% interval", "P(over-performer)"],
    ["Senegal", "+12.1", "[+9.5, +14.9]", "1.00"],
    ["Niger", "+11.2", "[+7.4, +15.2]", "1.00"],
    ["Bangladesh", "+9.1", "[+6.1, +12.1]", "1.00"],
    ["Viet Nam", "+7.1", "[+5.0, +9.3]", "1.00"],
    ["Kenya", "-22.1", "[-28.4, -15.7]", "0.00"],
    ["Haiti", "-38.8", "[-47.9, -29.6]", "0.00"],
]

TABLE4 = [
    ["Scenario", "P(collapse <= 20 y)", "Expected years underwater", "Buffer yr 20 (median, pp)"],
    ["Baseline (large shocks, slow recovery)", "41.3%", "1.01", "+6.6"],
    ["Resilience investment", "4.4%", "0.05", "+8.2"],
]

FIGURE_CAPTIONS = {
    1: (
        "Figure 1: Relationship between ND-GAIN climate vulnerability and FAO prevalence of "
        "undernourishment across 143 countries (N = 143). Each point is one country. The line shows "
        "the ordinary least squares fit (R2 = 0.423). Bangladesh is labeled as an example of a "
        "positive adaptation buffer (actual hunger below the prediction)."
    ),
    2: (
        "Figure 2: Adaptation buffer for selected countries. The buffer equals predicted minus actual "
        "undernourishment (percentage points). Positive values indicate countries outperforming their "
        "climate-only benchmark. Data from 143-country cross-section, 2022."
    ),
    3: (
        "Figure 3: Monte Carlo simulation of buffer collapse over 20 years (N = 20,000 trials). "
        "Starting buffer +9.1 percentage points (Bangladesh). Both scenarios face the same 20% annual "
        "shock probability; resilience investment reduces mean shock size and increases recovery rate."
    ),
}


def set_jei_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.size = Pt(12 if level == 1 else 11)
        h.font.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for p in paragraphs:
        doc.add_paragraph(p)


def add_table(doc: Document, caption: str, rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.bold = True
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = val
            if i == 0:
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.bold = True
    doc.add_paragraph()


def add_figure_placeholder(doc: Document, num: int) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[Figure {num} submitted as separate JPEG file: figures/Figure{num}.jpeg]")
    run.italic = True
    cap = doc.add_paragraph()
    cap_run = cap.add_run(FIGURE_CAPTIONS[num])
    cap_run.bold = True


def build_figures() -> None:
    OUT_FIG.mkdir(parents=True, exist_ok=True)
    wb, backbone = load_world_bank_backbone()
    cs = cross_section_2022(wb, backbone).dropna(subset=["undernourishment", "vulnerability"])
    x = cs["vulnerability"].to_numpy()
    y = cs["undernourishment"].to_numpy()
    b0, b1 = np.polyfit(x, y, 1)
    line_x = np.linspace(x.min(), x.max(), 100)
    line_y = b0 + b1 * line_x

    # Figure 1
    fig, ax = plt.subplots(figsize=(7, 5))
    ax.scatter(x, y, alpha=0.6, s=35, color="#2c5f8a", edgecolors="white", linewidth=0.3)
    ax.plot(line_x, line_y, color="#c44e52", linewidth=2, label=f"OLS fit (R² = 0.423)")
    bgd = cs[cs["iso3"] == "BGD"]
    if len(bgd):
        ax.annotate("Bangladesh", (bgd["vulnerability"].iloc[0], bgd["undernourishment"].iloc[0]),
                    textcoords="offset points", xytext=(8, 8), fontsize=11, fontweight="bold")
    ax.set_xlabel("ND-GAIN vulnerability (0-1)", fontsize=14, fontweight="bold")
    ax.set_ylabel("Prevalence of undernourishment (%)", fontsize=14, fontweight="bold")
    ax.legend(fontsize=12)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    fig.savefig(OUT_FIG / "Figure1.jpeg", dpi=300, bbox_inches="tight")
    plt.close(fig)

    # Figure 2
    mc = pd.read_csv(PAPER / "output" / "global_mc_buffer_uncertainty.csv")
    pick = ["Senegal", "Niger", "Bangladesh", "Viet Nam", "Kenya", "Haiti", "Botswana", "Madagascar"]
    sub = mc[mc["country"].isin(pick)].copy()
    sub = sub.set_index("country").loc[pick].reset_index()
    colors = ["#2a9d8f" if v > 0 else "#e76f51" for v in sub["buffer_mc_mean"]]
    fig, ax = plt.subplots(figsize=(8, 5))
    ax.barh(sub["country"], sub["buffer_mc_mean"], color=colors)
    ax.axvline(0, color="black", linewidth=0.8)
    ax.set_xlabel("Adaptation buffer (percentage points)", fontsize=14, fontweight="bold")
    ax.set_ylabel("")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    fig.savefig(OUT_FIG / "Figure2.jpeg", dpi=300, bbox_inches="tight")
    plt.close(fig)

    # Figure 3
    coll = pd.read_csv(PAPER / "output" / "global_mc_collapse_dynamics.csv")
    labels = ["Baseline", "Resilience\ninvestment"]
    vals = coll["P_collapse_within_20y"].to_numpy() * 100
    fig, ax = plt.subplots(figsize=(6, 5))
    bars = ax.bar(labels, vals, color=["#e76f51", "#2a9d8f"], width=0.55)
    ax.set_ylabel("P(buffer falls below 0 within 20 years, %)", fontsize=13, fontweight="bold")
    ax.set_ylim(0, max(vals) * 1.25)
    for bar, v in zip(bars, vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1, f"{v:.1f}%",
                ha="center", fontsize=13, fontweight="bold")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    fig.tight_layout()
    fig.savefig(OUT_FIG / "Figure3.jpeg", dpi=300, bbox_inches="tight")
    plt.close(fig)


def build_docx() -> None:
    doc = Document()
    set_jei_styles(doc)

    # Title page
    t = doc.add_paragraph()
    t.alignment = 1  # center
    tr = t.add_run(TITLE)
    tr.bold = True
    tr.font.size = Pt(12)
    doc.add_paragraph()
    for name, aff in AUTHORS:
        p = doc.add_paragraph()
        p.add_run(f"{name}{aff}")
    for aff in AFFILIATIONS:
        doc.add_paragraph(aff)
    doc.add_page_break()

    # Summary
    add_heading(doc, "Summary")
    doc.add_paragraph(SUMMARY)
    word_count = len(SUMMARY.split())
    note = doc.add_paragraph()
    note.add_run(f"[Summary word count: {word_count} — JEI limit is 250 words]").italic = True
    doc.add_page_break()

    # Introduction
    add_heading(doc, "Introduction")
    add_paragraphs(doc, INTRODUCTION)

    # Results
    add_heading(doc, "Results")
    add_paragraphs(doc, RESULTS[:1])
    add_table(doc, "Table 1: Experiment 1 regression of undernourishment on vulnerability.", TABLE1)
    add_figure_placeholder(doc, 1)
    add_paragraphs(doc, RESULTS[1:2])
    add_table(doc, "Table 2: Experiment 2 multivariate regression.", TABLE2)
    add_paragraphs(doc, RESULTS[2:3])
    add_table(doc, "Table 3: Monte Carlo buffer estimates for selected countries (Experiment 5).", TABLE3)
    add_figure_placeholder(doc, 2)
    add_paragraphs(doc, RESULTS[3:5])
    add_paragraphs(doc, RESULTS[5:6])
    add_table(doc, "Table 4: Buffer collapse simulation results (Experiment 6).", TABLE4)
    add_figure_placeholder(doc, 3)
    add_paragraphs(doc, RESULTS[6:])

    # Discussion
    add_heading(doc, "Discussion")
    add_paragraphs(doc, DISCUSSION)

    # Materials and Methods
    add_heading(doc, "Materials and Methods")
    add_paragraphs(doc, METHODS)

    # References
    add_heading(doc, "References")
    for i, ref in enumerate(REFERENCES, start=1):
        doc.add_paragraph(f"{i}. {ref}")

    # Acknowledgements
    add_heading(doc, "Acknowledgements")
    doc.add_paragraph(ACKNOWLEDGEMENTS)

    doc.save(OUT_DOC)


def main() -> int:
    assert len(TITLE) <= 110, f"Title too long: {len(TITLE)} chars"
    build_figures()
    build_docx()
    print(f"Written: {OUT_DOC}")
    print(f"Figures: {OUT_FIG}/Figure1.jpeg, Figure2.jpeg, Figure3.jpeg")
    print(f"Title length: {len(TITLE)} characters (max 110)")
    print(f"Summary length: {len(SUMMARY.split())} words (max 250)")
    print("\nBefore submitting:")
    print("  1. Download the official JEI template from https://emerginginvestigators.org/documents")
    print("  2. Fill in author names and affiliations on the title page")
    print("  3. Paste content into the template OR submit this file if your mentor confirms format")
    print("  4. Upload figures separately as JPEG files in Editorial Manager")
    return 0


if __name__ == "__main__":
    sys.exit(main())
