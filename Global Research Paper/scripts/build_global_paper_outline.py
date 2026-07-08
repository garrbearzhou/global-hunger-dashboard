#!/usr/bin/env python3
"""Generate Global_Paper_Outline.docx (and PDF via fpdf2)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches, Pt

from global_paper_outline_sections import build_expanded_document

PAPER = Path(__file__).resolve().parents[1]
OUT_DOCX = PAPER / "Global_Paper_Outline.docx"
OUT_PDF = PAPER / "Global_Paper_Outline.pdf"


def set_styles(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(4)
    for level in range(1, 4):
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Times New Roman"
        h.font.bold = True
        h.font.size = Pt(14 if level == 1 else 12 if level == 2 else 11)


def build() -> Document:
    doc = Document()
    set_styles(doc)
    build_expanded_document(doc)
    return doc


def try_pdf(docx_path: Path, pdf_path: Path) -> bool:
    import textwrap

    from docx import Document as DocxDocument
    from fpdf import FPDF

    try:
        doc = DocxDocument(str(docx_path))
        pdf = FPDF(format="Letter")
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_margins(15, 15, 15)
        pdf.add_page()

        def sanitize(text: str) -> str:
            replacements = {
                "—": "-", "−": "-", "→": "->", "≥": ">=", "≤": "<=",
                "≈": "~", "ρ": "rho", "β": "beta", "²": "2", "×": "x", "¶": "P",
            }
            for a, b in replacements.items():
                text = text.replace(a, b)
            return text.encode("latin-1", errors="replace").decode("latin-1")

        def write_wrapped(text: str, size: int = 10, bold: bool = False) -> None:
            pdf.set_font("Helvetica", "B" if bold else "", size)
            for line in textwrap.wrap(sanitize(text), width=95):
                pdf.cell(0, 5, line, new_x="LMARGIN", new_y="NEXT")

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                pdf.ln(2)
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style:
                pdf.ln(4)
                write_wrapped(text, 14, True)
            elif "Heading 2" in style:
                pdf.ln(2)
                write_wrapped(text, 12, True)
            elif "Heading 3" in style:
                write_wrapped(text, 11, True)
            elif "List Bullet" in style:
                write_wrapped("- " + text, 10)
            else:
                write_wrapped(text, 10)

        for table in doc.tables:
            pdf.ln(3)
            for ri, row in enumerate(table.rows):
                line = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
                write_wrapped(line, 9, bold=(ri == 0))

        pdf.output(str(pdf_path))
        return pdf_path.exists()
    except Exception as exc:
        print(f"PDF generation failed: {exc}")
        return False


def main() -> int:
    doc = build()
    doc.save(OUT_DOCX)
    print(f"Written: {OUT_DOCX} ({OUT_DOCX.stat().st_size // 1024} KB)")

    if try_pdf(OUT_DOCX, OUT_PDF):
        print(f"Written: {OUT_PDF} ({OUT_PDF.stat().st_size // 1024} KB)")
    else:
        print("PDF not generated. Open the .docx and Export as PDF in Word.")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
